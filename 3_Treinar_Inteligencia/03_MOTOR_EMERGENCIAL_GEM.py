# ==============================================================================
# MOTOR GEM - SMOSU OLIVEIRA/MG - EMERGENCIAL CONSOLIDADO v1.0
# ==============================================================================
# USO: copie para a pasta raiz do projeto e execute: python 03_MOTOR_EMERGENCIAL_GEM.py
# DEPENDENCIAS: pip install python-docx docx2pdf
# ==============================================================================


# ----------------------------------------------------------------------------
# MODULO: config.py
# ----------------------------------------------------------------------------

"""
Configuração central do sistema de geração de documentos — SMOSU Oliveira/MG.
Constantes, cores, fontes, caminhos e mapeamento de tipos.
"""

import os
from docx.shared import RGBColor

# ═══════════════════════════════════════════════════════════
#  CAMINHOS
# ═══════════════════════════════════════════════════════════
SCRIPT_DIR      = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR     = os.path.dirname(SCRIPT_DIR)  # pasta raiz: 02. Pareceres
LOGO_BRASAO     = os.path.join(SCRIPT_DIR, "logos", "logo_brasao.jpeg")
LOGO_PREFEITURA = os.path.join(SCRIPT_DIR, "logos", "logo_prefeitura.png")
TEMPLATES_DIR   = os.path.join(SCRIPT_DIR, "templates")

# ═══════════════════════════════════════════════════════════
#  TIPOGRAFIA
# ═══════════════════════════════════════════════════════════
# • Corpo do texto:   Calibri 11pt  — sans-serif moderna, excelente leitura
# • Títulos/Labels:   Cambria bold   — serifa clássica, peso institucional
# • Cabeçalho:        Cambria (nome) + Calibri (detalhes)
# • Tabelas:          Calibri 8-9pt  — compacta e limpa
FONT_CORPO   = "Calibri"     # sans-serif para texto corrido
FONT_TITULO  = "Cambria"     # serifa para títulos e peso institucional
FONT_HEADER  = "Cambria"     # serifa para nome da prefeitura
FONT_DETALHE = "Calibri"     # sans-serif para contato, notas

# ═══════════════════════════════════════════════════════════
#  TAMANHOS DE FONTE
# ═══════════════════════════════════════════════════════════
SZ_CORPO   = 11     # 11pt conforme solicitado
SZ_TABELA  = 9
SZ_CITACAO = 9
SZ_NOTA    = 8
SZ_RODAPE  = 8

# ═══════════════════════════════════════════════════════════
#  CORES PADRONIZADAS
# ═══════════════════════════════════════════════════════════
COR_INST         = '1F3864'
COR_LABEL_BG     = 'D6DCE4'
COR_BORDA_TABELA = 'C0C0C0'
COR_LABEL_FONT   = RGBColor(0x1F, 0x38, 0x64)
COR_CINZA_TEXTO  = RGBColor(0x44, 0x44, 0x44)
COR_CINZA_LEVE   = RGBColor(0x66, 0x66, 0x66)

# ═══════════════════════════════════════════════════════════
#  ESPAÇAMENTO PADRÃO
# ═══════════════════════════════════════════════════════════
PAR_AFTER = 120
LINE_SPC  = 276   # 1.15 linhas (276 twips)

# ═══════════════════════════════════════════════════════════
#  ASSINANTE PADRÃO
# ═══════════════════════════════════════════════════════════
ASSINANTE = {
    "nome":     "Diego Tarcísio Nunes Vilela",
    "titulo":   "Engenheiro Civil",
    "registro": "CREA 235.474/D",
}
CIDADE = "Oliveira"

# ═══════════════════════════════════════════════════════════
#  MAPEAMENTO tipo_relatorio → categoria de gerador
# ═══════════════════════════════════════════════════════════
TIPOS_DOCUMENTO = {
    # ── Pareceres Técnicos (completos, com dados do carimbo) ──
    "alvara_aprovacao":                   "parecer_tecnico",
    "alvara_regularizacao":               "parecer_tecnico",
    "alvara_ampliacao":                   "parecer_tecnico",
    "alvara_galpao_comercial":            "parecer_tecnico",
    "alvara_reforma_demolicao_ampliacao": "parecer_tecnico",
    "alvara_substituicao_projeto":        "parecer_tecnico",
    "regularizacao":                      "parecer_tecnico",  # compatibilidade

    # ── Pareceres Simples (sem dados do carimbo) ──
    "certidao_numero_2via":               "parecer_simples",
    "certidao_nome_rua":                  "parecer_simples",
    "certidao_localizacao":               "parecer_simples",
    "certidao_conjunta":                  "parecer_simples",
    "certidao_numero_comercial":          "parecer_simples",
    "habitese_comum":                     "parecer_simples",
    "habitese_multa":                     "parecer_simples",
    "certidao_averbacao_decadencia":      "parecer_simples",
    "habitese_2via":                      "parecer_simples",
    "habitese_inclusao_area":             "parecer_simples",
    "alvara_renovacao":                   "parecer_simples",
    "alvara_cancelamento":                "parecer_simples",
    "alvara_substituicao_titular":        "parecer_simples",
    "alvara_demolicao":                   "parecer_simples",
    "certidao_demolicao":                 "parecer_simples",
    "certidao_desmembramento":            "parecer_simples",
    "certidao_retificacao_area":          "parecer_simples",

    # ── Ofícios ──
    "oficio_meio_ambiente":               "oficio",
    "parecer_juridico":                   "oficio",
    "oficio_juridico_embargo":            "oficio",
    "oficio_interno_materiais":           "oficio",
    "oficio_decreto_utilidade":           "oficio",

    # ── Comunicados ──
    "comunicado_indeferimento":           "comunicado",
    "comunicado_pendencia":               "comunicado_pendencia",

    # ── Documentos de Emissão da Secretaria (Prontos / Balcão) ──
    "alvara_oficial":                     "documento_pronto",
    "carta_habitese_oficial":             "documento_pronto",
    "certidao_oficial":                   "documento_pronto",
}

# ----------------------------------------------------------------------------
# MODULO: formatacao.py
# ----------------------------------------------------------------------------

"""
Funções auxiliares de formatação DOCX.
Helpers de baixo nível para fontes, espaçamento, tabelas e campos.
"""

import re
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# [inline] from config import (
    FONT_CORPO, SZ_CORPO, COR_INST, COR_LABEL_BG,
    COR_BORDA_TABELA, COR_LABEL_FONT, PAR_AFTER, LINE_SPC,
)


# ═══════════════════════════════════════════════════════════
#  ESPAÇAMENTO E FONTE
# ═══════════════════════════════════════════════════════════

def set_spacing(p, line=276, before=0, after=0):
    """Define espaçamento de um parágrafo."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    pPr.append(sp)


def set_font(run, name=None, size=SZ_CORPO, bold=False, underline=False, italic=False):
    """Aplica fonte, tamanho e estilo a um run."""
    fname = name or FONT_CORPO
    run.font.name      = fname
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.underline = underline
    run.font.italic    = italic
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), fname)
    rf.set(qn('w:hAnsi'), fname)
    rf.set(qn('w:cs'), fname)
    rPr.insert(0, rf)


# ═══════════════════════════════════════════════════════════
#  TEXTO E PARÁGRAFOS
# ═══════════════════════════════════════════════════════════

def clean(t):
    """Remove colchetes espúrios do texto."""
    if not isinstance(t, str):
        return t
    return re.sub(r'\s*\]+\]\s*', ' ', t)


def add_run(p, text, bold=False, size=None, underline=False, italic=False):
    """Adiciona run formatado a um parágrafo."""
    r = p.add_run(clean(text))
    set_font(r, bold=bold, size=size or SZ_CORPO, underline=underline, italic=italic)
    return r


def add_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_SPC,
             before=0, after=PAR_AFTER, indent_cm=None):
    """Cria parágrafo formatado."""
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, line=line, before=before, after=after)
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p


def bold_segments(p, text, size=None, color=None):
    """Processa **negrito** inline em um parágrafo."""
    for i, part in enumerate(text.split('**')):
        if part:
            r = add_run(p, part, bold=(i % 2 == 1), size=size)
            if color is not None:
                r.font.color.rgb = color


def rich_segments(p, text, size=None, color=None):
    """Processa **negrito** e __itálico__ inline em um parágrafo."""
    bold_parts = text.split('**')
    for bi, bpart in enumerate(bold_parts):
        is_bold = (bi % 2 == 1)
        italic_parts = bpart.split('__')
        for ii, ipart in enumerate(italic_parts):
            if ipart:
                is_italic = (ii % 2 == 1)
                r = add_run(p, ipart, bold=is_bold, italic=is_italic, size=size)
                if color is not None:
                    r.font.color.rgb = color


# ═══════════════════════════════════════════════════════════
#  TABELAS E CÉLULAS
# ═══════════════════════════════════════════════════════════

def no_borders(cell):
    """Remove todas as bordas de uma célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        b.append(el)
    tcPr.append(b)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """Define margens internas de uma célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def apply_table_borders(tbl):
    """Aplica bordas padrão cinza a uma tabela."""
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), COR_BORDA_TABELA)
        tblBrd.append(el)
    tblPr.append(tblBrd)


def apply_label_cell(cell, width_twips):
    """Formata célula de rótulo (fundo cinza, texto azul)."""
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), COR_LABEL_BG)
    tcPr.append(shd)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def apply_value_cell(cell, width_twips):
    """Formata célula de valor (fundo branco)."""
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ═══════════════════════════════════════════════════════════
#  ELEMENTOS VISUAIS
# ═══════════════════════════════════════════════════════════

def add_separator(doc, color=COR_INST):
    """Adiciona linha separadora colorida."""
    p_sep = doc.add_paragraph()
    set_spacing(p_sep, line=60, before=80, after=60)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pb_sep = OxmlElement('w:pBdr')
    bot_sep = OxmlElement('w:bottom')
    bot_sep.set(qn('w:val'), 'single')
    bot_sep.set(qn('w:sz'), '4')
    bot_sep.set(qn('w:color'), color)
    pb_sep.append(bot_sep)
    pPr_sep.append(pb_sep)


def add_field(r, instr_text):
    """Insere campo do Word (PAGE, NUMPAGES, etc.)."""
    for tag, txt in [('begin', None), ('instrText', instr_text), ('separate', None), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = txt
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        r._r.append(el)


def add_page_break(doc):
    """Insere uma quebra de página explícita."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_spacing(p, line=0, before=0, after=0)

# ----------------------------------------------------------------------------
# MODULO: cabecalho.py
# ----------------------------------------------------------------------------

"""
Construção do cabeçalho institucional e rodapé com numeração de páginas.
"""

import os
from docx.shared import Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# [inline] from config import (
    LOGO_PREFEITURA, FONT_HEADER, FONT_DETALHE,
    COR_LABEL_FONT, COR_CINZA_LEVE, COR_INST, SZ_RODAPE,
)
# [inline] from formatacao import set_spacing, set_font, no_borders, add_field


# ═══════════════════════════════════════════════════════════
#  CABEÇALHO INSTITUCIONAL
# ═══════════════════════════════════════════════════════════

def build_header(doc):
    """Constrói cabeçalho premium com logo + dados da prefeitura."""
    section = doc.sections[0]
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p._element.getparent().remove(p._element)

    W_LOGO = 2400
    W_TEXT = 7800
    total = W_LOGO + W_TEXT

    tbl = hdr.add_table(rows=1, cols=2, width=Twips(total))
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remover bordas da tabela do cabeçalho
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBrd.append(el)
    tblPr.append(tblBrd)

    row = tbl.rows[0]
    row.height = Twips(1200)

    # ── Logo Prefeitura (à esquerda) ──
    c0 = row.cells[0]
    c0.width = Twips(W_LOGO)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_borders(c0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p0, line=240, before=0, after=0)
    if os.path.exists(LOGO_PREFEITURA):
        p0.add_run().add_picture(LOGO_PREFEITURA, width=Cm(3.3))

    # ── Bloco de texto (à direita do logo) ──
    c1 = row.cells[1]
    c1.width = Twips(W_TEXT)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_borders(c1)
    for extra in c1.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # L1: PREFEITURA MUNICIPAL DE OLIVEIRA
    p1a = c1.paragraphs[0]
    p1a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p1a, line=280, before=0, after=10)  # sem espaço acima
    r_pref = p1a.add_run("PREFEITURA MUNICIPAL DE OLIVEIRA")
    set_font(r_pref, name=FONT_HEADER, size=14, bold=True)
    r_pref.font.color.rgb = COR_LABEL_FONT

    # L2: Secretaria
    p1c = c1.add_paragraph()
    p1c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p1c, line=240, before=0, after=8)
    r_sec = p1c.add_run("Secretaria Municipal de Obras e Serviços Urbanos")
    set_font(r_sec, name=FONT_HEADER, size=8, bold=True)
    r_sec.font.color.rgb = COR_LABEL_FONT

    # L3: Endereço e telefone
    pb = c1.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(pb, line=220, before=0, after=10)
    r_info = pb.add_run(
        "Praça XV de Novembro, 127  •  Centro  •  Oliveira/MG  •  (37) 3331-9800"
    )
    set_font(r_info, name=FONT_DETALHE, size=7)
    r_info.font.color.rgb = COR_CINZA_LEVE

    # ── Linha separadora dupla (fina cinza + grossa azul) ──
    p_thin = hdr.add_paragraph()
    set_spacing(p_thin, line=20, before=8, after=0)  # linha logo após o endereço
    pPr_thin = p_thin._p.get_or_add_pPr()
    pb_thin = OxmlElement('w:pBdr')
    bot_thin = OxmlElement('w:bottom')
    bot_thin.set(qn('w:val'), 'single')
    bot_thin.set(qn('w:sz'), '4')
    bot_thin.set(qn('w:color'), 'AAAAAA')
    bot_thin.set(qn('w:space'), '1')
    pb_thin.append(bot_thin)
    pPr_thin.append(pb_thin)

    p_thick = hdr.add_paragraph()
    set_spacing(p_thick, line=20, before=4, after=0)
    pPr_thick = p_thick._p.get_or_add_pPr()
    pb_thick = OxmlElement('w:pBdr')
    bot_thick = OxmlElement('w:bottom')
    bot_thick.set(qn('w:val'), 'single')
    bot_thick.set(qn('w:sz'), '18')
    bot_thick.set(qn('w:color'), COR_INST)
    bot_thick.set(qn('w:space'), '1')
    pb_thick.append(bot_thick)
    pPr_thick.append(pb_thick)


# ═══════════════════════════════════════════════════════════
#  RODAPÉ COM NUMERAÇÃO
# ═══════════════════════════════════════════════════════════

def add_page_number_footer(doc):
    """Rodapé institucional: linha azul + texto secretaria à esquerda + página à direita."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Linha fina divisora no topo do rodapé ──
    p_linha = footer.add_paragraph()
    set_spacing(p_linha, line=100, before=80, after=20)  # maior respiro acima da linha
    pPr = p_linha._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'), 'single')
    top_el.set(qn('w:sz'), '12')
    top_el.set(qn('w:color'), COR_INST)
    top_el.set(qn('w:space'), '1')
    pBdr.append(top_el)
    pPr.append(pBdr)

    # ── Tabela 2 colunas: Texto à esquerda | Página à direita ──
    tbl = footer.add_table(rows=1, cols=2, width=Twips(9360))
    tbl.autofit = False
    # Sem bordas
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBrd.append(el)
    tblPr.append(tblBrd)

    row = tbl.rows[0]
    # Coluna esquerda: nome da secretaria
    c_left = row.cells[0]
    c_left.width = Twips(6000)
    p_left = c_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_left, line=260, before=8, after=0)
    r_left = p_left.add_run("SMOSU – Secretaria Municipal de Obras e Serviços Urbanos")
    set_font(r_left, size=9)
    r_left.font.color.rgb = COR_CINZA_LEVE

    # Coluna direita: Página X de Y
    c_right = row.cells[1]
    c_right.width = Twips(3360)
    fp = c_right.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, line=260, before=8, after=0)

    r1 = fp.add_run('Página ')
    set_font(r1, size=9)
    r1.font.color.rgb = COR_CINZA_LEVE
    r2 = fp.add_run()
    set_font(r2, size=9)
    r2.font.color.rgb = COR_CINZA_LEVE
    add_field(r2, ' PAGE ')
    r3 = fp.add_run(' de ')
    set_font(r3, size=9)
    r3.font.color.rgb = COR_CINZA_LEVE
    r4 = fp.add_run()
    set_font(r4, size=9)
    r4.font.color.rgb = COR_CINZA_LEVE
    add_field(r4, ' NUMPAGES ')

# ----------------------------------------------------------------------------
# MODULO: componentes.py
# ----------------------------------------------------------------------------

"""
Componentes reutilizáveis para construção de documentos.
Blocos de conteúdo: identificação, dados técnicos, corpo, conclusão, assinatura.
"""

import re
from datetime import datetime
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# [inline] from config import (
    FONT_TITULO, FONT_CORPO, SZ_CORPO, SZ_TABELA,
    COR_INST, COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE,
)
# [inline] from formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator,
)


# ═══════════════════════════════════════════════════════════
#  UTILITÁRIOS
# ═══════════════════════════════════════════════════════════

def _data_hoje_extenso():
    """Retorna data atual por extenso: '20 de abril de 2026'."""
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]
    d = datetime.now()
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _parse_numero(texto):
    """Extrai valor numérico de strings como '86,23%' ou '180,00m²' → float."""
    if not texto:
        return None
    try:
        limpo = re.sub(r'[^\d,.]', '', str(texto)).replace(',', '.')
        return float(limpo) if limpo else None
    except ValueError:
        return None


def _ensure_list(data):
    """Garante que o dado seja uma lista de strings. Se for string única, separa por quebra de linha."""
    if not data:
        return []
    if isinstance(data, str):
        return [item.strip() for item in data.split('\n') if item.strip()]
    return data


# ═══════════════════════════════════════════════════════════
#  TÍTULO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════

def build_titulo(doc, titulo="PARECER SETOR TÉCNICO - SMOSU"):
    """Título centralizado com fonte Cambria bold."""
    pt = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line=288, before=200, after=100)
    r_tit = add_run(pt, titulo, bold=True, size=16)
    set_font(r_tit, name=FONT_TITULO, size=16, bold=True)
    r_tit.font.color.rgb = COR_LABEL_FONT
    # Sem separador aqui — o cabeçalho já tem linha azul


# ═══════════════════════════════════════════════════════════
#  IDENTIFICAÇÃO (Processo, Assunto, Requerente)
# ═══════════════════════════════════════════════════════════

def build_identificacao(doc, d):
    """Tabela de identificação com 3 linhas: Processo, Assunto, Requerente."""
    W_LABEL = 2268
    W_VALUE = 7938
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    linhas = [
        (
            "Processo nº",
            f"{d.get('numero_processo', '')}  —  {d.get('data_processo', '')}"
            if d.get('numero_processo') else ""
        ),
        ("Assunto",    d.get("assunto", "")),
        ("Requerente", d.get("requerente", "")),
    ]
    for i, (rot, val) in enumerate(linhas):
        row = tbl.rows[i]
        c_rot = row.cells[0]
        apply_label_cell(c_rot, W_LABEL)
        p_rot = c_rot.paragraphs[0]
        p_rot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_rot, line=280, before=70, after=70)
        r_rot = p_rot.add_run(rot)
        set_font(r_rot, size=10, bold=True)
        r_rot.font.color.rgb = COR_LABEL_FONT

        c_val = row.cells[1]
        apply_value_cell(c_val, W_VALUE)
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_val, line=280, before=70, after=70)
        rich_segments(p_val, val, size=10)

    add_separator(doc, color='D0D0D0')   # único separador após identificação


# ═══════════════════════════════════════════════════════════
#  DESTINATÁRIO (para Ofícios)
# ═══════════════════════════════════════════════════════════

def build_destinatario(doc, d):
    """Bloco de destinatário para ofícios e comunicados."""
    if d.get("destinatario_titulo"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=200, after=40)
        add_run(p, d["destinatario_titulo"], bold=True, size=SZ_CORPO)

    if d.get("destinatario_para"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "Para: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_para"], size=SZ_CORPO)

    if d.get("destinatario_de"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "De: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_de"], size=SZ_CORPO)

    if d.get("assunto"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=80)
        add_run(p, "Assunto: ", bold=True, size=SZ_CORPO)
        add_run(p, d["assunto"], size=SZ_CORPO)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  DADOS DO CARIMBO TÉCNICO (Decreto 4.149/2019)
# ═══════════════════════════════════════════════════════════

def build_dados_carimbo(doc, d):
    """Tabela 4 colunas com dados técnicos do projeto."""
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=280, after=60)
    r_tit = add_run(p_tit, "DADOS TÉCNICOS DO PROJETO (Ref. Decreto nº 4.149/2019)",
                    bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    logradouro = d.get('logradouro', '')
    bairro = d.get('bairro', '')
    end_completo = (f"{logradouro} - {bairro}" if logradouro and bairro
                    else logradouro or bairro)

    # Detectar alertas urbanísticos
    to_val = _parse_numero(d.get('taxa_ocupacao', ''))
    perm_val = _parse_numero(d.get('taxa_permeabilidade', ''))
    terreno_val = _parse_numero(d.get('area_terreno', ''))
    alerta_zoneamento = (
        (to_val is not None and to_val > 70)
        and (perm_val is not None and perm_val < 20)
    )
    alerta_excecao = (terreno_val is not None and terreno_val < 220)

    linhas = [
        ("Endereço:",       end_completo,
         "Inscrição Mun.:", d.get("inscricao_municipal", "")),
        ("Proprietário:",   d.get("proprietario", d.get("requerente", "")),
         "Desenhista:",     d.get("desenhista", "")),
        ("Lote:",           d.get("lote", ""),
         "Quadra:",         d.get("quadra", "")),
        ("Área Terreno:",   d.get("area_terreno", ""),
         "Área Total:",     d.get("area_total_construida", "")),
        ("Taxa Ocupação:",  d.get("taxa_ocupacao", ""),
         "Coef. Aprov.:",   d.get("coef_aproveitamento", "")),
        ("Permeabilidade:", d.get("taxa_permeabilidade", ""),
         "Resp. Técnico:",  d.get("profissional_nome", "")),
    ]

    W_L1 = 2000
    W_V1 = 3500
    W_L2 = 2000
    W_V2 = 2700

    for lab1, val1, lab2, val2 in linhas:
        row = tbl.add_row()

        # Label 1
        c0 = row.cells[0]
        apply_label_cell(c0, W_L1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p0, line=220, before=25, after=25)
        r0 = p0.add_run(lab1)
        set_font(r0, size=SZ_TABELA, bold=True)
        r0.font.color.rgb = COR_LABEL_FONT

        # Valor 1
        c1 = row.cells[1]
        apply_value_cell(c1, W_V1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p1, line=220, before=25, after=25)

        cor_v1 = None
        alerta_msg_v1 = None

        # Alerta: Área do Terreno < 220m² → verde + "exceção da lei"
        if lab1 == "Área Terreno:" and alerta_excecao:
            cor_v1 = RGBColor(0x00, 0x80, 0x00)
            alerta_msg_v1 = "  (exceção da lei)"

        # Alerta: Taxa Ocupação > 70% → vermelho + "conferir zoneamento"
        if lab1 == "Taxa Ocupação:" and alerta_zoneamento:
            cor_v1 = RGBColor(0xCC, 0x00, 0x00)
            alerta_msg_v1 = "  (conferir zoneamento)"

        # Alerta: Permeabilidade < 20% → vermelho
        if lab1 == "Permeabilidade:" and alerta_zoneamento:
            cor_v1 = RGBColor(0xCC, 0x00, 0x00)
            alerta_msg_v1 = "  (conferir zoneamento)"

        rich_segments(p1, val1, size=SZ_TABELA, color=cor_v1)

        if alerta_msg_v1:
            r_alert = p1.add_run(alerta_msg_v1)
            set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = cor_v1

        # Label 2
        c2 = row.cells[2]
        apply_label_cell(c2, W_L2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p2, line=220, before=25, after=25)
        r2 = p2.add_run(lab2)
        set_font(r2, size=SZ_TABELA, bold=True)
        r2.font.color.rgb = COR_LABEL_FONT

        # Valor 2
        c3 = row.cells[3]
        apply_value_cell(c3, W_V2)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p3, line=220, before=25, after=25)
        rich_segments(p3, val2, size=SZ_TABELA)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  CORPO DO PARECER (abertura + considerandos)
# ═══════════════════════════════════════════════════════════

def build_corpo(doc, d):
    """Corpo do parecer: abertura, considerandos e parágrafos adicionais."""
    INDENT = 1.25

    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=300,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    if d.get("considerandos"):
        for cons in _ensure_list(d["considerandos"]):
            cons_limpo = cons
            # Remove o "Considerando que" se o GEM já tiver escrito na string
            if cons_limpo.lower().startswith("considerando que "):
                cons_limpo = cons_limpo[17:].strip()
            elif cons_limpo.lower().startswith("considerando "):
                cons_limpo = cons_limpo[13:].strip()

            p = add_para(doc, line=LINE_SPC, before=0,
                         after=PAR_AFTER, indent_cm=INDENT)
            add_run(p, "Considerando que ", bold=True, size=SZ_CORPO)
            rich_segments(p, cons_limpo, size=SZ_CORPO)

    if d.get("paragrafos_adicionais"):
        for txt in _ensure_list(d["paragrafos_adicionais"]):
            p_extra = add_para(doc, line=LINE_SPC, before=0,
                               after=PAR_AFTER, indent_cm=INDENT)
            rich_segments(p_extra, txt, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO LEGAL
# ═══════════════════════════════════════════════════════════

def build_fundamentacao(doc, d):
    """Seção de fundamentação legal com bullets."""
    if not d.get("fundamentacao_legal"):
        return

    INDENT = 1.25
    p_fund_tit = add_para(doc, line=LINE_SPC, before=200,
                          after=80, indent_cm=INDENT)
    r_tit = add_run(p_fund_tit, "Da Análise Legal e Técnica:",
                    bold=True, size=SZ_CORPO)
    r_tit.font.color.rgb = COR_LABEL_FONT

    for fund in _ensure_list(d["fundamentacao_legal"]):
        fund_limpo = fund.lstrip("•- \t") # Limpa marcadores se o GEM tiver colocado
        p_fund = add_para(doc, line=LINE_SPC, before=0,
                          after=PAR_AFTER, indent_cm=INDENT)
        r_bullet = add_run(p_fund, "▪ ", size=SZ_CORPO, bold=True)
        r_bullet.font.color.rgb = COR_LABEL_FONT
        rich_segments(p_fund, fund_limpo, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  CARDS DE DOCUMENTOS A EMITIR
# ═══════════════════════════════════════════════════════════

def add_doc_item(doc, tipo, obs=None):
    """Bloco de documento elegante: titulo destacado + observacao em estilo distinto."""
    from docx.shared import Pt, Cm

    # Linha do titulo: fundo azul claríssimo, bullet colorido, nome em negrito
    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_tipo.paragraph_format.left_indent  = Cm(0.3)
    p_tipo.paragraph_format.space_before = Pt(10)
    p_tipo.paragraph_format.space_after  = Pt(2)
    p_tipo.paragraph_format.line_spacing = Pt(14)

    # Fundo azul claríssimo via shading XML
    pPr = p_tipo._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF0FA')
    pPr.append(shd)

    r_bullet = p_tipo.add_run("\u2713  ")
    set_font(r_bullet, size=11, bold=True)
    r_bullet.font.color.rgb = COR_LABEL_FONT

    r_tipo = p_tipo.add_run(tipo)
    set_font(r_tipo, size=11, bold=True)
    r_tipo.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Observacao: prefixo azul + texto italico cinza discreto, recuado
    if obs:
        p_obs = doc.add_paragraph()
        p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_obs.paragraph_format.left_indent  = Cm(0.8)
        p_obs.paragraph_format.space_before = Pt(3)
        p_obs.paragraph_format.space_after  = Pt(10)
        p_obs.paragraph_format.line_spacing = Pt(13)

        r_prefix = p_obs.add_run("Obs.:  ")
        set_font(r_prefix, size=9, bold=True)
        r_prefix.font.color.rgb = COR_LABEL_FONT

        r_obs = p_obs.add_run(obs.strip())
        set_font(r_obs, size=9, italic=True)
        r_obs.font.color.rgb = COR_CINZA_TEXTO


# ═══════════════════════════════════════════════════════════
#  CONCLUSÃO + DOCUMENTOS + ASSINATURA
# ═══════════════════════════════════════════════════════════

def build_conclusao_e_docs(doc, d):
    """Conclusão completa: fundamentação + conclusão + lista de docs + assinatura."""
    # 1º Fundamentação Legal
    build_fundamentacao(doc, d)

    # 2º Conclusão ("Diante do exposto...")
    conclusao = d.get(
        "conclusao",
        "Diante do exposto, após análise das documentações apresentadas, conclui-se que o "
        "projeto está apto a ser aprovado quanto às questões técnicas pertinentes, podendo "
        "ser emitidos os seguintes documentos:"
    )

    p_conc = doc.add_paragraph()
    p_conc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_conc, line=LINE_SPC, before=200, after=80)

    r_conc_icon = p_conc.add_run("CONCLUSÃO TÉCNICA:\n")
    set_font(r_conc_icon, name=FONT_TITULO, size=12, bold=True)
    r_conc_icon.font.color.rgb = COR_LABEL_FONT

    rich_segments(p_conc, conclusao, size=SZ_CORPO)

    # 3º Título "Emissão de Documentos"
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=12, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        # 4º Lista de documentos
        for item in d.get("documentos_emitir", []):
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    # 5º Assinatura
    build_assinatura(doc, d)


def build_conclusao_simples(doc, d):
    """Conclusão simples sem lista de documentos em cards."""
    if d.get("conclusao"):
        p_conc = doc.add_paragraph()
        p_conc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p_conc, line=LINE_SPC, before=200, after=80)

        r_conc_icon = p_conc.add_run("CONCLUSÃO TÉCNICA:\n")
        set_font(r_conc_icon, name=FONT_TITULO, size=12, bold=True)
        r_conc_icon.font.color.rgb = COR_LABEL_FONT

        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # Se houver documentos
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=12, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        for item in d.get("documentos_emitir", []):
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  ASSINATURA
# ═══════════════════════════════════════════════════════════

def build_assinatura(doc, d):
    """Bloco de assinatura: data + linha + nome/título/registro."""
    assinante = d.get("assinante", ASSINANTE)
    cidade = d.get("cidade", CIDADE)

    # Local e data (centralizado, negrito)
    p_local = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       line=LINE_SPC, before=200, after=0)
    add_run(p_local, f"{cidade}, {_data_hoje_extenso()}.", bold=True, size=SZ_CORPO)
    p_local.paragraph_format.keep_with_next = True

    # Espaço
    for _ in range(2):
        pb = doc.add_paragraph()
        set_spacing(pb, line=LINE_SPC, before=0, after=0)
        pb.paragraph_format.keep_with_next = True

    # Linha de assinatura
    p_line = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=240, before=0, after=40)
    r_line = p_line.add_run('_' * 45)
    set_font(r_line, size=10)
    r_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p_line.paragraph_format.keep_with_next = True

    # Dados do assinante
    assinatura_items = [
        (assinante.get("nome",     ASSINANTE["nome"]),     True),
        (assinante.get("titulo",   ASSINANTE["titulo"]),   False),
        (assinante.get("registro", ASSINANTE["registro"]), False),
    ]
    for i, (texto, bold) in enumerate(assinatura_items):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=LINE_SPC, before=0, after=0)
        add_run(p, texto, bold=bold, size=SZ_CORPO)
        if i < len(assinatura_items) - 1:
            p.paragraph_format.keep_with_next = True


# ═══════════════════════════════════════════════════════════
#  COMUNICADO DE PENDÊNCIA — Design Visual Dedicado
# ═══════════════════════════════════════════════════════════

def _box_colorido(doc, fill_hex, borda_hex, borda_sz='12'):
    """Cria uma tabela-card de 1 célula com fundo colorido e borda lateral."""
    card = doc.add_table(rows=1, cols=1)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        card._tbl.insert(0, tblPr)

    # Largura total
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')   # ~16,5 cm em twips
    tblW.set(qn('w:type'), 'dxa')
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)

    # Bordas
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), borda_hex)
        el.set(qn('w:space'), '0')
        tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), borda_sz)
    left_b.set(qn('w:color'), borda_hex)
    left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=100)

    return card, cell


def build_comunicado_pendencia(doc, d):
    """
    Layout visual dedicado para Comunicado de Pendência Documental.

    Estrutura:
      ① Parágrafo de abertura (texto institucional normal)
      ② Box LARANJA — lista numerada de pendências
      ③ Box VERDE   — orientação de continuidade ao requerente
      ④ Conclusão
      ⑤ Assinatura
    """
    INDENT = 1.25
    COR_ALERTA_FILL  = 'FFFDF2'   # amarelo-âmbar super claro
    COR_ALERTA_BORDA = 'F2C94C'   # amarelo/âmbar suave pastel
    COR_OK_FILL      = 'F5FCF5'   # verde muito clarinho e suave
    COR_OK_BORDA     = '81C784'   # verde pastel suave

    # ── ① Parágrafo de abertura ────────────────────────────────────────────
    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=240,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    # ── ② Box de Pendências (laranja/âmbar) ───────────────────────────────
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    card_alert, cell_alert = _box_colorido(doc, COR_ALERTA_FILL, COR_ALERTA_BORDA, '18')

    # Título do box
    p_tit = cell_alert.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=260, before=60, after=60)
    r_icon = p_tit.add_run("⚠  DOCUMENTOS PENDENTES — ANÁLISE SUSPENSA")
    set_font(r_icon, name="Cambria", size=11, bold=True)
    r_icon.font.color.rgb = RGBColor(0x99, 0x65, 0x15)

    # Subtítulo
    p_sub = cell_alert.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_sub, line=240, before=0, after=80)
    r_sub = p_sub.add_run(
        "A análise técnica não pôde ser iniciada. "
        "Os documentos abaixo são imprescindíveis e devem ser apresentados "
        "para o prosseguimento do processo:"
    )
    set_font(r_sub, size=10, italic=True)
    r_sub.font.color.rgb = RGBColor(0x73, 0x5C, 0x1F)

    # Lista de pendências
    considerandos = _ensure_list(d.get("considerandos", []))
    for item in considerandos:
        p_item = cell_alert.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_item, line=260, before=20, after=20)
        # Ícone de marcador
        r_bullet = p_item.add_run("  ✗  ")
        set_font(r_bullet, size=10, bold=True)
        r_bullet.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
        # Texto do item (sem "Considerando que" — é uma lista de pendências)
        texto = item.lstrip("0123456789. ")  # remove "1. ", "2. " se houver
        if texto.lower().startswith("considerando que "):
            texto = texto[17:].strip()
        elif texto.lower().startswith("considerando "):
            texto = texto[13:].strip()

        rich_segments(p_item, texto, size=10, color=RGBColor(0x5C, 0x4A, 0x21))

    # Espaço após o box
    p_esp2 = doc.add_paragraph()
    set_spacing(p_esp2, line=80, before=0, after=0)

    # ── ③ Box de Orientação (verde) ────────────────────────────────────────
    card_ok, cell_ok = _box_colorido(doc, COR_OK_FILL, COR_OK_BORDA, '18')

    p_ok_tit = cell_ok.paragraphs[0]
    p_ok_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_ok_tit, line=260, before=60, after=40)
    r_ok_icon = p_ok_tit.add_run("✔  COMO REGULARIZAR E RETOMAR O PROCESSO")
    set_font(r_ok_icon, name="Cambria", size=11, bold=True)
    r_ok_icon.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    orientacoes = [
        "Reúna todos os documentos listados acima em sua versão original ou digitalizada com legibilidade plena.",
        "Dirija-se ao balcão do protocolo da Prefeitura Municipal de Oliveira ou acesse o portal digital (Atende.Net) para acostar os documentos ao processo.",
        "Após o recebimento e verificação documental, a análise técnica será reiniciada automaticamente, sem necessidade de novo requerimento.",
        "Em caso de dúvidas, entre em contato com o Setor Técnico da SMOSU (Secretaria Municipal de Obras e Serviços Urbanos).",
    ]
    for i, orient in enumerate(orientacoes, 1):
        p_or = cell_ok.add_paragraph()
        p_or.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_or, line=260, before=20, after=20)
        r_num = p_or.add_run(f"  {i}.  ")
        set_font(r_num, size=10, bold=True)
        r_num.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        r_or = p_or.add_run(orient)
        set_font(r_or, size=10)
        r_or.font.color.rgb = RGBColor(0x26, 0x4D, 0x26)

    # Espaço após box verde
    p_esp3 = doc.add_paragraph()
    set_spacing(p_esp3, line=120, before=0, after=0)

    # ── ④ Conclusão ────────────────────────────────────────────────────────
    if d.get("conclusao"):
        p_conc = add_para(doc, line=LINE_SPC, before=120,
                          after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # ── ⑤ Assinatura ───────────────────────────────────────────────────────
    build_assinatura(doc, d)


# ----------------------------------------------------------------------------
# MODULO: alvaras_finais.py
# ----------------------------------------------------------------------------

"""
Gerador de Layout Exclusivo para os Alvarás Oficiais (Secretaria)
Gera o Alvará de Construção e a Carta de Habite-se final com a tipografia
de apresentação de CPF, Matrizes e Fim de Emissão em Papel A4 timbrado.
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# [inline] from formatacao import set_font, add_run, apply_table_borders, set_spacing
# [inline] from config import FONT_CORPO, COR_INST, COR_LABEL_FONT

def build_alvara_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'alvara_oficial.json'.
    Foca nos dados: Nome, CPF, Autores, CREA e dados da Obra.
    """
    _build_cabecalho_simples(doc, dados)

    # Bloco "PROPRIETÁRIO:"
    _add_labeled_pair(doc, "PROPRIETÁRIO:", "NOME", dados.get("proprietario_nome"), "CPF/CNPJ", dados.get("proprietario_cpf_cnpj"))

    # Bloco "AUTOR DO PROJETO:"
    _add_labeled_pair(doc, "AUTOR DO PROJETO:", "NOME", dados.get("autor_projeto_nome"),
                      "CREA/ART", f"{dados.get('autor_projeto_crea', '')} / {dados.get('autor_projeto_art', '')}")

    # Bloco "RESPONSÁVEL TÉCNICO:"
    _add_labeled_pair(doc, "RESPONSÁVEL TÉCNICO:", "NOME", dados.get("responsavel_tecnico_nome"),
                      "CREA/ART", f"{dados.get('responsavel_tecnico_crea', '')} / {dados.get('responsavel_tecnico_art', '')}")

    # Bloco "CONSTRUTORA:"
    if dados.get("construtora_nome"):
        _add_labeled_pair(doc, "CONSTRUTORA OU RESPONSÁVEL PELA EXECUÇÃO:", "NOME", dados.get("construtora_nome"),
                          "CPF/CNPJ", dados.get("construtora_cpf_cnpj"))

    # Texto Livre (Concessão)
    t_processo = dados.get("numero_processo", "")
    t_data = dados.get("data_aprovacao", "")
    t_obra = dados.get("nome_obra", "")
    t_endereco = f"{dados.get('logradouro', '')}, {dados.get('bairro', '')}"
    texto_abertura = f"Tendo em vista o constante no processo nº {t_processo} fica concedida a licença para execução do projeto aprovado em {t_data} a obra denominada de {t_obra}, a ser executada no endereço: {t_endereco}."

    p = doc.add_paragraph()
    set_spacing(p, before=100, after=200, line=240)
    add_run(p, texto_abertura, size=11).font.color.rgb = RGBColor(0,0,0)

    # Matriz da Obra
    _build_matriz_areas(doc, dados)

    # Assinatura (Rodapé Simplificado do Alvará)
    _build_assinatura_balcao(doc)


def build_habitese_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'carta_habitese_oficial.json'.
    Foca no endereço em cima e despacho.
    """
    _build_cabecalho_simples(doc, dados)

    # Endereço
    p_end = doc.add_paragraph()
    add_run(p_end, "ENDEREÇO DA OBRA: ", bold=True)
    add_run(p_end, dados.get("logradouro", ""))
    p_end2 = doc.add_paragraph()
    add_run(p_end2, "BAIRRO: ", bold=True)
    add_run(p_end2, dados.get("bairro", ""))

    # Proprietário
    _add_labeled_pair(doc, "PROPRIETÁRIO DO IMÓVEL:", "NOME", dados.get("proprietario_nome"), "CPF", dados.get("proprietario_cpf_cnpj"))
    _add_labeled_pair(doc, "RESPONSÁVEL PELA EXECUÇÃO DA OBRA:", "NOME", dados.get("responsavel_execucao_nome"), "CPF", dados.get("responsavel_execucao_cpf_cnpj"))

    # Responsável
    p_resp = doc.add_paragraph()
    set_spacing(p_resp, before=100, after=50)
    add_run(p_resp, "RESPONSÁVEL TÉCNICO:", bold=True, size=11)
    ptx = doc.add_paragraph()
    set_spacing(ptx, before=0, after=200, line=240)
    add_run(ptx, dados.get("texto_despacho_responsavel_tecnico", ""))

    # Matriz de Áreas
    _build_matriz_areas(doc, dados)

    _build_assinatura_balcao(doc)


def build_certidao_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'certidao_oficial.json'.
    Foca no texto corrido justificado e assinaturas lado a lado.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=100, after=400)
    add_run(p, dados.get("titulo_documento", "CERTIDÃO"), size=16, bold=True).font.color.rgb = COR_LABEL_FONT

    p_text = doc.add_paragraph()
    p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_text, before=200, after=400, line=360)
    add_run(p_text, "        " + dados.get("texto_certidao", ""), size=12).font.color.rgb = RGBColor(0,0,0)

    from datetime import datetime
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    d = datetime.now()
    hoje = f"Oliveira, {d.day} de {meses[d.month - 1]} de {d.year}."

    p_dt = doc.add_paragraph()
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(p_dt, before=200, after=600)
    add_run(p_dt, hoje, size=11).font.color.rgb = RGBColor(0,0,0)

    assinantes = dados.get("assinantes", [])
    if assinantes:
        p_ass = doc.add_paragraph()
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p_ass, before=400, after=200)
        for ass in assinantes:
            add_run(p_ass, str(ass.get("nome", "")) + "\n", bold=True, size=11)
            add_run(p_ass, str(ass.get("titulo", "")) + "\n\n", size=10)

    obs = dados.get("observacoes_finais", [])
    if isinstance(obs, list):
        for o in obs:
            po = doc.add_paragraph()
            po.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(po, o, size=9, bold=True).font.color.rgb = RGBColor(0,0,0)


# Helpers Visuais de Guias Finais

def _build_cabecalho_simples(doc, dados):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "MUNICIPIO DE OLIVEIRA\nSecretaria Municipal de Obras", size=12, bold=True)

    num_doc = dados.get("numero_documento", "XXX/XXXX")
    tit = str(dados.get("titulo_documento", "DOCUMENTO Nº {numero}")).replace("{numero}", num_doc)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(pt, before=100, after=250)
    rt = add_run(pt, tit, size=14, bold=True)
    rt.font.color.rgb = COR_LABEL_FONT


def _add_labeled_pair(doc, main_label, l1, v1, l2, v2):
    p = doc.add_paragraph()
    set_spacing(p, before=150, after=30)
    add_run(p, main_label, bold=True, size=11)

    p1 = doc.add_paragraph()
    set_spacing(p1, before=0, after=0)
    add_run(p1, f"{l1}: ", bold=True, size=10)
    add_run(p1, v1 if v1 else "---", size=10)

    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=0)
    add_run(p2, f"{l2}: ", bold=True, size=10)
    add_run(p2, v2 if v2 else "---", size=10)


def _build_matriz_areas(doc, dados):
    p = doc.add_paragraph()
    set_spacing(p, before=150, after=50)
    add_run(p, "Dados da obra:\nÁreas principais", bold=True, size=11)

    matriz = dados.get("areas_matriz", [])
    if matriz:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        apply_table_borders(tbl)

        # Cabeçalhos
        for i, h in enumerate(["Categoria", "Destinação", "Tipo de Obra", "Área (m²)"]):
            c = tbl.cell(0, i)
            pt = c.paragraphs[0]
            add_run(pt, h, bold=True, size=9)

        for row in matriz:
            tr = tbl.add_row()
            if isinstance(row, dict):
                add_run(tr.cells[0].paragraphs[0], str(row.get("categoria","")), size=9)
                add_run(tr.cells[1].paragraphs[0], str(row.get("destinacao","")), size=9)
                add_run(tr.cells[2].paragraphs[0], str(row.get("tipo_obra","")), size=9)
                add_run(tr.cells[3].paragraphs[0], str(row.get("area_m2","")), size=9)

    p_tot = doc.add_paragraph()
    set_spacing(p_tot, before=150, after=50)
    add_run(p_tot, "Área total da obra: ", bold=True, size=11)
    add_run(p_tot, str(dados.get("area_total_obra", "")), size=11)

    if dados.get("observacoes"):
        po = doc.add_paragraph()
        set_spacing(po, before=100, after=200, line=240)
        add_run(po, "Observações: ", bold=True, size=10)
        add_run(po, dados.get("observacoes", ""), size=10)

def _build_assinatura_balcao(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=400, after=400)

    p_ass = doc.add_paragraph()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_ass, "_" * 45, size=10)
    p_as2 = doc.add_paragraph()
    p_as2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_as2, "SECRETARIA MUNICIPAL DE OBRAS / PMO", bold=True, size=10)

# ----------------------------------------------------------------------------
# MODULO: geradores\__init__.py
# ----------------------------------------------------------------------------

"""
Sistema de despacho de geradores de documentos — SMOSU Oliveira/MG.

Identifica o tipo_relatorio no JSON de entrada, carrega o template
correspondente e chama o gerador correto (parecer_tecnico, parecer_simples,
ofício ou comunicado).
"""

import os
import json

from docx import Document
from docx.shared import Pt, Cm

# [inline] from config import (
    TIPOS_DOCUMENTO, TEMPLATES_DIR, PROJECT_DIR,
    FONT_CORPO, SZ_CORPO,
)
# [inline] from cabecalho import build_header, add_page_number_footer
# [inline] from componentes import (
    build_titulo, build_identificacao, build_destinatario,
    build_dados_carimbo, build_corpo, build_fundamentacao,
    build_conclusao_e_docs, build_conclusao_simples, build_assinatura,
    build_comunicado_pendencia,
)
# [inline] from alvaras_finais import build_alvara_oficial, build_habitese_oficial, build_certidao_oficial


# ═══════════════════════════════════════════════════════════
#  TEMPLATES
# ═══════════════════════════════════════════════════════════

def carregar_template(tipo_relatorio):
    """Carrega o template JSON para o tipo de documento, se existir."""
    caminho = os.path.join(TEMPLATES_DIR, f"{tipo_relatorio}.json")
    if os.path.exists(caminho):
        with open(caminho, encoding="utf-8") as f:
            return json.load(f)
    return {}  # template vazio = usar defaults


# ═══════════════════════════════════════════════════════════
#  DOCUMENTO BASE
# ═══════════════════════════════════════════════════════════

def _criar_documento_base():
    """Cria documento com margens A4 e estilo padrão configurados."""
    doc = Document()
    section = doc.sections[0]
    # Margens 1.5cm + cabeçalho rente ao topo + rodapé com altura
    M = Cm(1.5)
    section.page_height      = Cm(29.7)
    section.page_width       = Cm(21.0)
    section.top_margin       = Cm(3.2)   # corpo começa após o cabeçalho (~2.8cm de altura)
    section.bottom_margin    = Cm(2.2)   # espaço generoso para o rodapé
    section.left_margin      = M
    section.right_margin     = M
    section.header_distance  = Cm(0.3)   # logo colado na borda superior — sem borda branca
    section.footer_distance  = Cm(0.5)   # rodapé colado na borda inferior

    style = doc.styles['Normal']
    style.font.name = FONT_CORPO
    style.font.size = Pt(SZ_CORPO)
    return doc


def _gerar_nome_saida(dados):
    """Gera nome de arquivo baseado nos dados do processo."""
    proc = str(dados.get("numero_processo", "RELATORIO")).replace("/", "-")

    # Nomenclatura Inteligente: busca nome na ordem de probabilidade das chaves
    nome_alvo = dados.get("requerente") or dados.get("proprietario_nome") or dados.get("interessado") or "Desconhecido"
    req_str = str(nome_alvo).split()[0]

    tipo = dados.get("tipo_relatorio", "PARECER").upper()
    nome = f"{tipo}_{proc} {req_str}.docx"
    for ch in ['<', '>', ':', '"', '|', '?', '*']:
        nome = nome.replace(ch, '')

    # PROJECT_DIR is now _Sistema_Interno. The root is its parent.
    root_dir = os.path.dirname(PROJECT_DIR)
    out_dir = os.path.join(root_dir, "2_Documentos_Prontos")
    os.makedirs(out_dir, exist_ok=True)
    return os.path.join(out_dir, nome)


# ═══════════════════════════════════════════════════════════
#  GERADORES POR CATEGORIA
# ═══════════════════════════════════════════════════════════

def gerar_parecer_tecnico(doc, dados, template):
    """
    Parecer técnico COMPLETO:
    Header → Título → Identificação → Dados Carimbo → Corpo → Conclusão+Docs
    """
    titulo = template.get("titulo_documento", "PARECER SETOR TÉCNICO - SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_dados_carimbo(doc, dados)
    build_corpo(doc, dados)
    build_conclusao_e_docs(doc, dados)


def gerar_parecer_simples(doc, dados, template):
    """
    Parecer SIMPLES (sem dados do carimbo):
    Header → Título → Identificação → Corpo → Conclusão
    """
    titulo = template.get("titulo_documento", "PARECER SETOR TÉCNICO - SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_corpo(doc, dados)

    if dados.get("documentos_emitir"):
        build_conclusao_e_docs(doc, dados)
    else:
        build_conclusao_simples(doc, dados)


def gerar_oficio(doc, dados, template):
    """
    OFÍCIO:
    Header → Título → Destinatário → Corpo → Assinatura
    """
    titulo = template.get("titulo_documento", "OFÍCIO")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_destinatario(doc, dados)
    build_corpo(doc, dados)
    build_assinatura(doc, dados)


def gerar_comunicado(doc, dados, template):
    """
    COMUNICADO genérico:
    Header → Título → Identificação → Corpo → Assinatura
    """
    titulo = template.get("titulo_documento", "COMUNICADO")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_corpo(doc, dados)
    build_assinatura(doc, dados)


def gerar_comunicado_pendencia(doc, dados, template):
    """
    COMUNICADO DE PENDÊNCIA — layout visual dedicado:
    Header → Título → Identificação → [Box Alerta + Box Orientação] → Assinatura
    """
    titulo = template.get("titulo_documento", "COMUNICADO DE PENDÊNCIA DOCUMENTAL")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_comunicado_pendencia(doc, dados)


def gerar_documento_pronto(doc, dados, template):
    """
    DOCUMENTOS DE SECRETARIA (ALVARÁ FINAL / HABITE-SE)
    Design totalmente único, usando o novo motor em alvaras_finais.py
    Não usa o Header e identificacao padrão.
    """
    tipo = dados.get("tipo_relatorio", "")
    if "habitese" in tipo:
        build_habitese_oficial(doc, dados)
    elif "certidao" in tipo:
        build_certidao_oficial(doc, dados)
    else:
        build_alvara_oficial(doc, dados)


# ═══════════════════════════════════════════════════════════
#  MAPEAMENTO E DESPACHO
# ═══════════════════════════════════════════════════════════

GERADORES = {
    "parecer_tecnico":      gerar_parecer_tecnico,
    "parecer_simples":      gerar_parecer_simples,
    "oficio":               gerar_oficio,
    "comunicado":           gerar_comunicado,
    "comunicado_pendencia": gerar_comunicado_pendencia,
    "documento_pronto":     gerar_documento_pronto,
}


def gerar(dados, caminho_saida=None):
    """
    Função principal: gera documento baseado nos dados e tipo_relatorio.

    Parâmetros:
        dados: dict com os dados do processo (JSON do Gemini)
        caminho_saida: caminho opcional para o arquivo .docx

    Retorna:
        Caminho do arquivo gerado.
    """
    tipo = dados.get("tipo_relatorio", "regularizacao")

    # ---------------------------------------------------------
    # RESILIÊNCIA: Se o LLM agrupar os campos em sub-dicionários
    # como 'campos_obrigatorios' ou 'campos_opcionais', nós os achata.
    # ---------------------------------------------------------
    if "campos_obrigatorios" in dados and isinstance(dados["campos_obrigatorios"], dict):
        dados.update(dados.pop("campos_obrigatorios"))
    if "campos_opcionais" in dados and isinstance(dados["campos_opcionais"], dict):
        dados.update(dados.pop("campos_opcionais"))

    # Tratamento caso ele faça lista de strings em documentos_emitir no invés de `{tipo:..}`
    if "documentos_emitir" in dados and isinstance(dados["documentos_emitir"], list):
        lista_corrigida = []
        for d in dados["documentos_emitir"]:
            if isinstance(d, str):
                lista_corrigida.append({"tipo": d})
            else:
                lista_corrigida.append(d)
        dados["documentos_emitir"] = lista_corrigida

    categoria = TIPOS_DOCUMENTO.get(tipo)

    if not categoria:
        print(f"[!] Tipo de documento desconhecido: '{tipo}'")
        print(f"    Tipos disponíveis: {', '.join(sorted(TIPOS_DOCUMENTO.keys()))}")
        raise ValueError(f"Tipo de documento desconhecido: {tipo}")

    print(f"[>] Tipo: {tipo} -> Categoria: {categoria}")

    # Carregar template (se existir)
    template = carregar_template(tipo)

    # === VALIDADOR DE INTEGRIDADE PARA AUTO-CORREÇÃO DO GEM ===
    obrigatorios = template.get("campos_obrigatorios", [])
    campos_faltantes = [c for c in obrigatorios if not dados.get(c)]

    if campos_faltantes:
        print("\n" + "="*70)
        print("  [ERRO ESTRUTURAL] VALIDAÇÃO DE CHAVES JSON FALHOU")
        print("="*70)
        print("  Parece que o Assistente GEM esqueceu de incluir algumas chaves obrigatórias!")
        print(f"  Tipo do Documento: {tipo}")
        print(f"  Chaves Faltantes : {', '.join(campos_faltantes)}")
        print("-" * 70)
        print("  [Ação de Correção] Copie e cole a mensagem abaixo de volta pro GEM:\n")
        print(f"  Gem, o JSON que você gerou falhou na validação estrutural do sistema.")
        print(f"  O template oficial para '{tipo}' exige o uso EXATO destas CHAVES JSON:")
        print(f"  >>> {', '.join(campos_faltantes)}")
        print(f"  Por favor, revise o PDF do processo, certifique-se de usar estes nomes exatos de chaves")
        print(f"  e gere o bloco de JSON corrigido e completo para eu compilar.")
        print("="*70 + "\n")
        raise ValueError(f"Faltam {len(campos_faltantes)} chaves obrigatórias: {','.join(campos_faltantes)}")


    # Criar documento base
    doc = _criar_documento_base()

    # Chamar gerador
    gerador_fn = GERADORES[categoria]
    gerador_fn(doc, dados, template)

    # Salvar
    if caminho_saida is None:
        caminho_saida = _gerar_nome_saida(dados)

    doc.save(caminho_saida)
    print(f"[+] Documento DOCX gerado: {caminho_saida}")

    # Gerar PDF automaticamente
    _gerar_pdf(caminho_saida)

    return caminho_saida


# ═══════════════════════════════════════════════════════════
#  GERAÇÃO DE PDF
# ═══════════════════════════════════════════════════════════

def _gerar_pdf(caminho_docx):
    """Tenta gerar PDF a partir do DOCX (docx2pdf ou comtypes)."""
    caminho_pdf = caminho_docx.replace('.docx', '.pdf')
    try:
        from docx2pdf import convert
        convert(caminho_docx, caminho_pdf)
        print(f"[+] Documento PDF  gerado: {caminho_pdf}")
    except ImportError:
        try:
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc_w = word.Documents.Open(os.path.abspath(caminho_docx))
            doc_w.SaveAs(os.path.abspath(caminho_pdf), FileFormat=17)
            doc_w.Close()
            word.Quit()
            print(f"[+] Documento PDF  gerado: {caminho_pdf}")
        except Exception as e:
            print(f"[!] PDF não gerado (instale docx2pdf ou comtypes): {e}")

# ----------------------------------------------------------------------------
# MODULO: compilador.py
# ----------------------------------------------------------------------------

"""
Gerador Automático de Documentos – Prefeitura de Oliveira / SMOSU
Orquestrador principal – despacha para os geradores especializados.

Uso: python _engine\\compilador.py dados.json [saida.docx]
"""

import sys
import json
import os
import glob

# Adicionar o diretório do script ao path para imports locais
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

# [inline] from config import TIPOS_DOCUMENTO
# [inline] from geradores import gerar


def main():
    args = sys.argv[1:]

    # Sem argumentos: tentar usar a pasta padrão 1_Colar_JSON_Aqui
    if not args:
        pasta_padrao = os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), "1_Colar_JSON_Aqui")
        if os.path.exists(pasta_padrao) and glob.glob(os.path.join(pasta_padrao, "*.json")):
            args = [pasta_padrao]
        else:
            print("╔══════════════════════════════════════════════════════════════════╗")
            print("║  Gerador de Documentos – SMOSU Oliveira/MG                     ║")
            print("╠══════════════════════════════════════════════════════════════════╣")
            print("║                                                                ║")
            print("║  Uso: python compilador.py dados.json [saida.docx]             ║")
            print("║                                                                ║")
            print("║  O JSON deve conter o campo 'tipo_relatorio' com um dos tipos  ║")
            print("║  listados abaixo.                                              ║")
            print("║                                                                ║")
            print("╠══════════════════════════════════════════════════════════════════╣")
            print("║  TIPOS DE DOCUMENTO DISPONÍVEIS:                               ║")
            print("╠══════════════════════════════════════════════════════════════════╣")

            # Agrupar por categoria
            categorias = {}
            for tipo, cat in sorted(TIPOS_DOCUMENTO.items()):
                categorias.setdefault(cat, []).append(tipo)

            for cat in ["parecer_tecnico", "parecer_simples", "oficio", "comunicado"]:
                tipos = categorias.get(cat, [])
                if tipos:
                    print(f"║  [{cat.upper():^20}]                                     ║")
                    for t in tipos:
                        print(f"║    • {t:<56} ║")
                    print("║                                                                ║")

            print("╚══════════════════════════════════════════════════════════════════╝")
            sys.exit(0)

    alvo = args[0]
    arquivos_para_processar = []

    if os.path.isdir(alvo):
        print(f"\n[>] MODO EM LOTE: Escaneando pasta '{alvo}'")
        arquivos_para_processar = glob.glob(os.path.join(alvo, "*.json"))
        if not arquivos_para_processar:
            print(f"[!] Nenhum arquivo .json encontrado na pasta de origem.")
            sys.exit(0)
    elif os.path.isfile(alvo) and alvo.endswith(".json"):
        arquivos_para_processar = [alvo]
    else:
            print(f"[!] Erro: O caminho não é um JSON válido ou não existe -> {alvo}")
            sys.exit(1)

    caminho_saida_fornecido = args[1] if len(args) > 1 else None

    if len(arquivos_para_processar) > 1 and caminho_saida_fornecido:
            print("[!] Aviso: Parâmetro de saída único ignorado devido ao processamento em lote.")
            caminho_saida_fornecido = None

    sucessos = 0
    erros = 0

    from traceback import print_exc

    for arquivo in arquivos_para_processar:
            # Ler JSON de entrada
            try:
                with open(arquivo, encoding="utf-8") as f:
                    dados = json.load(f)
                print(f"\n[>] Processando arquivo: {os.path.basename(arquivo)}")
            except json.JSONDecodeError as e:
                print(f"[X] Ignorando {os.path.basename(arquivo)}: JSON inválido ({e})")
                erros += 1
                continue

            # Verificar marcadores de segurança
            json_str = json.dumps(dados, ensure_ascii=False)
            if "⚠️ VERIFICAR" in json_str:
                print("  - [AVISO TÉCNICO]: O JSON contém marcações incompletas (VERIFICAR).")
                print("  - Dica para o GEM:")
                print("    'Gem, você marcou dados como 'VERIFICAR'. Releia os anexos")
                print("    com mais atenção e tente cruzar os dados pra ter a certeza.'\n")

            # Gerar documento
            try:
                gerar(dados, caminho_saida_fornecido)
                sucessos += 1
            except Exception as e:
                # Em modo lote, interceptar erro e continuar o loop sem crashar a prefeitura inteira
                print(f"  [ALERTA DE SISTEMA] Falha ao compilar {os.path.basename(arquivo)}: {e}")
                erros += 1

    print("\n" + "="*70)
    print(f" [V] COMPILAÇÃO ENCERRADA. Sucessos: {sucessos} | Falhas de Dados: {erros}")
    print("="*70)


if __name__ == "__main__":
    main()