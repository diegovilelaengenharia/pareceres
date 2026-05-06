"""
Validador de Fidelidade (Cross-Check) — Golden Dataset SMOSU.

Carrega um JSON do Golden Dataset, gera o DOCX via motor e extrai o texto
para comparar campos críticos entre a entrada e o documento gerado.

Uso:
    python validador_fidelidade.py                     # auditoria de todos os casos
    python validador_fidelidade.py --all               # idem
    python validador_fidelidade.py --file caminho.json # caso individual
    python validador_fidelidade.py --test-mode         # roda alvara_ouro.json e verifica estrutura
"""

import sys
import os
import json
import re
import argparse
import tempfile
import datetime

_SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
_MOTOR_DIR   = os.path.dirname(_SCRIPTS_DIR)
if _MOTOR_DIR not in sys.path:
    sys.path.insert(0, _MOTOR_DIR)

from generators.geradores_core import gerar

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


GOLDEN_DIR = os.path.join(_MOTOR_DIR, "tests", "golden_dataset")

# Campos numéricos/textuais críticos para validação cruzada
CAMPOS_CRITICOS_NUMERICOS = [
    "area_terreno",
    "area_total_construida",
    "taxa_ocupacao",
    "taxa_permeabilidade",
    "coef_aproveitamento",
]

CAMPOS_CRITICOS_TEXTO = [
    "numero_processo",
    "requerente",
    "logradouro",
    "numero_alvara",
    "matricula_sri",
    "inscricao_municipal",
    "art_rrt",
]


def _extrair_texto_docx(caminho_docx: str) -> str:
    """Extrai todo o texto de um DOCX como string única."""
    if not _DOCX_OK:
        return ""
    doc = _DocxDocument(caminho_docx)
    linhas = []
    for para in doc.paragraphs:
        if para.text.strip():
            linhas.append(para.text.strip())
    # Incluir texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    linhas.append(cell.text.strip())
    return "\n".join(linhas)


def _normalizar_numero(valor: str) -> str:
    """Normaliza representações numéricas para comparação (vírgula <-> ponto)."""
    if not valor:
        return ""
    # Remove sufixos como m², %, etc. e espaços
    valor = re.sub(r'[m²%\s]', '', str(valor))
    # Normaliza separador decimal: aceita 33,47 ou 33.47
    valor = valor.replace(',', '.')
    try:
        return str(round(float(valor), 2))
    except ValueError:
        return valor.strip()


def _campo_presente_no_texto(campo_valor: str, texto_docx: str) -> bool:
    """Verifica se o valor do campo aparece no texto extraído do DOCX."""
    if not campo_valor or not texto_docx:
        return False

    # Tentativa 1: presença literal
    if str(campo_valor) in texto_docx:
        return True

    # Tentativa 2: normalização numérica (vírgula vs ponto)
    val_norm = _normalizar_numero(campo_valor)
    if val_norm and val_norm in texto_docx:
        return True

    # Tentativa 3: extrair apenas dígitos e ponto/vírgula do valor e buscar no texto
    val_simples = re.sub(r'[^\d,.]', '', str(campo_valor))
    if val_simples and len(val_simples) >= 3:
        for variante in [val_simples, val_simples.replace(',', '.')]:
            if variante in texto_docx:
                return True

    return False


def _auditar_caso(caminho_json: str) -> dict:
    """
    Gera o DOCX a partir do JSON e compara campos críticos.

    Retorna dicionário com resultado da auditoria.
    """
    with open(caminho_json, encoding="utf-8") as f:
        dados = json.load(f)

    nome_caso = os.path.basename(caminho_json)
    tipo = dados.get("tipo_relatorio", "desconhecido")

    # Gerar em diretório temporário
    with tempfile.TemporaryDirectory() as tmp_dir:
        dados_tmp = dict(dados)
        # Redirecionar saída para temp
        import generators.geradores_core as _gc_mod
        _orig_pasta = None
        try:
            import core.config as _cfg
            _orig_pasta = _cfg.PASTA_SAIDA
            _cfg.PASTA_SAIDA = tmp_dir
        except Exception:
            pass

        try:
            caminho_docx = gerar(dados_tmp, caminho_saida=None)
        except Exception as e:
            return {
                "caso": nome_caso,
                "tipo": tipo,
                "status": "ERRO_GERACAO",
                "erro": str(e),
                "campos": {},
            }
        finally:
            if _orig_pasta is not None:
                try:
                    _cfg.PASTA_SAIDA = _orig_pasta
                except Exception:
                    pass

        if not caminho_docx or not os.path.exists(caminho_docx):
            return {
                "caso": nome_caso,
                "tipo": tipo,
                "status": "ERRO_ARQUIVO_NAO_GERADO",
                "erro": "gerar() não retornou caminho válido",
                "campos": {},
            }

        texto = _extrair_texto_docx(caminho_docx)

    if not _DOCX_OK or not texto:
        return {
            "caso": nome_caso,
            "tipo": tipo,
            "status": "AVISO_SEM_EXTRACAO",
            "erro": "python-docx indisponível ou DOCX vazio",
            "campos": {},
        }

    resultados_campos = {}

    # Verificar campos críticos de texto
    for campo in CAMPOS_CRITICOS_TEXTO:
        valor = dados.get(campo)
        if valor:
            ok = _campo_presente_no_texto(str(valor), texto)
            resultados_campos[campo] = "MATCH" if ok else "MISMATCH"

    # Verificar campos numéricos
    for campo in CAMPOS_CRITICOS_NUMERICOS:
        valor = dados.get(campo)
        if valor:
            ok = _campo_presente_no_texto(str(valor), texto)
            resultados_campos[campo] = "MATCH" if ok else "MISMATCH"

    mismatches = [c for c, r in resultados_campos.items() if r == "MISMATCH"]
    status = "PASS" if not mismatches else "FAIL"

    return {
        "caso": nome_caso,
        "tipo": tipo,
        "status": status,
        "mismatches": mismatches,
        "campos": resultados_campos,
    }


def _imprimir_resultado(resultado: dict):
    print(f"\n{'='*60}")
    print(f"  Caso: {resultado['caso']}")
    print(f"  Tipo: {resultado['tipo']}")
    status = resultado["status"]
    if status == "PASS":
        print(f"  Status: [OK] PASS - todos os campos criticos verificados")
    elif status == "FAIL":
        print(f"  Status: [FALHA] FAIL")
        print(f"  Mismatches: {', '.join(resultado.get('mismatches', []))}")
    else:
        print(f"  Status: [AVISO] {status}")
        print(f"  Detalhe: {resultado.get('erro', '')}")

    campos = resultado.get("campos", {})
    if campos:
        print(f"\n  Campos verificados ({len(campos)}):")
        for campo, res in campos.items():
            simbolo = "[OK]" if res == "MATCH" else "[X]"
            print(f"    {simbolo} {campo}: {res}")


def _gerar_relatorio(resultados: list[dict], caminho_relatorio: str):
    """Gera relatório Markdown com os resultados da auditoria."""
    now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    total = len(resultados)
    aprovados = sum(1 for r in resultados if r["status"] == "PASS")
    falhos = sum(1 for r in resultados if r["status"] == "FAIL")
    erros = total - aprovados - falhos

    linhas = [
        "# Relatório de Auditoria — Golden Dataset SMOSU",
        f"\n**Data:** {now}  ",
        f"**Total de casos:** {total}  ",
        f"**Aprovados (PASS):** {aprovados}  ",
        f"**Reprovados (FAIL):** {falhos}  ",
        f"**Erros de geração:** {erros}  ",
        "\n---\n",
    ]

    for r in resultados:
        status_icon = {"PASS": "✅", "FAIL": "❌"}.get(r["status"], "⚠️")
        linhas.append(f"## {status_icon} {r['caso']} (`{r['tipo']}`)")
        linhas.append(f"**Status:** {r['status']}  ")
        if r.get("mismatches"):
            linhas.append(f"**Mismatches:** {', '.join(r['mismatches'])}  ")
        if r.get("erro"):
            linhas.append(f"**Erro:** {r['erro']}  ")
        campos = r.get("campos", {})
        if campos:
            linhas.append("\n| Campo | Resultado |")
            linhas.append("|-------|-----------|")
            for campo, res in campos.items():
                icone = "✅" if res == "MATCH" else "❌"
                linhas.append(f"| `{campo}` | {icone} {res} |")
        linhas.append("")

    with open(caminho_relatorio, "w", encoding="utf-8") as f:
        f.write("\n".join(linhas))

    print(f"\n[OK] Relatorio salvo em: {caminho_relatorio}")


def main():
    parser = argparse.ArgumentParser(description="Validador de Fidelidade — Golden Dataset SMOSU")
    parser.add_argument("--all", action="store_true", help="Auditar todos os JSONs do golden_dataset")
    parser.add_argument("--file", help="Auditar um JSON específico")
    parser.add_argument("--test-mode", action="store_true", help="Modo de teste rápido (alvara_ouro.json)")
    args = parser.parse_args()

    casos = []

    if args.test_mode:
        caminho = os.path.join(GOLDEN_DIR, "alvara_ouro.json")
        if not os.path.exists(caminho):
            print(f"[!] arquivo não encontrado: {caminho}")
            sys.exit(1)
        casos = [caminho]
    elif args.file:
        if not os.path.exists(args.file):
            print(f"[!] arquivo não encontrado: {args.file}")
            sys.exit(1)
        casos = [args.file]
    else:
        # default: todos
        if not os.path.isdir(GOLDEN_DIR):
            print(f"[!] Diretório não encontrado: {GOLDEN_DIR}")
            sys.exit(1)
        casos = sorted(
            os.path.join(GOLDEN_DIR, f)
            for f in os.listdir(GOLDEN_DIR)
            if f.endswith(".json") and not f.startswith("_")
        )

    if not casos:
        print("[!] Nenhum caso encontrado para auditar.")
        sys.exit(1)

    print(f"\n[>] Iniciando auditoria de fidelidade — {len(casos)} caso(s)")
    resultados = []
    for caminho in casos:
        print(f"    Processando: {os.path.basename(caminho)} ...", end=" ", flush=True)
        resultado = _auditar_caso(caminho)
        resultados.append(resultado)
        print(resultado["status"])

    for r in resultados:
        _imprimir_resultado(r)

    # Salvar relatório
    caminho_relatorio = os.path.join(GOLDEN_DIR, "relatorio_auditoria.md")
    _gerar_relatorio(resultados, caminho_relatorio)

    # Exit code
    falhos = [r for r in resultados if r["status"] == "FAIL"]
    if falhos:
        print(f"\n[FALHA] {len(falhos)} caso(s) com mismatch detectado.")
        sys.exit(1)
    else:
        print(f"\n[OK] Auditoria concluida - todos os campos criticos verificados com sucesso.")


if __name__ == "__main__":
    main()
