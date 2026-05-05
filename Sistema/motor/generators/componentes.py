"""
Componentes reutilizáveis para construção de documentos.
Blocos de conteúdo: identificação, dados técnicos, corpo, conclusão, assinatura.
"""

import re
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)


# ═══════════════════════════════════════════════════════════
#  UTILITÁRIOS
# ═══════════════════════════════════════════════════════════

def _data_hoje_extenso():
    """Retorna data atual por extenso: '20 de abril de 2026'."""
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]
    d = datetime.now()
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _parse_numero(texto):
    """Extrai valor numérico de strings como '86,23%' ou '180,00m²' → float."""
    if not texto:
        return None
    try:
        limpo = re.sub(r'[^\d,.]', '', str(texto)).replace(',', '.')
        return float(limpo) if limpo else None
    except ValueError:
        return None


def _ensure_list(data):
    """Garante que o dado seja uma lista de strings. Se for string única, separa por quebra de linha."""
    if not data:
        return []
    if isinstance(data, str):
        return [item.strip() for item in data.split('\n') if item.strip()]
    return data


# ═══════════════════════════════════════════════════════════
#  TÍTULO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════

def build_titulo(doc, titulo="PARECER SETOR TÉCNICO - SMOSU"):
    """Título centralizado com fonte Cambria bold."""
    pt = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line=288, before=200, after=100)
    r_tit = add_run(pt, titulo, bold=True, size=16)
    set_font(r_tit, name=FONT_TITULO, size=16, bold=True)
    r_tit.font.color.rgb = COR_LABEL_FONT
    # Sem separador aqui — o cabeçalho já tem linha azul


# ═══════════════════════════════════════════════════════════
#  IDENTIFICAÇÃO (Processo, Assunto, Requerente)
# ═══════════════════════════════════════════════════════════

def build_identificacao(doc, d):
    """Tabela de identificação dinâmica (Processo, Assunto, Requerente)."""
    
    # Montar lista de linhas candidatas filtrando valores vazios/nulos per D-04
    linhas = []
    
    # 1. Processo
    if d.get('numero_processo'):
        val_proc = f"{d.get('numero_processo', '')}"
        if d.get('data_processo'):
            val_proc += f"  —  {d.get('data_processo')}"
        linhas.append(("Processo nº", val_proc))
        
    # 2. Assunto
    if d.get("assunto"):
        linhas.append(("Assunto", d.get("assunto")))
        
    # 3. Requerente
    if d.get("requerente"):
        linhas.append(("Requerente", d.get("requerente")))

    # Se não houver nada, não gera tabela
    if not linhas:
        return

    # Criar tabela com o número exato de linhas filtradas
    tbl = doc.add_table(rows=len(linhas), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    for i, (rot, val) in enumerate(linhas):
        row = tbl.rows[i]
        c_rot = row.cells[0]
        apply_label_cell(c_rot, W_IDENT_LABEL)
        p_rot = c_rot.paragraphs[0]
        p_rot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_rot, line=280, before=70, after=70)
        r_rot = p_rot.add_run(rot)
        set_font(r_rot, size=10, bold=True)
        r_rot.font.color.rgb = COR_LABEL_FONT

        c_val = row.cells[1]
        apply_value_cell(c_val, W_IDENT_VALUE)
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_val, line=280, before=70, after=70)
        rich_segments(p_val, val, size=10)

    add_separator(doc, color='D0D0D0')   # único separador após identificação


# ═══════════════════════════════════════════════════════════
#  DESTINATÁRIO (para Ofícios)
# ═══════════════════════════════════════════════════════════

def build_destinatario(doc, d):
    """Bloco de destinatário para ofícios e comunicados."""
    if d.get("destinatario_titulo"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=200, after=40)
        add_run(p, d["destinatario_titulo"], bold=True, size=SZ_CORPO)

    if d.get("destinatario_para"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "Para: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_para"], size=SZ_CORPO)

    if d.get("destinatario_de"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "De: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_de"], size=SZ_CORPO)

    if d.get("assunto"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=80)
        add_run(p, "Assunto: ", bold=True, size=SZ_CORPO)
        add_run(p, d["assunto"], size=SZ_CORPO)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  DADOS DO CARIMBO TÉCNICO (Decreto 4.149/2019)
# ═══════════════════════════════════════════════════════════

def build_dados_carimbo(doc, d):
    """Tabela 4 colunas com dados técnicos do projeto."""
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=280, after=60)
    r_tit = add_run(p_tit, "DADOS TÉCNICOS DO PROJETO (Ref. Decreto nº 4.149/2019)",
                    bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    logradouro = d.get('logradouro', '')
    bairro = d.get('bairro', '')
    end_completo = (f"{logradouro} - {bairro}" if logradouro and bairro
                    else logradouro or bairro)

    # Detectar alertas urbanísticos
    to_val = _parse_numero(d.get('taxa_ocupacao', ''))
    perm_val = _parse_numero(d.get('taxa_permeabilidade', ''))
    terreno_val = _parse_numero(d.get('area_terreno', ''))
    alerta_zoneamento = (
        (to_val is not None and to_val > 70)
        and (perm_val is not None and perm_val < 20)
    )
    alerta_excecao = (terreno_val is not None and terreno_val < 220)

    linhas = [
        ("Endereço:",       end_completo,
         "Inscrição Mun.:", d.get("inscricao_municipal", "")),
        ("Proprietário:",   d.get("proprietario", d.get("requerente", "")),
         "Resp. Técnico:",  d.get("profissional_nome", d.get("responsavel_tecnico", ""))),
        ("Lote:",           d.get("lote", ""),
         "Quadra:",         d.get("quadra", "")),
        ("Área Terreno:",   d.get("area_terreno", ""),
         "Área Total:",     d.get("area_total_construida", "")),
        ("Taxa Ocupação:",  d.get("taxa_ocupacao", ""),
         "Coef. Aprov.:",   d.get("coef_aproveitamento", "")),
        ("Permeabilidade:", d.get("taxa_permeabilidade", ""),
         "Pavimentos:",     d.get("pavimentos", "")),
        ("Vagas Garagem:",  d.get("vagas_garagem", ""),
         "Zona de Uso:",    d.get("zona_uso", "")),
        ("ART/RRT:",        d.get("art_rrt_numero", d.get("art_rrt", "")),
         "Multa Específica:", d.get("tipo_multa_especifica", "")),
    ]

    W_L1 = W_CARIMBO_L1
    W_V1 = W_CARIMBO_V1
    W_L2 = W_CARIMBO_L2
    W_V2 = W_CARIMBO_V2

    for lab1, val1, lab2, val2 in linhas:
        row = tbl.add_row()

        # Label 1
        c0 = row.cells[0]
        apply_label_cell(c0, W_L1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p0, line=220, before=25, after=25)
        r0 = p0.add_run(lab1)
        set_font(r0, size=SZ_TABELA, bold=True)
        r0.font.color.rgb = COR_LABEL_FONT

        # Valor 1
        c1 = row.cells[1]
        apply_value_cell(c1, W_V1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p1, line=220, before=25, after=25)
        
        cor_v1 = None
        alerta_msg_v1 = None

        # Alerta: Área do Terreno < 220m² → verde + "exceção da lei"
        if lab1 == "Área Terreno:" and alerta_excecao:
            cor_v1 = COR_ALERTA_GREEN
            alerta_msg_v1 = "  (exceção da lei)"

        # Alerta: Taxa Ocupação > 70% → vermelho + "conferir zoneamento"
        if lab1 == "Taxa Ocupação:" and alerta_zoneamento:
            cor_v1 = COR_ALERTA_RED
            alerta_msg_v1 = "  (conferir zoneamento)"

        # Alerta: Permeabilidade < 20% → vermelho
        if lab1 == "Permeabilidade:" and alerta_zoneamento:
            cor_v1 = COR_ALERTA_RED
            alerta_msg_v1 = "  (conferir zoneamento)"

        rich_segments(p1, val1, size=SZ_TABELA, color=cor_v1)
        
        if alerta_msg_v1:
            r_alert = p1.add_run(alerta_msg_v1)
            set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = cor_v1

        # Label 2
        c2 = row.cells[2]
        apply_label_cell(c2, W_L2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p2, line=220, before=25, after=25)
        r2 = p2.add_run(lab2)
        set_font(r2, size=SZ_TABELA, bold=True)
        r2.font.color.rgb = COR_LABEL_FONT

        # Valor 2
        c3 = row.cells[3]
        apply_value_cell(c3, W_V2)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p3, line=220, before=25, after=25)
        rich_segments(p3, val2, size=SZ_TABELA)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  CORPO DO PARECER (abertura + considerandos)
# ═══════════════════════════════════════════════════════════

def build_corpo(doc, d):
    """Corpo do parecer: abertura, considerandos e parágrafos adicionais."""
    INDENT = INDENT_PADRAO

    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=300,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    if d.get("considerandos"):
        add_section_heading(doc, "Considerandos")
        for cons in _ensure_list(d["considerandos"]):
            cons_limpo = cons.strip()
            
            # 1. Remoção robusta de prefixos redundantes (Considerando que, Que, etc)
            # Remove "Considerando que " ou "Considerando " (case insensitive)
            cons_limpo = re.sub(r'^(considerando\s+(que\s+)?)+', '', cons_limpo, flags=re.IGNORECASE).strip()
            # Remove "Que " se tiver sobrado no início
            cons_limpo = re.sub(r'^que\s+', '', cons_limpo, flags=re.IGNORECASE).strip()
            
            # 2. Primeira letra em maiúscula se não for negrito ou marcador
            if cons_limpo and not cons_limpo.startswith('**') and not re.match(r'^(\d+\.|[A-Za-z]\.|\d+\s*\-)', cons_limpo):
                cons_limpo = cons_limpo[0].upper() + cons_limpo[1:]

            p = add_para(doc, line=LINE_SPC, before=0,
                         after=PAR_AFTER, indent_cm=INDENT)
            
            # 3. Inteligência de Prefixo: 
            # NÃO adiciona "Considerando que " se:
            # - O texto original já tinha (removido no passo 1, então checamos se o usuário forneceu um texto finalizado)
            # - Começa com numeração (1., 2.) ou letras (a., b.)
            # - Começa com negrito técnico (**Art.**)
            # - Começa com marcador de lista (- ou *)
            # - O usuário sinalizou que é um texto final
            ja_tem_prefixo = re.match(r'^(considerando|que)\b', cons.strip(), re.I)
            eh_lista = re.match(r'^(\d+\.|[A-Za-z]\.|\*\*|\d+\s*\-|\*|\-)', cons_limpo)
            
            if not ja_tem_prefixo and not eh_lista:
                add_run(p, "Considerando que ", bold=True, size=SZ_CORPO)
                
            rich_segments(p, cons_limpo, size=SZ_CORPO)

    if d.get("multas_aplicaveis"):
        add_section_heading(doc, "Multas e Penalidades Aplicáveis", fill='7B1A0A')
        for multa in _ensure_list(d["multas_aplicaveis"]):
            p_multa = add_para(doc, line=LINE_SPC, before=40, after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_multa, "▪  ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_ALERTA_RED
            rich_segments(p_multa, multa, size=SZ_CORPO)

    if d.get("condicionantes_aprovacao"):
        add_section_heading(doc, "Condicionantes de Aprovação", fill='1A5C1A')
        for cond in _ensure_list(d["condicionantes_aprovacao"]):
            p_cond = add_para(doc, line=LINE_SPC, before=40, after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_cond, "▪  ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_ALERTA_GREEN
            rich_segments(p_cond, cond, size=SZ_CORPO)

    if d.get("paragrafos_adicionais"):
        for txt in _ensure_list(d["paragrafos_adicionais"]):
            p_extra = add_para(doc, line=LINE_SPC, before=0,
                               after=PAR_AFTER, indent_cm=INDENT)
            rich_segments(p_extra, txt, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO LEGAL
# ═══════════════════════════════════════════════════════════

def build_fundamentacao(doc, d):
    """Seção de fundamentação legal com bullets e hierarquia correta."""
    if not d.get("fundamentacao_legal") and not d.get("tipo_relatorio") in ["alvara_aprovacao", "alvara_regularizacao", "alvara_ampliacao"]:
        return

    INDENT = INDENT_PADRAO
    add_section_heading(doc, "Da Análise Legal e Técnica", fill='1F3864') # Azul Institucional

    # Base legal padrão e simplificada — DECRETO 4.149 em primeiro lugar
    p_padrao = add_para(doc, line=LINE_SPC, before=80, after=PAR_AFTER, indent_cm=INDENT)
    r_bullet_padrao = add_run(p_padrao, "▪ ", size=SZ_CORPO, bold=True)
    r_bullet_padrao.font.color.rgb = COR_LABEL_FONT
    texto_padrao = "A presente análise técnica pautou-se primordialmente no **Decreto Municipal nº 4.149/2019**, bem como no **Código de Obras** (Lei nº 1.544/1986) e na **LUOS** (Lei Complementar nº 267/2019)."
    rich_segments(p_padrao, texto_padrao, size=SZ_CORPO)

    if d.get("fundamentacao_legal"):
        for fund in _ensure_list(d["fundamentacao_legal"]):
            fund_limpo = fund.lstrip("•- \t")
            # Se já for o Decreto ou Lei que citamos acima, pula para não duplicar
            if re.search(r'4\.149|1\.544|267', fund_limpo) and len(fund_limpo) < 60:
                continue
                
            p_fund = add_para(doc, line=LINE_SPC, before=0,
                              after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_fund, "▪ ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_LABEL_FONT
            rich_segments(p_fund, fund_limpo, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  CARDS DE DOCUMENTOS A EMITIR
# ═══════════════════════════════════════════════════════════

def add_doc_item(doc, tipo, obs=None):
    """Bloco de documento com borda lateral azul e fundo claríssimo."""
    from docx.shared import Pt, Cm

    # Linha do título: fundo azul claríssimo, bullet colorido, nome em negrito
    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_tipo.paragraph_format.left_indent  = Cm(0.3)
    p_tipo.paragraph_format.space_before = Pt(10)
    p_tipo.paragraph_format.space_after  = Pt(2)
    p_tipo.paragraph_format.line_spacing = Pt(15)

    # Fundo azul claríssimo via shading XML
    pPr = p_tipo._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), COR_DOC_BOX_FILL)
    pPr.append(shd)

    # Borda lateral esquerda azul institucional
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '12')
    left_b.set(qn('w:color'), COR_INST)
    left_b.set(qn('w:space'), '8')
    pBdr.append(left_b)
    pPr.append(pBdr)

    r_bullet = p_tipo.add_run("\u2713  ")
    set_font(r_bullet, size=11, bold=True)
    r_bullet.font.color.rgb = COR_LABEL_FONT

    r_tipo = p_tipo.add_run(tipo)
    set_font(r_tipo, size=11, bold=True)
    r_tipo.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Observa\u00e7\u00e3o: prefixo azul + texto it\u00e1lico cinza, recuado
    if obs:
        p_obs = doc.add_paragraph()
        p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_obs.paragraph_format.left_indent  = Cm(0.8)
        p_obs.paragraph_format.space_before = Pt(3)
        p_obs.paragraph_format.space_after  = Pt(10)
        p_obs.paragraph_format.line_spacing = Pt(14)

        r_prefix = p_obs.add_run("Obs.:  ")
        set_font(r_prefix, size=SZ_CORPO, bold=True)
        r_prefix.font.color.rgb = COR_LABEL_FONT

        rich_segments(p_obs, obs.strip(), size=SZ_CORPO, color=COR_CINZA_TEXTO, base_italic=True)


# ═══════════════════════════════════════════════════════════
#  CONCLUSÃO + DOCUMENTOS + ASSINATURA
# ═══════════════════════════════════════════════════════════

def _build_conclusao_bloco(doc, conclusao_text):
    """Heading 'CONCLUSÃO TÉCNICA' + parágrafo justificado de conclusão."""
    add_section_heading(doc, "Conclusão Técnica")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, line=LINE_SPC, before=100, after=80)
    p.paragraph_format.left_indent = Pt(18)
    rich_segments(p, conclusao_text, size=SZ_CORPO)


def build_conclusao_e_docs(doc, d):
    """Conclusão completa: fundamentação + conclusão + lista de docs + assinatura."""
    # 1º Fundamentação Legal
    build_fundamentacao(doc, d)

    # 2º Conclusão ("Diante do exposto...")
    conclusao = d.get(
        "conclusao",
        "Diante do exposto, após análise das documentações apresentadas, conclui-se que o "
        "projeto está apto a ser aprovado quanto às questões técnicas pertinentes, podendo "
        "ser emitidos os seguintes documentos:"
    )
    
    _build_conclusao_bloco(doc, conclusao)

    # 3º Título "Emissão de Documentos"
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        # 4º Lista de documentos
        _map_nomes = {
            "parecer_tecnico": "Parecer Técnico (SMOSU)",
            "habitese_comum": "Carta de Habite-se",
            "certidao_averbacao": "Certidão de Averbação de Área",
            "alvara_aprovacao": "Alvará de Construção",
            "alvara_regularizacao": "Alvará de Regularização As-Built",
            "comunicado_pendencia": "Comunicado de Pendência Documental"
        }
        for item in d.get("documentos_emitir", []):
            tipo_id = item if isinstance(item, str) else item.get("tipo", "")
            descricao = ""
            if isinstance(item, dict):
                descricao = item.get("descricao") or item.get("nome")
            
            if not descricao:
                descricao = _map_nomes.get(tipo_id, tipo_id.replace("_", " ").title())
                
            obs = item.get("obs") if isinstance(item, dict) else None
            add_doc_item(doc, descricao, obs)

    # 5º Assinatura
    build_assinatura(doc, d)


def build_conclusao_simples(doc, d):
    """Conclusão simples sem lista de documentos em cards."""
    if d.get("conclusao"):
        _build_conclusao_bloco(doc, d["conclusao"])

    # Se houver documentos
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        for item in d.get("documentos_emitir", []):
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  ASSINATURA
# ═══════════════════════════════════════════════════════════

def build_assinatura(doc, d):
    """Bloco de assinatura: data + linha tracejada + nome/título/registro."""
    assinante = d.get("assinante", ASSINANTE)
    cidade    = d.get("cidade", CIDADE)

    # Local e data
    p_local = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       line=LINE_SPC, before=280, after=0)
    add_run(p_local, f"{cidade}, {_data_hoje_extenso()}.", bold=False, size=SZ_CORPO)
    p_local.paragraph_format.keep_with_next = True

    # Espaço para assinar
    for _ in range(3):
        pb = doc.add_paragraph()
        set_spacing(pb, line=240, before=0, after=0)
        pb.paragraph_format.keep_with_next = True

    # Linha de assinatura via borda inferior do parágrafo (mais elegante que underscores)
    p_line = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=60, before=0, after=60)
    # Largura limitada via indentação
    p_line.paragraph_format.left_indent  = Pt(100)
    p_line.paragraph_format.right_indent = Pt(100)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), COR_INST)
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)
    p_line.paragraph_format.keep_with_next = True

    # Nome em negrito (destaque)
    nome = d.get("assinante_parecer") or assinante.get("nome", ASSINANTE["nome"])
    p_nome = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=LINE_SPC, before=40, after=0)
    r_nome = add_run(p_nome, nome, bold=True, size=SZ_CORPO)
    r_nome.font.color.rgb = COR_LABEL_FONT
    p_nome.paragraph_format.keep_with_next = True

    # Título e Registro em tamanho menor
    titulo   = assinante.get("titulo",   ASSINANTE["titulo"])
    registro = assinante.get("registro", ASSINANTE["registro"])
    for texto in [titulo, registro]:
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=240, before=0, after=0)
        r = add_run(p, texto, bold=False, size=9)
        r.font.color.rgb = COR_CINZA_TEXTO
        p.paragraph_format.keep_with_next = True

    # Secretaria
    p_sec = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=240, before=0, after=120)
    r_sec = add_run(p_sec, "Secretaria Municipal de Obras e Serviços Urbanos", bold=False, size=9)
    r_sec.font.color.rgb = COR_CINZA_TEXTO


# ═══════════════════════════════════════════════════════════
#  UTILITÁRIOS INTERNOS DE CÉLULA
# ═══════════════════════════════════════════════════════════

def _apply_cell_fill(cell, fill_hex: str):
    """Aplica cor de fundo a uma célula de tabela via XML."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_width(cell, width_twips: int):
    """Define largura de célula em twips via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ═══════════════════════════════════════════════════════════
#  PARTES E RESPONSÁVEIS DO PROCESSO
# ═══════════════════════════════════════════════════════════

def build_partes_envolvidas(doc, partes: dict):
    """
    Tabela 'PARTES E RESPONSÁVEIS' (2 colunas: PAPEL | IDENTIFICAÇÃO).
    Gerada quando o JSON contém o campo 'partes_envolvidas'.
    """
    if not partes:
        return

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=240, after=60)
    r_tit = add_run(p_tit, "PARTES E RESPONSÁVEIS DO PROCESSO", bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    linhas = []

    req = partes.get("requerente") or {}
    if isinstance(req, str):
        req = {"nome": req}
    nome_req = req.get("nome", "")
    qualidade_req = req.get("qualidade", "")
    if nome_req:
        val = f"{nome_req} ({qualidade_req})" if qualidade_req else nome_req
        linhas.append(("Requerente", val))

    prop = partes.get("proprietario") or {}
    if isinstance(prop, str):
        prop = {"nome": prop}
    nome_prop = prop.get("nome", "")
    mat_prop = prop.get("matricula_imovel", "")
    if nome_prop and nome_prop != nome_req:
        val = nome_prop + (f" — Matrícula nº {mat_prop}" if mat_prop else "")
        linhas.append(("Proprietário", val))

    rt = partes.get("responsavel_tecnico") or {}
    if isinstance(rt, str):
        rt = {"nome": rt}
    nome_rt = rt.get("nome", "")
    if nome_rt:
        conselho = rt.get("conselho", "")
        tipo_rt = rt.get("tipo_rt", "")
        num_rt = rt.get("numero_rt", "")
        val = nome_rt
        if conselho:
            val += f" — {conselho}"
        if tipo_rt and num_rt:
            val += f" — {tipo_rt} nº {num_rt}"
        linhas.append(("Resp. Técnico", val))

    for fiscal in partes.get("agentes_fiscais") or []:
        if isinstance(fiscal, str):
            linhas.append(("Fiscal Vistoriador", fiscal))
        else:
            nome_f = fiscal.get("nome", "")
            mat_f = fiscal.get("matricula_funcional", "")
            val = nome_f + (f" (Mat. {mat_f})" if mat_f else "")
            if nome_f:
                linhas.append(("Fiscal Vistoriador", val))

    ass = partes.get("assinante_parecer") or {}
    if isinstance(ass, str):
        ass = {"nome": ass}
    nome_ass = ass.get("nome", "")
    if nome_ass:
        titulo_ass = ass.get("titulo", "")
        reg_ass = ass.get("registro", "")
        val = nome_ass
        if titulo_ass:
            val += f" — {titulo_ass}"
        if reg_ass:
            val += f" — {reg_ass}"
        linhas.append(("Signatário", val))

    if not linhas:
        return

    tbl = doc.add_table(rows=len(linhas), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    for i, (papel, identificacao) in enumerate(linhas):
        fill_hex = 'EEF2FB' if i % 2 == 0 else 'FFFFFF'
        row = tbl.rows[i]

        c_papel = row.cells[0]
        _set_cell_width(c_papel, W_PARTES_LABEL)
        _apply_cell_fill(c_papel, fill_hex)
        set_cell_margins(c_papel, top=40, bottom=40, left=80, right=40)
        p_papel = c_papel.paragraphs[0]
        p_papel.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_papel, line=220, before=20, after=20)
        r_p = p_papel.add_run(papel)
        set_font(r_p, size=SZ_TABELA, bold=True)
        r_p.font.color.rgb = COR_LABEL_FONT

        c_id = row.cells[1]
        _set_cell_width(c_id, W_PARTES_VALUE)
        _apply_cell_fill(c_id, fill_hex)
        set_cell_margins(c_id, top=40, bottom=40, left=80, right=40)
        p_id = c_id.paragraphs[0]
        p_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_id, line=220, before=20, after=20)
        rich_segments(p_id, identificacao, size=SZ_TABELA)

    add_separator(doc, color='D0D0D0')


# ═══════════════════════════════════════════════════════════
#  HISTÓRICO CRONOLÓGICO DO PROCESSO
# ═══════════════════════════════════════════════════════════

def build_historico_cronologico(doc, historico: list):
    """
    Tabela 'HISTÓRICO CRONOLÓGICO' (3 colunas: DATA | EVENTO | REFERÊNCIA/AGENTES).
    Gerada quando o JSON contém o campo 'historico_cronologico'.
    Os eventos devem estar em ordem cronológica (do mais antigo ao mais recente).
    """
    if not historico:
        return

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=240, after=60)
    r_tit = add_run(p_tit, "HISTÓRICO CRONOLÓGICO DO PROCESSO", bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=len(historico) + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    # Cabeçalho
    headers = [("DATA", W_HIST_DATA, WD_ALIGN_PARAGRAPH.CENTER),
               ("EVENTO / DESCRIÇÃO", W_HIST_EVENT, WD_ALIGN_PARAGRAPH.LEFT),
               ("REFERÊNCIA / AGENTES", W_HIST_REF, WD_ALIGN_PARAGRAPH.LEFT)]
    hrow = tbl.rows[0]
    for j, (hdr, w, align) in enumerate(headers):
        cell = hrow.cells[j]
        _set_cell_width(cell, w)
        _apply_cell_fill(cell, COR_INST)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=40)
        p = cell.paragraphs[0]
        p.alignment = align
        set_spacing(p, line=200, before=20, after=20)
        r = p.add_run(hdr)
        set_font(r, size=SZ_TABELA, bold=True)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas de dados
    for i, ev in enumerate(historico):
        fill_hex = 'F0F4FD' if i % 2 == 0 else 'FFFFFF'
        row = tbl.rows[i + 1]

        # DATA
        c_data = row.cells[0]
        _set_cell_width(c_data, W_HIST_DATA)
        _apply_cell_fill(c_data, fill_hex)
        set_cell_margins(c_data, top=40, bottom=40, left=80, right=40)
        p_d = c_data.paragraphs[0]
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p_d, line=220, before=20, after=20)
        r_d = p_d.add_run(str(ev.get("data", "")))
        set_font(r_d, size=SZ_TABELA, bold=True)
        r_d.font.color.rgb = COR_LABEL_FONT

        # EVENTO
        c_ev = row.cells[1]
        _set_cell_width(c_ev, W_HIST_EVENT)
        _apply_cell_fill(c_ev, fill_hex)
        set_cell_margins(c_ev, top=40, bottom=40, left=80, right=40)
        p_ev = c_ev.paragraphs[0]
        p_ev.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_ev, line=220, before=20, after=20)
        rich_segments(p_ev, str(ev.get("evento", "")), size=SZ_TABELA)

        # REFERÊNCIA / AGENTES
        c_ref = row.cells[2]
        _set_cell_width(c_ref, W_HIST_REF)
        _apply_cell_fill(c_ref, fill_hex)
        set_cell_margins(c_ref, top=40, bottom=40, left=80, right=40)
        p_ref = c_ref.paragraphs[0]
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_ref, line=220, before=20, after=20)

        agentes = ev.get("agentes", [])
        ref_text = ev.get("referencia", "")
        if agentes:
            if isinstance(agentes, list):
                ref_text = "; ".join(str(a) for a in agentes)
            else:
                ref_text = str(agentes)

        rich_segments(p_ref, ref_text, size=SZ_TABELA)

    add_separator(doc, color='D0D0D0')


# ═══════════════════════════════════════════════════════════
#  COMUNICADO DE PENDÊNCIA — Design Visual Dedicado
# ═══════════════════════════════════════════════════════════

def _box_colorido(doc, fill_hex, borda_hex, borda_sz='12'):
    """Cria uma tabela-card de 1 célula com fundo colorido e borda lateral."""
    card = doc.add_table(rows=1, cols=1)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        card._tbl.insert(0, tblPr)

    # Largura total
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(W_CARD_TOTAL))   # ~16,5 cm em twips
    tblW.set(qn('w:type'), 'dxa')
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)

    # Bordas
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), borda_hex)
        el.set(qn('w:space'), '0')
        tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), borda_sz)
    left_b.set(qn('w:color'), borda_hex)
    left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=100)

    return card, cell


def build_comunicado_pendencia(doc, d):
    """
    Layout visual dedicado para Comunicado de Pendência Documental.

    Estrutura:
      ① Parágrafo de abertura (texto institucional normal)
      ② Box LARANJA — lista numerada de pendências
      ③ Box VERDE   — orientação de continuidade ao requerente
      ④ Conclusão
      ⑤ Assinatura
    """
    INDENT = INDENT_PADRAO

    # ── ① Parágrafo de abertura ────────────────────────────────────────────
    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=240,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    # ── ② Box de Pendências (laranja/âmbar) ───────────────────────────────
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    card_alert, cell_alert = _box_colorido(doc, COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA, '18')

    # Título do box
    p_tit = cell_alert.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=260, before=60, after=60)
    r_icon = p_tit.add_run("⚠  DOCUMENTOS PENDENTES — ANÁLISE SUSPENSA")
    set_font(r_icon, name="Cambria", size=11, bold=True)
    r_icon.font.color.rgb = COR_PENDENCIA_ICON

    # Subtítulo
    p_sub = cell_alert.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_sub, line=240, before=0, after=80)
    r_sub = p_sub.add_run(
        "A análise técnica não pôde ser iniciada. "
        "Os documentos abaixo são imprescindíveis e devem ser apresentados "
        "para o prosseguimento do processo:"
    )
    set_font(r_sub, size=10, italic=True)
    r_sub.font.color.rgb = COR_PENDENCIA_ICON

    # Lista de pendências
    considerandos = _ensure_list(d.get("considerandos", []))
    for item in considerandos:
        p_item = cell_alert.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_item, line=260, before=20, after=20)
        # Ícone de marcador
        r_bullet = p_item.add_run("  ✗  ")
        set_font(r_bullet, size=10, bold=True)
        r_bullet.font.color.rgb = COR_PENDENCIA_BORDA # Use orange for bullet
        # Texto do item (sem "Considerando que" — é uma lista de pendências)
        texto = item.lstrip("0123456789. ")  # remove "1. ", "2. " se houver
        if texto.lower().startswith("considerando que "):
            texto = texto[17:].strip()
        elif texto.lower().startswith("considerando "):
            texto = texto[13:].strip()
            
        rich_segments(p_item, texto, size=10, color=COR_PENDENCIA_TEXTO)

    # Espaço após o box
    p_esp2 = doc.add_paragraph()
    set_spacing(p_esp2, line=80, before=0, after=0)

    # ── ③ Box de Orientação (verde) ────────────────────────────────────────
    card_ok, cell_ok = _box_colorido(doc, COR_SUCESSO_FILL, COR_SUCESSO_BORDA, '18')

    p_ok_tit = cell_ok.paragraphs[0]
    p_ok_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_ok_tit, line=260, before=60, after=40)
    r_ok_icon = p_ok_tit.add_run("✔  COMO REGULARIZAR E RETOMAR O PROCESSO")
    set_font(r_ok_icon, name="Cambria", size=11, bold=True)
    r_ok_icon.font.color.rgb = COR_SUCESSO_ICON

    orientacoes = [
        "Reúna todos os documentos listados acima em sua versão original ou digitalizada com legibilidade plena.",
        "Dirija-se ao balcão do protocolo da Prefeitura Municipal de Oliveira ou acesse o portal digital (Atende.Net) para acostar os documentos ao processo.",
        "Após o recebimento e verificação documental, a análise técnica será reiniciada automaticamente, sem necessidade de novo requerimento.",
        "Em caso de dúvidas, entre em contato com o Setor Técnico da SMOSU (Secretaria Municipal de Obras e Serviços Urbanos).",
    ]
    for i, orient in enumerate(orientacoes, 1):
        p_or = cell_ok.add_paragraph()
        p_or.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_or, line=260, before=20, after=20)
        r_num = p_or.add_run(f"  {i}.  ")
        set_font(r_num, size=10, bold=True)
        r_num.font.color.rgb = COR_SUCESSO_ICON
        r_or = p_or.add_run(orient)
        set_font(r_or, size=10)
        r_or.font.color.rgb = COR_SUCESSO_TEXTO

    # Espaço após box verde
    p_esp3 = doc.add_paragraph()
    set_spacing(p_esp3, line=120, before=0, after=0)

    # ── ④ Conclusão ────────────────────────────────────────────────────────
    if d.get("conclusao"):
        p_conc = add_para(doc, line=LINE_SPC, before=120,
                          after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # ── ⑤ Assinatura ───────────────────────────────────────────────────────
    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  MEMÓRIA DE CÁLCULO
# ═══════════════════════════════════════════════════════════

def build_memoria_calculo(doc, d):
    """
    Renderiza a seção de Memória de Cálculo em caixa sombreada.
    Gerado quando o JSON contém o campo 'memoria_de_calculo'.
    """
    memoria = d.get("memoria_de_calculo")
    if not memoria:
        return

    add_section_heading(doc, "Memória de Cálculo e Índices Urbanísticos")

    # Caixa cinza-azulada para isolar o bloco de cálculo
    card, cell = _box_colorido(doc, 'EFF2F8', 'B0BAD0', '8')

    # Substituir os parágrafos internos por linhas da memória
    for p_old in list(cell.paragraphs[1:]):
        p_old._element.getparent().remove(p_old._element)

    p_mem = cell.paragraphs[0]
    p_mem.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_mem, line=LINE_SPC, before=0, after=0)
    rich_segments(p_mem, memoria, size=SZ_TABELA)

    # Para cada quebra de linha no texto, cria parágrafo separado
    linhas = [l.strip() for l in memoria.split('\n') if l.strip()]
    if len(linhas) > 1:
        # Substitui o parágrafo único por múltiplos parágrafos
        p_mem.clear()
        rich_segments(p_mem, linhas[0], size=SZ_TABELA)
        for linha in linhas[1:]:
            p_extra = cell.add_paragraph()
            p_extra.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p_extra, line=LINE_SPC, before=0, after=0)
            rich_segments(p_extra, linha, size=SZ_TABELA)

    # Respiro após a caixa
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    add_separator(doc, color='D0D0D0')


