import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

from generators.componentes.assinatura import build_assinatura
from generators.componentes.corpo import _ensure_list

def build_fundamentacao(doc, d):
    """Seção de fundamentação legal com bullets e hierarquia correta."""
    if not d.get("fundamentacao_legal") and not d.get("tipo_relatorio") in ["alvara_aprovacao", "alvara_regularizacao", "alvara_ampliacao"]:
        return

    INDENT = INDENT_PADRAO
    add_section_heading(doc, "Da Análise Legal e Técnica", fill='1F3864') # Azul Institucional

    # Base legal padrão e simplificada — DECRETO 4.149 em primeiro lugar
    p_padrao = add_para(doc, line=LINE_SPC, before=80, after=PAR_AFTER, indent_cm=INDENT)
    r_bullet_padrao = add_run(p_padrao, "▪ ", size=SZ_CORPO, bold=True)
    r_bullet_padrao.font.color.rgb = COR_LABEL_FONT
    texto_padrao = "A presente análise técnica pautou-se primordialmente no **Decreto Municipal nº 4.149/2019**, bem como no **Código de Obras** (Lei nº 1.544/1986) e na **LUOS** (Lei Complementar nº 267/2019)."
    rich_segments(p_padrao, texto_padrao, size=SZ_CORPO)

    if d.get("fundamentacao_legal"):
        for fund in _ensure_list(d["fundamentacao_legal"]):
            fund_limpo = fund.lstrip("•- \t")
            # Se já for o Decreto ou Lei que citamos acima, pula para não duplicar
            if re.search(r'4\.149|1\.544|267', fund_limpo) and len(fund_limpo) < 60:
                continue
                
            p_fund = add_para(doc, line=LINE_SPC, before=0,
                              after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_fund, "▪ ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_LABEL_FONT
            rich_segments(p_fund, fund_limpo, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  CARDS DE DOCUMENTOS A EMITIR
# ═══════════════════════════════════════════════════════════

def add_doc_item(doc, tipo, obs=None):
    """Bloco de documento com borda lateral azul e fundo claríssimo."""
    from docx.shared import Pt, Cm

    # Linha do título: fundo azul claríssimo, bullet colorido, nome em negrito
    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_tipo.paragraph_format.left_indent  = Cm(0.3)
    p_tipo.paragraph_format.space_before = Pt(10)
    p_tipo.paragraph_format.space_after  = Pt(2)
    p_tipo.paragraph_format.line_spacing = Pt(15)

    # Fundo azul claríssimo via shading XML
    pPr = p_tipo._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), COR_DOC_BOX_FILL)
    pPr.append(shd)

    # Borda lateral esquerda azul institucional
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '12')
    left_b.set(qn('w:color'), COR_INST)
    left_b.set(qn('w:space'), '8')
    pBdr.append(left_b)
    pPr.append(pBdr)

    r_bullet = p_tipo.add_run("\u2713  ")
    set_font(r_bullet, size=11, bold=True)
    r_bullet.font.color.rgb = COR_LABEL_FONT

    r_tipo = p_tipo.add_run(tipo)
    set_font(r_tipo, size=11, bold=True)
    r_tipo.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Observa\u00e7\u00e3o: prefixo azul + texto it\u00e1lico cinza, recuado
    if obs:
        p_obs = doc.add_paragraph()
        p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_obs.paragraph_format.left_indent  = Cm(0.8)
        p_obs.paragraph_format.space_before = Pt(3)
        p_obs.paragraph_format.space_after  = Pt(10)
        p_obs.paragraph_format.line_spacing = Pt(14)

        r_prefix = p_obs.add_run("Obs.:  ")
        set_font(r_prefix, size=SZ_CORPO, bold=True)
        r_prefix.font.color.rgb = COR_LABEL_FONT

        rich_segments(p_obs, obs.strip(), size=SZ_CORPO, color=COR_CINZA_TEXTO, base_italic=True)


# ═══════════════════════════════════════════════════════════
#  CONCLUSÃO + DOCUMENTOS + ASSINATURA
# ═══════════════════════════════════════════════════════════

def _build_conclusao_bloco(doc, conclusao_text):
    """Heading 'CONCLUSÃO TÉCNICA' + parágrafo justificado de conclusão."""
    add_section_heading(doc, "Conclusão Técnica")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, line=LINE_SPC, before=100, after=80)
    p.paragraph_format.left_indent = Pt(18)
    rich_segments(p, conclusao_text, size=SZ_CORPO)


def _build_merito_bloco(doc, merito_text):
    """Heading 'ANÁLISE DE MÉRITO TÉCNICO' + texto em itálico com borda lateral."""
    add_section_heading(doc, "Análise de Mérito Técnico", fill='4f9cff') # Azul Acento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, line=LINE_SPC, before=100, after=100)
    p.paragraph_format.left_indent = Pt(18)
    
    # Adiciona borda lateral sutil para destacar o mérito
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '4')
    left_b.set(qn('w:color'), '4f9cff')
    left_b.set(qn('w:space'), '8')
    pBdr.append(left_b)
    pPr.append(pBdr)

    rich_segments(p, merito_text, size=SZ_CORPO, base_italic=True)


def build_conclusao_e_docs(doc, d):
    """Conclusão completa: fundamentação + conclusão + lista de docs + assinatura."""
    # 1º Fundamentação Legal
    build_fundamentacao(doc, d)

    # 2º Análise de Mérito (v3.0)
    if d.get("analise_merito"):
        _build_merito_bloco(doc, d["analise_merito"])

    # 3º Conclusão ("Diante do exposto...")
    conclusao = d.get(
        "conclusao",
        "Diante do exposto, após análise das documentações apresentadas, conclui-se que o "
        "projeto está apto a ser aprovado quanto às questões técnicas pertinentes, podendo "
        "ser emitidos os seguintes documentos:"
    )
    
    _build_conclusao_bloco(doc, conclusao)

    # 3º Título "Emissão de Documentos"
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        # 4º Lista de documentos
        _map_nomes = {
            "parecer_tecnico": "Parecer Técnico (SMOSU)",
            "habitese_comum": "Carta de Habite-se",
            "certidao_averbacao": "Certidão de Averbação de Área",
            "alvara_aprovacao": "Alvará de Construção",
            "alvara_regularizacao": "Alvará de Regularização As-Built",
            "comunicado_pendencia": "Comunicado de Pendência Documental"
        }
        for item in d.get("documentos_emitir", []):
            tipo_id = item if isinstance(item, str) else item.get("tipo", "")
            descricao = ""
            if isinstance(item, dict):
                descricao = item.get("descricao") or item.get("nome")
            
            if not descricao:
                descricao = _map_nomes.get(tipo_id, tipo_id.replace("_", " ").title())
                
            obs = (item.get("obs") or item.get("observacao")) if isinstance(item, dict) else None
            add_doc_item(doc, descricao, obs)

    # 5º Assinatura
    build_assinatura(doc, d)


def build_conclusao_simples(doc, d):
    """Conclusão simples sem lista de documentos em cards."""
    if d.get("conclusao"):
        _build_conclusao_bloco(doc, d["conclusao"])

    # Se houver documentos
    if d.get("documentos_emitir"):
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        for item in d.get("documentos_emitir", []):
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or item.get("observacao") or None)

    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  ASSINATURA
# ═══════════════════════════════════════════════════════════

