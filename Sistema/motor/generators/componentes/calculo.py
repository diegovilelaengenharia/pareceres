import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

from generators.componentes.comunicado import _box_colorido

def build_memoria_calculo(doc, d):
    """
    Renderiza a seção de Memória de Cálculo em caixa sombreada.
    Gerado quando o JSON contém o campo 'memoria_de_calculo'.
    """
    memoria = d.get("memoria_de_calculo")
    if not memoria:
        return

    add_section_heading(doc, "Memória de Cálculo e Índices Urbanísticos")

    # Caixa cinza-azulada para isolar o bloco de cálculo
    card, cell = _box_colorido(doc, 'EFF2F8', 'B0BAD0', '8')

    # Substituir os parágrafos internos por linhas da memória
    for p_old in list(cell.paragraphs[1:]):
        p_old._element.getparent().remove(p_old._element)

    p_mem = cell.paragraphs[0]
    p_mem.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_mem, line=LINE_SPC, before=0, after=0)
    rich_segments(p_mem, memoria, size=SZ_TABELA)

    # Para cada quebra de linha no texto, cria parágrafo separado
    linhas = [l.strip() for l in memoria.split('\n') if l.strip()]
    if len(linhas) > 1:
        # Substitui o parágrafo único por múltiplos parágrafos
        p_mem.clear()
        rich_segments(p_mem, linhas[0], size=SZ_TABELA)
        for linha in linhas[1:]:
            p_extra = cell.add_paragraph()
            p_extra.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p_extra, line=LINE_SPC, before=0, after=0)
            rich_segments(p_extra, linha, size=SZ_TABELA)

    # Respiro após a caixa
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    add_separator(doc, color='D0D0D0')


