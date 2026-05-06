import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

from generators.componentes.corpo import _ensure_list
from generators.componentes.assinatura import build_assinatura

def _box_colorido(doc, fill_hex, borda_hex, borda_sz='12'):
    """Cria uma tabela-card de 1 célula com fundo colorido e borda lateral."""
    card = doc.add_table(rows=1, cols=1)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        card._tbl.insert(0, tblPr)

    # Largura total
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(W_CARD_TOTAL))   # ~16,5 cm em twips
    tblW.set(qn('w:type'), 'dxa')
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)

    # Bordas
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), borda_hex)
        el.set(qn('w:space'), '0')
        tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), borda_sz)
    left_b.set(qn('w:color'), borda_hex)
    left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=100)

    return card, cell


def build_comunicado_pendencia(doc, d):
    """
    Layout visual dedicado para Comunicado de Pendência Documental.

    Estrutura:
      ① Parágrafo de abertura (texto institucional normal)
      ② Box LARANJA — lista numerada de pendências
      ③ Box VERDE   — orientação de continuidade ao requerente
      ④ Conclusão
      ⑤ Assinatura
    """
    INDENT = INDENT_PADRAO

    # ── ① Parágrafo de abertura ────────────────────────────────────────────
    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=240,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    # ── ② Box de Pendências (laranja/âmbar) ───────────────────────────────
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    card_alert, cell_alert = _box_colorido(doc, COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA, '18')

    # Título do box
    p_tit = cell_alert.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=260, before=60, after=60)
    r_icon = p_tit.add_run("⚠  DOCUMENTOS PENDENTES — ANÁLISE SUSPENSA")
    set_font(r_icon, name="Cambria", size=11, bold=True)
    r_icon.font.color.rgb = COR_PENDENCIA_ICON

    # Subtítulo
    p_sub = cell_alert.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_sub, line=240, before=0, after=80)
    r_sub = p_sub.add_run(
        "A análise técnica não pôde ser iniciada. "
        "Os documentos abaixo são imprescindíveis e devem ser apresentados "
        "para o prosseguimento do processo:"
    )
    set_font(r_sub, size=10, italic=True)
    r_sub.font.color.rgb = COR_PENDENCIA_ICON

    # Lista de pendências
    considerandos = _ensure_list(d.get("considerandos", []))
    for item in considerandos:
        p_item = cell_alert.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_item, line=260, before=20, after=20)
        # Ícone de marcador
        r_bullet = p_item.add_run("  ✗  ")
        set_font(r_bullet, size=10, bold=True)
        r_bullet.font.color.rgb = RGBColor(0xF2, 0xC9, 0x4C)
        # Texto do item (sem "Considerando que" — é uma lista de pendências)
        texto = item.lstrip("0123456789. ")  # remove "1. ", "2. " se houver
        if texto.lower().startswith("considerando que "):
            texto = texto[17:].strip()
        elif texto.lower().startswith("considerando "):
            texto = texto[13:].strip()
            
        rich_segments(p_item, texto, size=10, color=COR_PENDENCIA_TEXTO)

    # Espaço após o box
    p_esp2 = doc.add_paragraph()
    set_spacing(p_esp2, line=80, before=0, after=0)

    # ── ③ Box de Orientação (verde) ────────────────────────────────────────
    card_ok, cell_ok = _box_colorido(doc, COR_SUCESSO_FILL, COR_SUCESSO_BORDA, '18')

    p_ok_tit = cell_ok.paragraphs[0]
    p_ok_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_ok_tit, line=260, before=60, after=40)
    r_ok_icon = p_ok_tit.add_run("✔  COMO REGULARIZAR E RETOMAR O PROCESSO")
    set_font(r_ok_icon, name="Cambria", size=11, bold=True)
    r_ok_icon.font.color.rgb = COR_SUCESSO_ICON

    orientacoes = [
        "Reúna todos os documentos listados acima em sua versão original ou digitalizada com legibilidade plena.",
        "Dirija-se ao balcão do protocolo da Prefeitura Municipal de Oliveira ou acesse o portal digital (Atende.Net) para acostar os documentos ao processo.",
        "Após o recebimento e verificação documental, a análise técnica será reiniciada automaticamente, sem necessidade de novo requerimento.",
        "Em caso de dúvidas, entre em contato com o Setor Técnico da SMOSU (Secretaria Municipal de Obras e Serviços Urbanos).",
    ]
    for i, orient in enumerate(orientacoes, 1):
        p_or = cell_ok.add_paragraph()
        p_or.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_or, line=260, before=20, after=20)
        r_num = p_or.add_run(f"  {i}.  ")
        set_font(r_num, size=10, bold=True)
        r_num.font.color.rgb = COR_SUCESSO_ICON
        r_or = p_or.add_run(orient)
        set_font(r_or, size=10)
        r_or.font.color.rgb = COR_SUCESSO_TEXTO

    # Espaço após box verde
    p_esp3 = doc.add_paragraph()
    set_spacing(p_esp3, line=120, before=0, after=0)

    # ── ④ Conclusão ────────────────────────────────────────────────────────
    if d.get("conclusao"):
        p_conc = add_para(doc, line=LINE_SPC, before=120,
                          after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # ── ⑤ Assinatura ───────────────────────────────────────────────────────
    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  MEMÓRIA DE CÁLCULO
# ═══════════════════════════════════════════════════════════

