import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

def _ensure_list(data):
    """Garante que o dado seja uma lista de strings. Se for string única, separa por quebra de linha."""
    if not data:
        return []
    if isinstance(data, str):
        return [item.strip() for item in data.split('\n') if item.strip()]
    return data


# ═══════════════════════════════════════════════════════════
#  TÍTULO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════

def build_corpo(doc, d):
    """Corpo do parecer: abertura, considerandos e parágrafos adicionais."""
    INDENT = INDENT_PADRAO

    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=300,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    if d.get("considerandos"):
        add_section_heading(doc, "Considerandos")
        for cons in _ensure_list(d["considerandos"]):
            cons_limpo = cons.strip()
            
            # 1. Remoção robusta de prefixos redundantes (Considerando que, Que, etc)
            # Remove "Considerando que " ou "Considerando " (case insensitive)
            cons_limpo = re.sub(r'^(considerando\s+(que\s+)?)+', '', cons_limpo, flags=re.IGNORECASE).strip()
            # Remove "Que " se tiver sobrado no início
            cons_limpo = re.sub(r'^que\s+', '', cons_limpo, flags=re.IGNORECASE).strip()
            
            # 2. Primeira letra em maiúscula se não for negrito ou marcador
            if cons_limpo and not cons_limpo.startswith('**') and not re.match(r'^(\d+\.|[A-Za-z]\.|\d+\s*\-)', cons_limpo):
                cons_limpo = cons_limpo[0].upper() + cons_limpo[1:]

            p = add_para(doc, line=LINE_SPC, before=0,
                         after=PAR_AFTER, indent_cm=INDENT)
            
            # 3. Inteligência de Prefixo: 
            # NÃO adiciona "Considerando que " se:
            # - O texto original já tinha (verificamos antes da limpeza regex)
            # - O texto (após limpeza) começa com letra MAIÚSCULA (indica frase independente)
            # - O campo 'paragrafo_abertura' está presente (IA optou por narrativa fluida)
            # - Começa com numeração (1., 2.) ou letras (a., b.)
            # - Começa com negrito técnico (**Art.**)
            # - Começa com marcador de lista (- ou *)
            
            ja_tem_prefixo = re.match(r'^considerando\b', cons.strip(), re.I)
            eh_lista = re.match(r'^(\d+\.|[A-Za-z]\.|\*\*|\d+\s*\-|\*|\-)', cons_limpo)
            comeca_maiuscula = cons_limpo and cons_limpo[0].isupper() and not cons_limpo.startswith('**')
            tem_abertura = bool(d.get("paragrafo_abertura"))

            # Inteligência de Prefixo v2.0 (Motor de Pareceres Robustos)
            if not ja_tem_prefixo and not eh_lista and not comeca_maiuscula:
                add_run(p, "Considerando que ", bold=True, size=SZ_CORPO)
                
            rich_segments(p, cons_limpo, size=SZ_CORPO)

    if d.get("multas_aplicaveis"):
        add_section_heading(doc, "Multas e Penalidades Aplicáveis", fill='7B1A0A')
        for multa in _ensure_list(d["multas_aplicaveis"]):
            p_multa = add_para(doc, line=LINE_SPC, before=40, after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_multa, "▪  ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_ALERTA_RED
            rich_segments(p_multa, multa, size=SZ_CORPO)

    if d.get("condicionantes_aprovacao"):
        add_section_heading(doc, "Condicionantes de Aprovação", fill='1A5C1A')
        for cond in _ensure_list(d["condicionantes_aprovacao"]):
            p_cond = add_para(doc, line=LINE_SPC, before=40, after=PAR_AFTER, indent_cm=INDENT)
            r_bullet = add_run(p_cond, "▪  ", size=SZ_CORPO, bold=True)
            r_bullet.font.color.rgb = COR_ALERTA_GREEN
            rich_segments(p_cond, cond, size=SZ_CORPO)

    if d.get("paragrafos_adicionais"):
        for txt in _ensure_list(d["paragrafos_adicionais"]):
            p_extra = add_para(doc, line=LINE_SPC, before=0,
                               after=PAR_AFTER, indent_cm=INDENT)
            rich_segments(p_extra, txt, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO LEGAL
# ═══════════════════════════════════════════════════════════

