import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

def _data_hoje_extenso():
    """Retorna data atual por extenso: '20 de abril de 2026'."""
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]
    d = datetime.now()
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _parse_numero(texto):
    """Extrai valor numérico de strings como '86,23%' ou '180,00m²' → float."""
    if not texto:
        return None
    try:
        limpo = re.sub(r'[^\d,.]', '', str(texto)).replace(',', '.')
        return float(limpo) if limpo else None
    except ValueError:
        return None


def build_titulo(doc, titulo="PARECER SETOR TÉCNICO - SMOSU"):
    """Título centralizado com fonte Cambria bold."""
    pt = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line=288, before=200, after=100)
    r_tit = add_run(pt, titulo, bold=True, size=16)
    set_font(r_tit, name=FONT_TITULO, size=16, bold=True)
    r_tit.font.color.rgb = COR_LABEL_FONT
    # Sem separador aqui — o cabeçalho já tem linha azul


# ═══════════════════════════════════════════════════════════
#  IDENTIFICAÇÃO (Processo, Assunto, Requerente)
# ═══════════════════════════════════════════════════════════

def build_identificacao(doc, d):
    """Tabela de identificação dinâmica (Processo, Assunto, Requerente)."""
    
    # Montar lista de linhas candidatas filtrando valores vazios/nulos per D-04
    linhas = []
    
    # 1. Processo
    if d.get('numero_processo'):
        val_proc = f"{d.get('numero_processo', '')}"
        if d.get('data_processo'):
            val_proc += f" de {d.get('data_processo')}"
        linhas.append(("Processo nº", val_proc))
        
    # 2. Assunto
    if d.get("assunto"):
        linhas.append(("Assunto", d.get("assunto")))
        
    # 3. Requerente
    if d.get("requerente"):
        linhas.append(("Requerente", d.get("requerente")))

    # Se não houver nada, não gera nada
    if not linhas:
        return

    # Estilo "Dotted" solicitado para Pareceres Administrativos/Retificação
    if d.get("tipo_relatorio") == "certidoes_separadas_localizacao_confrontacao" or "administrativo" in d.get("tipo_relatorio", ""):
        for rot, val in linhas:
            p = add_para(doc, line=240, before=60, after=60)
            # Rótulo com pontos
            r_rot = p.add_run(f"{rot: <15}: \t")
            set_font(r_rot, size=10, bold=True)
            
            # Adicionar pontos manualmente para simular o modelo
            dots = "." * (60 - len(rot))
            r_dots = p.add_run(dots + " ")
            set_font(r_dots, size=10, bold=False)
            
            rich_segments(p, val, size=10)
        add_separator(doc, color='D0D0D0')
        return

    # Criar tabela padrão (para outros tipos)
    tbl = doc.add_table(rows=len(linhas), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    for i, (rot, val) in enumerate(linhas):
        row = tbl.rows[i]
        c_rot = row.cells[0]
        apply_label_cell(c_rot, W_IDENT_LABEL)
        p_rot = c_rot.paragraphs[0]
        p_rot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_rot, line=280, before=70, after=70)
        r_rot = p_rot.add_run(rot)
        set_font(r_rot, size=10, bold=True)
        r_rot.font.color.rgb = COR_LABEL_FONT

        c_val = row.cells[1]
        apply_value_cell(c_val, W_IDENT_VALUE)
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_val, line=280, before=70, after=70)
        rich_segments(p_val, val, size=10)

    add_separator(doc, color='D0D0D0')   # único separador após identificação


# ═══════════════════════════════════════════════════════════
#  DESTINATÁRIO (para Ofícios)
# ═══════════════════════════════════════════════════════════

def build_destinatario(doc, d):
    """Bloco de destinatário para ofícios e comunicados."""
    if d.get("destinatario_titulo"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=200, after=40)
        add_run(p, d["destinatario_titulo"], bold=True, size=SZ_CORPO)

    if d.get("destinatario_para"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "Para: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_para"], size=SZ_CORPO)

    if d.get("destinatario_de"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "De: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_de"], size=SZ_CORPO)

    if d.get("assunto"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=80)
        add_run(p, "Assunto: ", bold=True, size=SZ_CORPO)
        add_run(p, d["assunto"], size=SZ_CORPO)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  DADOS DO CARIMBO TÉCNICO (Decreto 4.149/2019)
# ═══════════════════════════════════════════════════════════

def build_dados_carimbo(doc, d):
    """Tabela Inteligente: Oculta campos vazios e agrupa os preenchidos em 4 colunas."""
    # 1. Preparação dos dados e alertas
    logradouro = d.get('logradouro', '')
    bairro = d.get('bairro', '')
    end_completo = (f"{logradouro} - {bairro}" if logradouro and bairro
                    else logradouro or bairro)

    to_val = _parse_numero(d.get('taxa_ocupacao', ''))
    perm_val = _parse_numero(d.get('taxa_permeabilidade', ''))
    terreno_val = _parse_numero(d.get('area_terreno', ''))
    alerta_zoneamento = (
        (to_val is not None and to_val > 70)
        and (perm_val is not None and perm_val < 20)
    )
    alerta_excecao = (terreno_val is not None and terreno_val < 220)

    # 2. Lista de candidatos (Label, Valor)
    candidatos = [
        ("Endereço:",       end_completo),
        ("Inscrição Mun.:", d.get("inscricao_municipal", "")),
        ("Proprietário:",   d.get("proprietario", d.get("requerente", ""))),
        ("Resp. Técnico:",  d.get("profissional_nome", d.get("responsavel_tecnico", ""))),
        ("Lote:",           d.get("lote", "")),
        ("Quadra:",         d.get("quadra", "")),
        ("Área Terreno:",   d.get("area_terreno", "")),
        ("Área Total:",     d.get("area_total_construida", "")),
        ("Taxa Ocupação:",  d.get("taxa_ocupacao", "")),
        ("Coef. Aprov.:",   d.get("coef_aproveitamento", "")),
        ("Permeabilidade:", d.get("taxa_permeabilidade", "")),
        ("Pavimentos:",     d.get("pavimentos", "")),
        ("Vagas Garagem:",  d.get("vagas_garagem", "")),
        ("Zona de Uso:",    d.get("zona_uso", "")),
        ("ART/RRT:",        d.get("art_rrt_numero", d.get("art_rrt", ""))),
        ("Multa Específica:", d.get("tipo_multa_especifica", "")),
    ]

    # 3. Filtragem (Remove o que estiver vazio ou com marcador de preenchimento)
    def eh_valido(v):
        if not v: return False
        v_str = str(v).strip().lower()
        return v_str not in ["", "-", "—", "[preencher]", "n/a", "não informado"]

    validos = [(l, v) for l, v in candidatos if eh_valido(v)]
    if not validos:
        return # Se não há dados técnicos, não gera a seção

    # 4. Agrupamento em pares (L1, V1, L2, V2)
    linhas_finais = []
    for i in range(0, len(validos), 2):
        l1, v1 = validos[i]
        l2, v2 = validos[i+1] if (i+1) < len(validos) else ("", "")
        linhas_finais.append((l1, v1, l2, v2))

    # 5. Geração Visual
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=280, after=60)
    r_tit = add_run(p_tit, "DADOS TÉCNICOS DO PROJETO (Ref. Decreto nº 4.149/2019)",
                    bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    W_L1, W_V1 = W_CARIMBO_L1, W_CARIMBO_V1
    W_L2, W_V2 = W_CARIMBO_L2, W_CARIMBO_V2

    for lab1, val1, lab2, val2 in linhas_finais:
        row = tbl.add_row()

        # Célula Label 1
        c0 = row.cells[0]
        apply_label_cell(c0, W_L1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p0, line=220, before=25, after=25)
        r0 = p0.add_run(lab1)
        set_font(r0, size=SZ_TABELA, bold=True)
        r0.font.color.rgb = COR_LABEL_FONT

        # Célula Valor 1
        c1 = row.cells[1]
        apply_value_cell(c1, W_V1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p1, line=220, before=25, after=25)
        
        cor_v1, alerta_msg_v1 = None, None
        if lab1 == "Área Terreno:" and alerta_excecao:
            cor_v1, alerta_msg_v1 = COR_ALERTA_GREEN, "  (exceção da lei)"
        elif lab1 in ["Taxa Ocupação:", "Permeabilidade:"] and alerta_zoneamento:
            cor_v1, alerta_msg_v1 = COR_ALERTA_RED, "  (conferir zoneamento)"

        rich_segments(p1, val1, size=SZ_TABELA, color=cor_v1)
        if alerta_msg_v1:
            r_alert = p1.add_run(alerta_msg_v1)
            set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = cor_v1

        # Célula Label 2
        c2 = row.cells[2]
        apply_label_cell(c2, W_L2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p2, line=220, before=25, after=25)
        r2 = p2.add_run(lab2)
        set_font(r2, size=SZ_TABELA, bold=True)
        r2.font.color.rgb = COR_LABEL_FONT

        # Célula Valor 2
        c3 = row.cells[3]
        apply_value_cell(c3, W_V2)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p3, line=220, before=25, after=25)
        
        if lab2: # Apenas se houver par preenchido
            rich_segments(p3, val2, size=SZ_TABELA)

    add_separator(doc)



# ═══════════════════════════════════════════════════════════
#  CORPO DO PARECER (abertura + considerandos)
# ═══════════════════════════════════════════════════════════

def _apply_cell_fill(cell, fill_hex: str):
    """Aplica cor de fundo a uma célula de tabela via XML."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_width(cell, width_twips: int):
    """Define largura de célula em twips via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ═══════════════════════════════════════════════════════════
#  PARTES E RESPONSÁVEIS DO PROCESSO
# ═══════════════════════════════════════════════════════════

def build_partes_envolvidas(doc, partes: dict):
    """
    Tabela 'PARTES E RESPONSÁVEIS' (2 colunas: PAPEL | IDENTIFICAÇÃO).
    Gerada quando o JSON contém o campo 'partes_envolvidas'.
    """
    if not partes:
        return

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=240, after=60)
    r_tit = add_run(p_tit, "PARTES E RESPONSÁVEIS DO PROCESSO", bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    linhas = []

    req = partes.get("requerente") or {}
    if isinstance(req, str):
        req = {"nome": req}
    nome_req = req.get("nome", "")
    qualidade_req = req.get("qualidade", "")
    if nome_req:
        val = f"{nome_req} ({qualidade_req})" if qualidade_req else nome_req
        linhas.append(("Requerente", val))

    prop = partes.get("proprietario") or {}
    if isinstance(prop, str):
        prop = {"nome": prop}
    nome_prop = prop.get("nome", "")
    mat_prop = prop.get("matricula_imovel", "")
    if nome_prop and nome_prop != nome_req:
        val = nome_prop + (f" — Matrícula nº {mat_prop}" if mat_prop else "")
        linhas.append(("Proprietário", val))

    rt = partes.get("responsavel_tecnico") or {}
    if isinstance(rt, str):
        rt = {"nome": rt}
    nome_rt = rt.get("nome", "")
    if nome_rt:
        conselho = rt.get("conselho", "")
        tipo_rt = rt.get("tipo_rt", "")
        num_rt = rt.get("numero_rt", "")
        val = nome_rt
        if conselho:
            val += f" — {conselho}"
        if tipo_rt and num_rt:
            val += f" — {tipo_rt} nº {num_rt}"
        linhas.append(("Resp. Técnico", val))

    for fiscal in partes.get("agentes_fiscais") or []:
        if isinstance(fiscal, str):
            linhas.append(("Fiscal Vistoriador", fiscal))
        else:
            nome_f = fiscal.get("nome", "")
            mat_f = fiscal.get("matricula_funcional", "")
            val = nome_f + (f" (Mat. {mat_f})" if mat_f else "")
            if nome_f:
                linhas.append(("Fiscal Vistoriador", val))

    ass = partes.get("assinante_parecer") or {}
    if isinstance(ass, str):
        ass = {"nome": ass}
    nome_ass = ass.get("nome", "")
    if nome_ass:
        titulo_ass = ass.get("titulo", "")
        reg_ass = ass.get("registro", "")
        val = nome_ass
        if titulo_ass:
            val += f" — {titulo_ass}"
        if reg_ass:
            val += f" — {reg_ass}"
        linhas.append(("Signatário", val))

    if not linhas:
        return

    tbl = doc.add_table(rows=len(linhas), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    for i, (papel, identificacao) in enumerate(linhas):
        fill_hex = 'EEF2FB' if i % 2 == 0 else 'FFFFFF'
        row = tbl.rows[i]

        c_papel = row.cells[0]
        _set_cell_width(c_papel, W_PARTES_LABEL)
        _apply_cell_fill(c_papel, fill_hex)
        set_cell_margins(c_papel, top=40, bottom=40, left=80, right=40)
        p_papel = c_papel.paragraphs[0]
        p_papel.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_papel, line=220, before=20, after=20)
        r_p = p_papel.add_run(papel)
        set_font(r_p, size=SZ_TABELA, bold=True)
        r_p.font.color.rgb = COR_LABEL_FONT

        c_id = row.cells[1]
        _set_cell_width(c_id, W_PARTES_VALUE)
        _apply_cell_fill(c_id, fill_hex)
        set_cell_margins(c_id, top=40, bottom=40, left=80, right=40)
        p_id = c_id.paragraphs[0]
        p_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_id, line=220, before=20, after=20)
        rich_segments(p_id, identificacao, size=SZ_TABELA)

    add_separator(doc, color='D0D0D0')


# ═══════════════════════════════════════════════════════════
#  HISTÓRICO CRONOLÓGICO DO PROCESSO
# ═══════════════════════════════════════════════════════════

def build_historico_cronologico(doc, historico: list):
    """
    Tabela 'HISTÓRICO CRONOLÓGICO' (3 colunas: DATA | EVENTO | REFERÊNCIA/AGENTES).
    Gerada quando o JSON contém o campo 'historico_cronologico'.
    Os eventos devem estar em ordem cronológica (do mais antigo ao mais recente).
    """
    if not historico:
        return

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=240, after=60)
    r_tit = add_run(p_tit, "HISTÓRICO CRONOLÓGICO DO PROCESSO", bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=len(historico) + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    # Cabeçalho
    headers = [("DATA", W_HIST_DATA, WD_ALIGN_PARAGRAPH.CENTER),
               ("EVENTO / DESCRIÇÃO", W_HIST_EVENT, WD_ALIGN_PARAGRAPH.LEFT),
               ("REFERÊNCIA / AGENTES", W_HIST_REF, WD_ALIGN_PARAGRAPH.LEFT)]
    hrow = tbl.rows[0]
    for j, (hdr, w, align) in enumerate(headers):
        cell = hrow.cells[j]
        _set_cell_width(cell, w)
        _apply_cell_fill(cell, COR_INST)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=40)
        p = cell.paragraphs[0]
        p.alignment = align
        set_spacing(p, line=200, before=20, after=20)
        r = p.add_run(hdr)
        set_font(r, size=SZ_TABELA, bold=True)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas de dados
    for i, ev in enumerate(historico):
        fill_hex = 'F0F4FD' if i % 2 == 0 else 'FFFFFF'
        row = tbl.rows[i + 1]

        # DATA
        c_data = row.cells[0]
        _set_cell_width(c_data, W_HIST_DATA)
        _apply_cell_fill(c_data, fill_hex)
        set_cell_margins(c_data, top=40, bottom=40, left=80, right=40)
        p_d = c_data.paragraphs[0]
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p_d, line=220, before=20, after=20)
        r_d = p_d.add_run(str(ev.get("data", "")))
        set_font(r_d, size=SZ_TABELA, bold=True)
        r_d.font.color.rgb = COR_LABEL_FONT

        # EVENTO
        c_ev = row.cells[1]
        _set_cell_width(c_ev, W_HIST_EVENT)
        _apply_cell_fill(c_ev, fill_hex)
        set_cell_margins(c_ev, top=40, bottom=40, left=80, right=40)
        p_ev = c_ev.paragraphs[0]
        p_ev.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_ev, line=220, before=20, after=20)
        rich_segments(p_ev, str(ev.get("evento", "")), size=SZ_TABELA)

        # REFERÊNCIA / AGENTES
        c_ref = row.cells[2]
        _set_cell_width(c_ref, W_HIST_REF)
        _apply_cell_fill(c_ref, fill_hex)
        set_cell_margins(c_ref, top=40, bottom=40, left=80, right=40)
        p_ref = c_ref.paragraphs[0]
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_ref, line=220, before=20, after=20)

        agentes = ev.get("agentes", [])
        ref_text = ev.get("referencia", "")
        if agentes:
            if isinstance(agentes, list):
                ref_text = "; ".join(str(a) for a in agentes)
            else:
                ref_text = str(agentes)

        rich_segments(p_ref, ref_text, size=SZ_TABELA)

    add_separator(doc, color='D0D0D0')


# ═══════════════════════════════════════════════════════════
#  COMUNICADO DE PENDÊNCIA — Design Visual Dedicado
# ═══════════════════════════════════════════════════════════

