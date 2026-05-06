import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

import re
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    FONT_TITULO, SZ_CORPO, SZ_TABELA,
    COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE, INDENT_PADRAO,
    W_IDENT_LABEL, W_IDENT_VALUE,
    W_CARIMBO_L1, W_CARIMBO_V1, W_CARIMBO_L2, W_CARIMBO_V2,
    W_PARTES_LABEL, W_PARTES_VALUE,
    W_HIST_DATA, W_HIST_EVENT, W_HIST_REF,
    W_CARD_TOTAL, AREA_UTIL_TWIPS,
    COR_PENDENCIA_FILL, COR_PENDENCIA_BORDA,
    COR_PENDENCIA_TEXTO, COR_PENDENCIA_ICON,
    COR_SUCESSO_FILL, COR_SUCESSO_BORDA,
    COR_SUCESSO_TEXTO, COR_SUCESSO_ICON,
    COR_ALERTA_RED, COR_ALERTA_GREEN, COR_DOC_BOX_FILL,
    COR_INST,
)
from generators.formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator, add_section_heading,
)

from generators.componentes.tabelas import _data_hoje_extenso

def build_assinatura(doc, d):
    """Bloco de assinatura: data + linha tracejada + nome/título/registro."""
    assinante = d.get("assinante", ASSINANTE)
    cidade    = d.get("cidade", CIDADE)

    # Local e data
    p_local = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       line=LINE_SPC, before=280, after=0)
    add_run(p_local, f"{cidade}, {_data_hoje_extenso()}.", bold=False, size=SZ_CORPO)
    p_local.paragraph_format.keep_with_next = True

    # Espaço para assinar
    for _ in range(3):
        pb = doc.add_paragraph()
        set_spacing(pb, line=240, before=0, after=0)
        pb.paragraph_format.keep_with_next = True

    # Linha de assinatura via borda inferior do parágrafo (mais elegante que underscores)
    p_line = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=60, before=0, after=60)
    # Largura limitada via indentação
    p_line.paragraph_format.left_indent  = Pt(100)
    p_line.paragraph_format.right_indent = Pt(100)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), COR_INST)
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)
    p_line.paragraph_format.keep_with_next = True

    # Nome em negrito (destaque)
    nome = d.get("assinante_parecer") or assinante.get("nome", ASSINANTE["nome"])
    p_nome = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=LINE_SPC, before=40, after=0)
    r_nome = add_run(p_nome, nome, bold=True, size=SZ_CORPO)
    r_nome.font.color.rgb = COR_LABEL_FONT
    p_nome.paragraph_format.keep_with_next = True

    # Título e Registro em tamanho menor
    titulo   = assinante.get("titulo",   ASSINANTE["titulo"])
    registro = assinante.get("registro", ASSINANTE["registro"])
    for texto in [titulo, registro]:
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=240, before=0, after=0)
        r = add_run(p, texto, bold=False, size=9)
        r.font.color.rgb = COR_CINZA_TEXTO
        p.paragraph_format.keep_with_next = True

    # Secretaria
    p_sec = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=240, before=0, after=120)
    r_sec = add_run(p_sec, "Secretaria Municipal de Obras e Serviços Urbanos", bold=False, size=9)
    r_sec.font.color.rgb = COR_CINZA_TEXTO


# ═══════════════════════════════════════════════════════════
#  UTILITÁRIOS INTERNOS DE CÉLULA
# ═══════════════════════════════════════════════════════════
