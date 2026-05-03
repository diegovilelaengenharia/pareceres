"""
Construção do cabeçalho institucional e rodapé com numeração de páginas.
"""

import os
from docx.shared import Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.config import (
    LOGO_PREFEITURA, FONT_HEADER, FONT_DETALHE,
    COR_LABEL_FONT, COR_CINZA_LEVE, COR_INST, SZ_RODAPE,
)
from generators.formatacao import set_spacing, set_font, no_borders, add_field


# ═══════════════════════════════════════════════════════════
#  CABEÇALHO INSTITUCIONAL
# ═══════════════════════════════════════════════════════════

def build_header(doc):
    """Constrói cabeçalho premium com logo + dados da prefeitura."""
    section = doc.sections[0]
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p._element.getparent().remove(p._element)

    W_LOGO = 2400
    W_TEXT = 7800
    total = W_LOGO + W_TEXT

    tbl = hdr.add_table(rows=1, cols=2, width=Twips(total))
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remover bordas da tabela do cabeçalho
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBrd.append(el)
    tblPr.append(tblBrd)

    row = tbl.rows[0]
    row.height = Twips(1200)

    # ── Logo Prefeitura (à esquerda) ──
    c0 = row.cells[0]
    c0.width = Twips(W_LOGO)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_borders(c0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p0, line=240, before=0, after=0)
    if os.path.exists(LOGO_PREFEITURA):
        p0.add_run().add_picture(LOGO_PREFEITURA, width=Cm(3.3))

    # ── Bloco de texto (à direita do logo) ──
    c1 = row.cells[1]
    c1.width = Twips(W_TEXT)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_borders(c1)
    for extra in c1.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # L1: PREFEITURA MUNICIPAL DE OLIVEIRA
    p1a = c1.paragraphs[0]
    p1a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p1a, line=280, before=0, after=10)  # sem espaço acima
    r_pref = p1a.add_run("PREFEITURA MUNICIPAL DE OLIVEIRA")
    set_font(r_pref, name=FONT_HEADER, size=14, bold=True)
    r_pref.font.color.rgb = COR_LABEL_FONT

    # L2: Secretaria
    p1c = c1.add_paragraph()
    p1c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p1c, line=240, before=0, after=8)
    r_sec = p1c.add_run("Secretaria Municipal de Obras e Serviços Urbanos")
    set_font(r_sec, name=FONT_HEADER, size=8, bold=True)
    r_sec.font.color.rgb = COR_LABEL_FONT

    # L3: Endereço e telefone
    pb = c1.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(pb, line=220, before=0, after=10)
    r_info = pb.add_run(
        "Praça XV de Novembro, 127  •  Centro  •  Oliveira/MG  •  (37) 3331-9800"
    )
    set_font(r_info, name=FONT_DETALHE, size=7)
    r_info.font.color.rgb = COR_CINZA_LEVE

    # ── Linha separadora dupla (fina cinza + grossa azul) ──
    p_thin = hdr.add_paragraph()
    set_spacing(p_thin, line=20, before=8, after=0)  # linha logo após o endereço
    pPr_thin = p_thin._p.get_or_add_pPr()
    pb_thin = OxmlElement('w:pBdr')
    bot_thin = OxmlElement('w:bottom')
    bot_thin.set(qn('w:val'), 'single')
    bot_thin.set(qn('w:sz'), '4')
    bot_thin.set(qn('w:color'), 'AAAAAA')
    bot_thin.set(qn('w:space'), '1')
    pb_thin.append(bot_thin)
    pPr_thin.append(pb_thin)

    p_thick = hdr.add_paragraph()
    set_spacing(p_thick, line=20, before=4, after=0)
    pPr_thick = p_thick._p.get_or_add_pPr()
    pb_thick = OxmlElement('w:pBdr')
    bot_thick = OxmlElement('w:bottom')
    bot_thick.set(qn('w:val'), 'single')
    bot_thick.set(qn('w:sz'), '18')
    bot_thick.set(qn('w:color'), COR_INST)
    bot_thick.set(qn('w:space'), '1')
    pb_thick.append(bot_thick)
    pPr_thick.append(pb_thick)


# ═══════════════════════════════════════════════════════════
#  RODAPÉ COM NUMERAÇÃO
# ═══════════════════════════════════════════════════════════

def add_page_number_footer(doc):
    """Rodapé institucional: linha azul + texto secretaria à esquerda + página à direita."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Linha fina divisora no topo do rodapé ──
    p_linha = footer.add_paragraph()
    set_spacing(p_linha, line=100, before=80, after=20)  # maior respiro acima da linha
    pPr = p_linha._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'), 'single')
    top_el.set(qn('w:sz'), '12')
    top_el.set(qn('w:color'), COR_INST)
    top_el.set(qn('w:space'), '1')
    pBdr.append(top_el)
    pPr.append(pBdr)

    # ── Tabela 2 colunas: Texto à esquerda | Página à direita ──
    tbl = footer.add_table(rows=1, cols=2, width=Twips(9360))
    tbl.autofit = False
    # Sem bordas
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBrd.append(el)
    tblPr.append(tblBrd)

    row = tbl.rows[0]
    # Coluna esquerda: nome da secretaria
    c_left = row.cells[0]
    c_left.width = Twips(6000)
    p_left = c_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_left, line=260, before=8, after=0)
    r_left = p_left.add_run("SMOSU – Secretaria Municipal de Obras e Serviços Urbanos")
    set_font(r_left, size=9)
    r_left.font.color.rgb = COR_CINZA_LEVE

    # Coluna direita: Página X de Y
    c_right = row.cells[1]
    c_right.width = Twips(3360)
    fp = c_right.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, line=260, before=8, after=0)

    r1 = fp.add_run('Página ')
    set_font(r1, size=9)
    r1.font.color.rgb = COR_CINZA_LEVE
    r2 = fp.add_run()
    set_font(r2, size=9)
    r2.font.color.rgb = COR_CINZA_LEVE
    add_field(r2, ' PAGE ')
    r3 = fp.add_run(' de ')
    set_font(r3, size=9)
    r3.font.color.rgb = COR_CINZA_LEVE
    r4 = fp.add_run()
    set_font(r4, size=9)
    r4.font.color.rgb = COR_CINZA_LEVE
    add_field(r4, ' NUMPAGES ')

