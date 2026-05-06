"""
Sistema de despacho de geradores de documentos — SMOSU Oliveira/MG.

Identifica o tipo_relatorio no JSON de entrada, carrega o template
correspondente e chama o gerador correto (parecer_tecnico, parecer_simples,
ofício ou comunicado).
"""

import os
import json
import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys as _sys
import os as _os
_MOTOR_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
if _MOTOR_DIR not in _sys.path:
    _sys.path.insert(0, _MOTOR_DIR)

from core.config import (
    TIPOS_DOCUMENTO, TEMPLATES_DIR,
    FONT_CORPO, SZ_CORPO, PASTA_SAIDA,
)
from generators.cabecalho import build_header, add_page_number_footer
from generators.componentes import (
    build_titulo, build_identificacao, build_destinatario,
    build_dados_carimbo, build_corpo,
    build_conclusao_e_docs, build_conclusao_simples, build_assinatura,
    build_comunicado_pendencia,
    build_partes_envolvidas, build_historico_cronologico,
    build_memoria_calculo,
)
from generators.enricher import enriquecer_dados

# ═══════════════════════════════════════════════════════════
#  TEMPLATES
# ═══════════════════════════════════════════════════════════

def carregar_template(tipo_relatorio):
    """Carrega o template JSON para o tipo de documento, se existir."""
    caminho = os.path.join(TEMPLATES_DIR, f"{tipo_relatorio}.json")
    if os.path.exists(caminho):
        with open(caminho, encoding="utf-8") as f:
            return json.load(f)
    return {}  # template vazio = usar defaults


# ═══════════════════════════════════════════════════════════
#  DOCUMENTO BASE
# ═══════════════════════════════════════════════════════════

def _criar_documento_base():
    """Cria documento com margens A4 e estilo padrão configurados."""
    doc = Document()
    section = doc.sections[0]
    # Margens 1.5cm + cabeçalho rente ao topo + rodapé com altura
    M = Cm(1.5)
    section.page_height      = Cm(29.7)
    section.page_width       = Cm(21.0)
    section.top_margin       = Cm(3.2)   # corpo começa após o cabeçalho (~2.8cm de altura)
    section.bottom_margin    = Cm(2.2)   # espaço generoso para o rodapé
    section.left_margin      = M
    section.right_margin     = M
    section.header_distance  = Cm(0.3)   # logo colado na borda superior — sem borda branca
    section.footer_distance  = Cm(0.5)   # rodapé colado na borda inferior

    style = doc.styles['Normal']
    style.font.name = FONT_CORPO
    style.font.size = Pt(SZ_CORPO)
    return doc


def _gerar_nome_saida(dados):
    """Gera nome de arquivo e pasta baseados nos dados do processo."""
    now = datetime.datetime.now()

    numero_raw = str(dados.get("numero_processo", "S-N"))
    proc_limpo = numero_raw.replace("/", "-").replace("\\", "-")

    # Separa número e ano — barra tem precedência; hífen só se ambas as partes forem dígitos
    proc_num = numero_raw
    proc_ano = now.strftime("%Y")
    if "/" in numero_raw:
        partes_proc = numero_raw.split("/")
        proc_num = partes_proc[0]
        proc_ano = partes_proc[-1]
    elif "-" in numero_raw:
        partes_proc = numero_raw.split("-")
        if len(partes_proc) == 2 and partes_proc[0].isdigit() and partes_proc[1].isdigit():
            proc_num = partes_proc[0]
            proc_ano = partes_proc[-1]

    nome_alvo = dados.get("requerente") or dados.get("proprietario_nome") or dados.get("interessado") or "Desconhecido"
    nome_completo_title = str(nome_alvo).strip().title()

    tipo = str(dados.get("tipo_relatorio", "Parecer")).replace("_", " ").title()

    nome_arquivo = f"{tipo} - {proc_num}-{proc_ano} - {nome_completo_title}.docx"
    nome_pasta = f"Processo {proc_num}-{proc_ano} - {nome_completo_title}"

    for ch in ['<', '>', ':', '"', '|', '?', '*', '\\']:
        nome_arquivo = nome_arquivo.replace(ch, '')
        nome_pasta = nome_pasta.replace(ch, '')

    final_out_dir = os.path.join(PASTA_SAIDA, nome_pasta)
    os.makedirs(final_out_dir, exist_ok=True)
    return os.path.join(final_out_dir, nome_arquivo)


# ═══════════════════════════════════════════════════════════
#  GERADORES POR CATEGORIA
# ═══════════════════════════════════════════════════════════

def gerar_parecer_tecnico(doc, dados, template):
    """
    Parecer técnico COMPLETO:
    Header → Título → Identificação → Dados Carimbo → [Partes] → [Histórico] → Corpo → Conclusão+Docs
    """
    titulo = template.get("titulo_documento", "PARECER SETOR TÉCNICO - SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_dados_carimbo(doc, dados)
    
    partes = dados.get("partes_envolvidas")
    if not partes and (dados.get("agentes_fiscais") or dados.get("responsavel_tecnico")):
        partes = {
            "requerente": {"nome": dados.get("requerente", "")},
            "responsavel_tecnico": {"nome": dados.get("responsavel_tecnico", dados.get("profissional_nome", ""))},
            "agentes_fiscais": dados.get("agentes_fiscais") if isinstance(dados.get("agentes_fiscais"), list) else [dados.get("agentes_fiscais")],
            "assinante_parecer": {"nome": dados.get("assinante_parecer", "")}
        }
    if partes:
        build_partes_envolvidas(doc, partes)
        
    if dados.get("historico_cronologico"):
        build_historico_cronologico(doc, dados["historico_cronologico"])
    build_corpo(doc, dados)
    build_memoria_calculo(doc, dados)
    build_conclusao_e_docs(doc, dados)


def gerar_parecer_administrativo(doc, dados, template):
    """
    Parecer administrativo (Layout limpo para certidões):
    Header → Título → Identificação → [Partes] → [Histórico] → Corpo → Conclusão+Docs
    (Omite Dados do Carimbo e Memória de Cálculo)
    """
    titulo = template.get("titulo_documento", "PARECER ADMINISTRATIVO - SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    
    partes = dados.get("partes_envolvidas")
    if not partes and (dados.get("agentes_fiscais") or dados.get("responsavel_tecnico")):
        partes = {
            "requerente": {"nome": dados.get("requerente", "")},
            "responsavel_tecnico": {"nome": dados.get("responsavel_tecnico", dados.get("profissional_nome", ""))},
            "agentes_fiscais": dados.get("agentes_fiscais") if isinstance(dados.get("agentes_fiscais"), list) else [dados.get("agentes_fiscais")],
            "assinante_parecer": {"nome": dados.get("assinante_parecer", "")}
        }
    if partes:
        build_partes_envolvidas(doc, partes)
        
    if dados.get("historico_cronologico"):
        build_historico_cronologico(doc, dados["historico_cronologico"])
    build_corpo(doc, dados)
    
    # build_memoria_calculo omitida propositalmente per D-02
    build_conclusao_e_docs(doc, dados)


def gerar_parecer_simples(doc, dados, template):
    """
    Parecer SIMPLES (sem dados do carimbo):
    Header → Título → Identificação → Corpo → Conclusão
    """
    titulo = template.get("titulo_documento", "PARECER SETOR TÉCNICO - SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_corpo(doc, dados)
    build_memoria_calculo(doc, dados)

    if dados.get("documentos_emitir"):
        build_conclusao_e_docs(doc, dados)
    else:
        build_conclusao_simples(doc, dados)


def gerar_oficio(doc, dados, template):
    """
    OFÍCIO:
    Header → Título → Destinatário → Corpo → Assinatura
    """
    titulo = template.get("titulo_documento", "OFÍCIO")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_destinatario(doc, dados)
    build_corpo(doc, dados)
    build_assinatura(doc, dados)


def gerar_comunicado(doc, dados, template):
    """
    COMUNICADO genérico:
    Header → Título → Identificação → Corpo → Assinatura
    """
    titulo = template.get("titulo_documento", "COMUNICADO")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_corpo(doc, dados)
    build_assinatura(doc, dados)


def gerar_comunicado_pendencia(doc, dados, template):
    """
    COMUNICADO DE PENDÊNCIA — layout visual dedicado:
    Header → Título → Identificação → [Box Alerta + Box Orientação] → Assinatura
    """
    titulo = template.get("titulo_documento", "COMUNICADO DE PENDÊNCIA DOCUMENTAL")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)
    build_comunicado_pendencia(doc, dados)


def gerar_triagem(doc, dados, template):
    """
    TRIAGEM SMOSU (Pedro Barros):
    Header → Título → Parágrafo fixo → Checkbox → [Pendências] → Assinatura
    """
    from generators.formatacao import add_para, rich_segments

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, "TRIAGEM - SMOSU")

    # Parágrafo fixo sobre o Decreto
    p_dec = add_para(doc, after=120, indent_cm=0)
    p_dec.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rich_segments(p_dec,
        "Conforme Decreto n°4.149 de 19 de dezembro de 2019, estabelece procedimentos "
        "e critérios para a concessão da aprovação de projetos e dá outras providências, "
        "e vistoria do setor de Fiscalização.", size=11)

    # Linha de documentação
    p_doc = add_para(doc, before=160, after=80, indent_cm=0)
    rich_segments(p_doc, "Documentação para análise técnica:", size=11)

    # Checkbox deferido / indeferido
    deferido = dados.get("deferido", False)
    if deferido:
        checkbox = "( X ) Deferida           (    ) Indeferida"
    else:
        checkbox = "(    ) Deferida           (   X   ) Indeferida"
    p_cb = add_para(doc, after=120, indent_cm=0)
    rich_segments(p_cb, checkbox, size=11)

    # Pendências (somente se indeferido)
    if not deferido:
        pendencias = dados.get("pendencias", [])
        if pendencias:
            p_comm = add_para(doc, before=120, after=80, indent_cm=0)
            rich_segments(p_comm, "Comunicado visto o Indeferimento:", size=11)
            for item in pendencias:
                p_it = add_para(doc, after=80, indent_cm=0.63)
                rich_segments(p_it, f"- {item}", size=11)

    build_assinatura(doc, dados)


def gerar_devolutiva_retificacao(doc, dados, template):
    """
    Gera a Devolutiva de Retificação com layout visual completo:
    1. DOCUMENTAÇÃO VERIFICADA E CRONOLOGIA
    2. FUNDAMENTAÇÃO TÉCNICA
    3. CONCLUSÃO

    Lê conteúdo de `dados` com fallback para `template`.
    """
    from generators.formatacao import add_section_heading, add_para, rich_segments, set_font
    from generators.componentes import _ensure_list, add_doc_item

    titulo = dados.get("titulo_documento") or template.get("titulo_documento", "PARECER SETOR TÉCNICO – SMOSU")

    build_header(doc)
    add_page_number_footer(doc)
    build_titulo(doc, titulo)
    build_identificacao(doc, dados)

    # ── 1. DOCUMENTAÇÃO VERIFICADA E CRONOLOGIA ──────────────────────────────
    add_section_heading(doc, "1. Documentação Verificada e Cronologia")

    abertura = dados.get("paragrafo_abertura") or template.get("paragrafo_abertura", "")
    if abertura:
        for linha in [l.strip() for l in abertura.split('\n') if l.strip()]:
            p_ab = add_para(doc, after=120, indent_cm=1.25)
            p_ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rich_segments(p_ab, linha, size=11)

    # ── 2. FUNDAMENTAÇÃO TÉCNICA ─────────────────────────────────────────────
    add_section_heading(doc, "2. Fundamentação Técnica")

    considerandos = _ensure_list(dados.get("considerandos") or template.get("considerandos", []))
    for cons in considerandos:
        p_cons = add_para(doc, after=120, indent_cm=1.25)
        p_cons.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rich_segments(p_cons, cons, size=11)

    # ── 3. CONCLUSÃO ─────────────────────────────────────────────────────────
    add_section_heading(doc, "3. Conclusão")

    concl = dados.get("conclusao") or template.get("conclusao", "")
    if concl:
        p_c = add_para(doc, after=120, indent_cm=1.25)
        p_c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rich_segments(p_c, concl, size=11)

    # ── Emissão de Documentos + Assinatura ───────────────────────────────────
    docs = dados.get("documentos_emitir") or template.get("documentos_emitir", [])
    if docs:
        from core.config import FONT_TITULO, SZ_CORPO, COR_LABEL_FONT
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        from generators.formatacao import set_spacing
        set_spacing(p_doc_tit, line=276, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        for item in docs:
            if isinstance(item, dict):
                descricao = item.get("descricao") or item.get("tipo", "")
                obs = item.get("obs") or item.get("observacao") or None
            else:
                descricao = str(item)
                obs = None
            add_doc_item(doc, descricao, obs)

    build_assinatura(doc, dados)


# ═══════════════════════════════════════════════════════════
#  MAPEAMENTO E DESPACHO
# ═══════════════════════════════════════════════════════════

GERADORES = {
    "parecer_tecnico":        gerar_parecer_tecnico,
    "parecer_simples":        gerar_parecer_simples,
    "parecer_administrativo": gerar_parecer_administrativo,
    "devolutiva_retificacao": gerar_devolutiva_retificacao,
    "oficio":                 gerar_oficio,
    "comunicado":             gerar_comunicado,
    "comunicado_pendencia":   gerar_comunicado_pendencia,
    "triagem":                gerar_triagem,
}


def gerar(dados, caminho_saida=None):
    """
    Função principal: gera documento baseado nos dados e tipo_relatorio.

    Parâmetros:
        dados: dict com os dados do processo (JSON do Gemini)
        caminho_saida: caminho opcional para o arquivo .docx

    Retorna:
        Caminho do arquivo gerado.
    """
    tipo = dados.get("tipo_relatorio")
    if not tipo:
        raise ValueError("O campo 'tipo_relatorio' é obrigatório e não foi encontrado no JSON de entrada.")
    
    # ---------------------------------------------------------
    # RESILIÊNCIA: Se o LLM agrupar os campos em sub-dicionários
    # como 'campos_obrigatorios' ou 'campos_opcionais', nós os achata.
    # ---------------------------------------------------------
    if "campos_obrigatorios" in dados and isinstance(dados["campos_obrigatorios"], dict):
        dados.update(dados.pop("campos_obrigatorios"))
    if "campos_opcionais" in dados and isinstance(dados["campos_opcionais"], dict):
        dados.update(dados.pop("campos_opcionais"))
    
    # Tratamento caso ele faça lista de strings em documentos_emitir no invés de `{tipo:..}`
    if "documentos_emitir" in dados and isinstance(dados["documentos_emitir"], list):
        lista_corrigida = []
        for d in dados["documentos_emitir"]:
            if isinstance(d, str):
                lista_corrigida.append({"tipo": d})
            else:
                lista_corrigida.append(d)
        dados["documentos_emitir"] = lista_corrigida

    # Injetar obs SERO automaticamente em habite-se com sero_metadata
    _HABITE_SE_TIPOS = {"habitese_comum", "habitese_multa", "habitese_inclusao_area"}
    if tipo in _HABITE_SE_TIPOS and dados.get("sero_metadata"):
        try:
            import gerador_sero
            obs_sero = gerador_sero.gerar_obs_sero(
                dados["sero_metadata"],
                dados.get("area_total_construida", "")
            )
            docs = dados.get("documentos_emitir", [])
            for doc_item in docs:
                if isinstance(doc_item, dict):
                    nome_tipo = str(doc_item.get("tipo", "")).lower()
                    if "habite" in nome_tipo or not doc_item.get("obs"):
                        doc_item["obs"] = obs_sero
                        break
        except Exception as _e:
            print(f"[!] Injeção de obs SERO falhou: {_e}")

    categoria = TIPOS_DOCUMENTO.get(tipo)

    if not categoria:
        print(f"[!] Tipo de documento inválido ou não mapeado: '{tipo}'")
        print(f"    Verifique se o 'tipo_relatorio' no JSON corresponde a um ID técnico no config.py.")
        print(f"    IDs técnicos disponíveis: {', '.join(sorted(TIPOS_DOCUMENTO.keys()))}")
        raise ValueError(f"Tipo de documento inválido: {tipo}")

    print(f"[>] Tipo: {tipo} -> Categoria: {categoria}")

    # Carregar template (se existir)
    template = carregar_template(tipo)

    # === VALIDADOR DE INTEGRIDADE PARA AUTO-CORREÇÃO DO GEM ===
    obrigatorios = template.get("campos_obrigatorios", [])
    
    # Resolver aliases antes de checar campos faltantes
    from generators._aliases import normalizar_dados
    dados = normalizar_dados(dados)

    # Enriquecer dados ausentes com modelos narrativos do template (Fase 16)
    dados = enriquecer_dados(dados, template)

    # Campos que se faltarem, o documento fica "quebrado" ou ilegal (identificação)
    CRITICOS = {
        "numero_processo", "requerente", "logradouro", "bairro", 
        "profissional_nome", "art_rrt", "zona_uso"
    }

    campos_faltantes = [c for c in obrigatorios if not dados.get(c)]
    
    if campos_faltantes:
        faltantes_criticos = [c for c in campos_faltantes if c in CRITICOS]
        
        if faltantes_criticos:
            print("\n" + "="*70)
            print("  [AVISO ESTRUTURAL] DADOS CRÍTICOS AUSENTES NO JSON")
            print("="*70)
            print(f"  Tipo do Documento: {tipo}")
            print(f"  Campos Críticos Faltantes: {', '.join(faltantes_criticos)}")
            print("-" * 70)
            print("  [Ação] O sistema irá preencher esses campos com marcadores [PREENCHER].")
            print("="*70 + "\n")
            
            for c in faltantes_criticos:
                dados[c] = f"[PREENCHER: {c}]"
        
        # Para campos de texto (considerandos, etc), não injetamos marcadores se houver 
        # algum conteúdo substancial no documento, assumindo que o GEM consolidou o texto.
        tem_texto_base = bool(dados.get("paragrafo_abertura") or dados.get("considerandos"))
        
        for c in campos_faltantes:
            if c in CRITICOS: continue # já tratado acima
            
            if c == "documentos_emitir":
                if not dados.get("documentos_emitir"):
                    dados[c] = [{"tipo": "[PREENCHER: documento a emitir]"}]
            elif c in ["considerandos", "fundamentacao_legal", "paragrafos_adicionais"]:
                # Se já tem abertura ou considerandos, não polui o documento com "[PREENCHER]" nesses campos narrativos
                if not tem_texto_base:
                    dados[c] = [f"[PREENCHER: {c}]"]
                else:
                    # Deixa como lista vazia se for lista, ou string vazia se for string, para não sujar o parecer do GEM
                    if isinstance(template.get(c), list) or c in ["considerandos", "fundamentacao_legal"]:
                         dados[c] = []
                    else:
                         dados[c] = ""
            else:
                # Outros campos técnicos menores
                if c not in ["inscricao_municipal", "matricula_sri"]:
                    dados[c] = f"[PREENCHER: {c}]"
                else:
                    dados[c] = "" # Omitir se for opcional/técnico e faltar


    # Criar documento base
    doc = _criar_documento_base()

    # Chamar gerador
    gerador_fn = GERADORES[categoria]
    gerador_fn(doc, dados, template)

    # Salvar
    if caminho_saida is None:
        caminho_saida = _gerar_nome_saida(dados)

    doc.save(caminho_saida)
    print(f"[+] Documento DOCX gerado: {caminho_saida}")

    # Geração de PDF removida por solicitação do usuário
    # _gerar_pdf(caminho_saida)

    return caminho_saida


# ═══════════════════════════════════════════════════════════
#  GERAÇÃO DE PDF
# ═══════════════════════════════════════════════════════════

def _gerar_pdf(caminho_docx: str) -> str | None:
    """
    Tenta converter o DOCX em PDF.
    Pipeline: Word COM (Windows) -> LibreOffice Headless (Multiplataforma) -> docx2pdf.
    """
    import os, shutil, tempfile, subprocess
    base, _ = os.path.splitext(caminho_docx)
    caminho_pdf = base + ".pdf"
    abs_docx    = os.path.abspath(caminho_docx)
    abs_pdf     = os.path.abspath(caminho_pdf)

    # 1. Tentar via Word COM (Apenas Windows)
    if os.name == 'nt':
        print(f"  [.] Convertendo DOCX para PDF (via Word COM)...")
        # Word COM não suporta paths > 255 chars; usa temp se necessário
        usar_temp = False
        docx_conv, pdf_conv = abs_docx, abs_pdf
        if len(abs_docx) > 220 or len(abs_pdf) > 220:
            tmp_dir  = tempfile.gettempdir()
            tmp_docx = os.path.join(tmp_dir, "gem_conv.docx")
            tmp_pdf  = os.path.join(tmp_dir, "gem_conv.pdf")
            shutil.copy2(abs_docx, tmp_docx)
            docx_conv, pdf_conv = tmp_docx, tmp_pdf
            usar_temp = True

        try:
            import comtypes.client
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                doc_w = word.Documents.Open(docx_conv, ConfirmConversions=False, ReadOnly=True, AddToRecentFiles=False)
                doc_w.SaveAs(pdf_conv, FileFormat=17)  # 17 = wdFormatPDF
                doc_w.Close(0)
                if usar_temp:
                    shutil.move(pdf_conv, abs_pdf)
                print(f"  [+] Documento PDF gerado: {caminho_pdf}")
                return caminho_pdf
            finally:
                word.Quit(0)
                if usar_temp and os.path.exists(tmp_docx): os.remove(tmp_docx)
        except (ImportError, Exception) as e:
            print(f"  [!] Word COM falhou ou não disponível: {e}")

    # 2. Fallback: LibreOffice Headless (soffice)
    print(f"  [.] Tentando conversão via LibreOffice Headless...")
    try:
        # Comando varia conforme o SO, soffice é comum em ambos
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(abs_pdf), abs_docx]
        # Tenta rodar de forma silenciosa
        result = subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=30)
        if result.returncode == 0 and os.path.exists(abs_pdf):
            print(f"  [+] Documento PDF gerado via LibreOffice: {caminho_pdf}")
            return caminho_pdf
    except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
        print(f"  [!] LibreOffice (soffice) não encontrado ou falhou.")

    # 3. Fallback Final: docx2pdf (wrapper de Word/Mac)
    try:
        print(f"  [.] Tentando conversão via docx2pdf...")
        from docx2pdf import convert
        convert(abs_docx, abs_pdf)
        if os.path.exists(abs_pdf):
            print(f"  [+] Documento PDF gerado via docx2pdf: {caminho_pdf}")
            return caminho_pdf
    except (ImportError, Exception):
        print(f"  [!] docx2pdf falhou.")

    print("\n[!] AVISO: PDF NÃO FOI GERADO AUTOMATICAMENTE.")
    print("    O documento DOCX está íntegro na pasta de saída.")
    print("    Para suporte a PDF em servidores Linux, instale o LibreOffice (sudo apt install libreoffice).")
    return None

