"""
compilador_livre.py  —  Modo Livre de Geração de Pareceres
Prefeitura de Oliveira / SMOSU

Aceita um JSON com "texto_livre" (texto narrativo da IA sem estrutura rígida)
e gera o DOCX/PDF no mesmo layout visual do motor padrão (compilador.py).

Campos do JSON:
  - tipo_relatorio:      igual ao motor (ex: "alvara_aprovacao")
  - numero_processo:     ex: "2765/2026"
  - data_processo:       ex: "03 de março de 2026"
  - assunto:             resumo do pedido
  - requerente:          nome em maiúsculas
  - logradouro:          endereço
  - bairro:              bairro
  - inscricao_municipal: número da inscrição
  - lote / quadra:       (opcionais)
  - area_terreno:        ex: "200,01m²"
  - area_total_construida
  - taxa_ocupacao
  - coef_aproveitamento
  - taxa_permeabilidade
  - profissional_nome:   nome do RT
  - texto_livre:         TEXTO COMPLETO DO PARECER (parágrafos narrativos, suporta **negrito** e __itálico__)
  - documentos_emitir:   lista [{tipo, obs}] (igual ao motor)
"""

import sys
import json
import os
from pathlib import Path

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

from docx import Document
from docx.shared import Pt, Cm

from cabecalho     import build_header
from componentes   import (
    build_titulo, build_identificacao, build_dados_carimbo,
    build_conclusao_e_docs, build_assinatura,
    add_doc_item, _ensure_list, _build_conclusao_bloco
)
from formatacao    import add_para, add_run, rich_segments, set_spacing
from config        import (
    FONT_TITULO, SZ_CORPO, COR_LABEL_FONT,
    PAR_AFTER, LINE_SPC,
    PASTA_ENTRADA, PASTA_SAIDA,
)


# ── PASTA DE SAÍDA ─────────────────────────────────────────────────────────────

def _caminho_saida_completo(dados: dict, sufixo="") -> str:
    now = __import__("datetime").datetime.now()
    numero_raw = str(dados.get("numero_processo", "S-N"))
    proc_num = numero_raw
    proc_ano = now.strftime("%Y")
    if "/" in numero_raw:
        partes = numero_raw.split("/")
        proc_num, proc_ano = partes[0], partes[-1]
    elif "-" in numero_raw:
        partes = numero_raw.split("-")
        if len(partes) == 2 and partes[0].isdigit() and partes[1].isdigit():
            proc_num, proc_ano = partes[0], partes[-1]

    req = str(dados.get("requerente", "Sem Requerente")).strip().title()
    s = f" - Livre" if not sufixo else f" - {sufixo}"
    
    tipo = str(dados.get("tipo_relatorio", "Parecer")).replace("_", " ").title()
    nome_arquivo = f"{tipo} - {proc_num}-{proc_ano} - {req}{s}.docx"
    nome_pasta = f"Processo {proc_num}-{proc_ano} - {req}"

    for ch in ['<', '>', ':', '"', '|', '?', '*', '\\']:
        nome_arquivo = nome_arquivo.replace(ch, '')
        nome_pasta = nome_pasta.replace(ch, '')

    base = Path(PASTA_SAIDA) / nome_pasta
    base.mkdir(parents=True, exist_ok=True)
    
    return str(base / nome_arquivo)


# ── CONSTRUTOR DO CORPO LIVRE ──────────────────────────────────────────────────

def _build_corpo_livre(doc, dados: dict):
    """
    Renderiza o 'texto_livre' como corpo do parecer.
    Cada linha em branco = novo parágrafo justificado com recuo.
    Suporta **negrito** e __itálico__ inline.
    """
    INDENT = 1.25
    texto = dados.get("texto_livre", "").strip()

    if not texto:
        p = add_para(doc, line=LINE_SPC, before=300, after=PAR_AFTER, indent_cm=INDENT)
        add_run(p, "[Texto livre não fornecido]", size=SZ_CORPO)
        return

    # Divide por linhas em branco (parágrafos)
    paragrafos = [bloco.strip() for bloco in texto.split("\n\n") if bloco.strip()]

    for bloco in paragrafos:
        # Linha simples (título de seção) → negrito centralizado
        linhas = bloco.splitlines()
        if len(linhas) == 1 and bloco.isupper():
            p = add_para(doc, line=LINE_SPC, before=200, after=80)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(bloco)
            from formatacao import set_font
            set_font(r, name=FONT_TITULO, size=SZ_CORPO, bold=True)
            r.font.color.rgb = COR_LABEL_FONT
        else:
            # Parágrafo normal justificado com recuo
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p = add_para(doc, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=INDENT)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Junta as linhas internas com espaço
            texto_bloco = " ".join(l.strip() for l in linhas)
            rich_segments(p, texto_bloco, size=SZ_CORPO)


# ── GERADOR PRINCIPAL ─────────────────────────────────────────────────────────

def gerar_livre(dados: dict, caminho_saida: str = None) -> str:
    """
    Gera o DOCX no modo livre usando texto_livre como corpo.
    Mantém o mesmo layout visual do compilador.py.
    """
    doc = Document()

    # ── Margens ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── Blocos padrão do layout ────────────────────────────────────
    build_header(doc)
    build_titulo(doc, dados.get("titulo_documento", "PARECER SETOR TÉCNICO - SMOSU"))
    build_identificacao(doc, dados)
    build_dados_carimbo(doc, dados)

    # ── Corpo livre ────────────────────────────────────────────────
    _build_corpo_livre(doc, dados)

    # ── Conclusão e documentos ─────────────────────────────────────
    conclusao = dados.get(
        "conclusao",
        "Diante do exposto, verificado e aprovado o projeto, poderá ser emitido:"
    )
    _build_conclusao_bloco(doc, conclusao)

    if dados.get("documentos_emitir"):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p_doc_tit = doc.add_paragraph()
        p_doc_tit.paragraph_format.page_break_before = True
        set_spacing(p_doc_tit, line=LINE_SPC, before=200, after=80)
        r_dt = p_doc_tit.add_run("Emissão de Documentos:")
        from formatacao import set_font
        set_font(r_dt, name=FONT_TITULO, size=SZ_CORPO, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

        for item in dados["documentos_emitir"]:
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    build_assinatura(doc, dados)

    # ── Salvar ─────────────────────────────────────────────────────
    if not caminho_saida:
        caminho_saida = _caminho_saida_completo(dados)

    doc.save(caminho_saida)
    print(f"[+] Documento DOCX (modo livre) gerado: {caminho_saida}")

    # ── Converter para PDF ─────────────────────────────────────────
    try:
        import subprocess
        pdf_path = caminho_saida.replace(".docx", ".pdf")
        subprocess.run(
            ["python", "-c",
             f"from docx2pdf import convert; convert(r'{caminho_saida}', r'{pdf_path}')"],
            capture_output=True, timeout=60
        )
        if Path(pdf_path).exists():
            print(f"[+] Documento PDF  gerado: {pdf_path}")
    except Exception as e:
        print(f"[!] PDF não gerado: {e}")

    return caminho_saida


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    import glob

    args = [a for a in sys.argv[1:] if a != "--sem-preview"]

    if not args:
        # Tenta pasta padrão
        pasta_padrao = Path(PASTA_ENTRADA)
        jsons = list(pasta_padrao.glob("*.json")) if pasta_padrao.exists() else []
        if not jsons:
            print(f"[!] Nenhum JSON encontrado em {os.path.basename(PASTA_ENTRADA)}")
            print("    Uso: python compilador_livre.py dados.json [saida.docx]")
            sys.exit(0)
        args = [str(pasta_padrao)]

    alvo = args[0]
    arquivos = []

    if os.path.isdir(alvo):
        arquivos = list(Path(alvo).glob("*.json"))
    elif os.path.isfile(alvo) and alvo.endswith(".json"):
        arquivos = [Path(alvo)]
    else:
        print(f"[!] Caminho inválido: {alvo}")
        sys.exit(1)

    saida_fornecida = args[1] if len(args) > 1 else None

    for arq in arquivos:
        try:
            with open(arq, encoding="utf-8") as f:
                dados = json.load(f)

            # Valida mínimo
            if not dados.get("texto_livre") and not dados.get("paragrafo_abertura"):
                print(f"[!] {arq.name}: JSON sem 'texto_livre'. Use o compilador.py padrão.")
                continue

            # Se tem texto_livre, garante que funcione como corpo livre
            if not dados.get("texto_livre") and dados.get("paragrafo_abertura"):
                # Compatibilidade: usa paragrafo_abertura como texto_livre
                dados["texto_livre"] = dados["paragrafo_abertura"]

            print(f"\n[>] Modo Livre — {arq.name}")
            print(f"    Processo: {dados.get('numero_processo', '?')}")
            print(f"    Requerente: {dados.get('requerente', '?')}")

            gerar_livre(dados, saida_fornecida)

        except Exception as e:
            print(f"[X] Erro em {arq.name}: {e}")
            import traceback
            traceback.print_exc()

    print("\n" + "=" * 60)
    print("  MODO LIVRE — COMPILAÇÃO ENCERRADA")
    print("=" * 60)


if __name__ == "__main__":
    main()
