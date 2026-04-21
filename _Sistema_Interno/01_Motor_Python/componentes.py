"""
Componentes reutilizáveis para construção de documentos.
Blocos de conteúdo: identificação, dados técnicos, corpo, conclusão, assinatura.
"""

import re
from datetime import datetime
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import (
    FONT_TITULO, FONT_CORPO, SZ_CORPO, SZ_TABELA,
    COR_INST, COR_LABEL_FONT, COR_CINZA_TEXTO,
    PAR_AFTER, LINE_SPC, ASSINANTE, CIDADE,
)
from formatacao import (
    set_spacing, set_font, add_run, add_para, rich_segments,
    apply_table_borders, apply_label_cell, apply_value_cell,
    set_cell_margins, add_separator,
)


# ═══════════════════════════════════════════════════════════
#  UTILITÁRIOS
# ═══════════════════════════════════════════════════════════

def _data_hoje_extenso():
    """Retorna data atual por extenso: '20 de abril de 2026'."""
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]
    d = datetime.now()
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _parse_numero(texto):
    """Extrai valor numérico de strings como '86,23%' ou '180,00m²' → float."""
    if not texto:
        return None
    try:
        limpo = re.sub(r'[^\d,.]', '', str(texto)).replace(',', '.')
        return float(limpo) if limpo else None
    except ValueError:
        return None


def _ensure_list(data):
    """Garante que o dado seja uma lista de strings. Se for string única, separa por quebra de linha."""
    if not data:
        return []
    if isinstance(data, str):
        return [item.strip() for item in data.split('\n') if item.strip()]
    return data


# ═══════════════════════════════════════════════════════════
#  TÍTULO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════

def build_titulo(doc, titulo="PARECER SETOR TÉCNICO - SMOSU"):
    """Título centralizado com fonte Cambria bold."""
    pt = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line=288, before=280, after=220)
    r_tit = add_run(pt, titulo, bold=True, size=14)
    set_font(r_tit, name=FONT_TITULO, size=14, bold=True)
    r_tit.font.color.rgb = COR_LABEL_FONT


# ═══════════════════════════════════════════════════════════
#  IDENTIFICAÇÃO (Processo, Assunto, Requerente)
# ═══════════════════════════════════════════════════════════

def build_identificacao(doc, d):
    """Tabela de identificação com 3 linhas: Processo, Assunto, Requerente."""
    W_LABEL = 2268
    W_VALUE = 7938
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    linhas = [
        (
            "Processo nº",
            f"{d.get('numero_processo', '')}  —  {d.get('data_processo', '')}"
            if d.get('numero_processo') else ""
        ),
        ("Assunto",    d.get("assunto", "")),
        ("Requerente", d.get("requerente", "")),
    ]
    for i, (rot, val) in enumerate(linhas):
        row = tbl.rows[i]
        c_rot = row.cells[0]
        apply_label_cell(c_rot, W_LABEL)
        p_rot = c_rot.paragraphs[0]
        p_rot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_rot, line=280, before=70, after=70)
        r_rot = p_rot.add_run(rot)
        set_font(r_rot, size=SZ_TABELA, bold=True)
        r_rot.font.color.rgb = COR_LABEL_FONT

        c_val = row.cells[1]
        apply_value_cell(c_val, W_VALUE)
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_val, line=280, before=70, after=70)
        rich_segments(p_val, val, size=SZ_TABELA)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  DESTINATÁRIO (para Ofícios)
# ═══════════════════════════════════════════════════════════

def build_destinatario(doc, d):
    """Bloco de destinatário para ofícios e comunicados."""
    if d.get("destinatario_titulo"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=200, after=40)
        add_run(p, d["destinatario_titulo"], bold=True, size=SZ_CORPO)

    if d.get("destinatario_para"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "Para: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_para"], size=SZ_CORPO)

    if d.get("destinatario_de"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=40)
        add_run(p, "De: ", bold=True, size=SZ_CORPO)
        add_run(p, d["destinatario_de"], size=SZ_CORPO)

    if d.get("assunto"):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                     line=LINE_SPC, before=0, after=80)
        add_run(p, "Assunto: ", bold=True, size=SZ_CORPO)
        add_run(p, d["assunto"], size=SZ_CORPO)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  DADOS DO CARIMBO TÉCNICO (Decreto 4.149/2019)
# ═══════════════════════════════════════════════════════════

def build_dados_carimbo(doc, d):
    """Tabela 4 colunas com dados técnicos do projeto."""
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=240, before=280, after=60)
    r_tit = add_run(p_tit, "DADOS TÉCNICOS DO PROJETO (Ref. Decreto nº 4.149/2019)",
                    bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(tbl)

    logradouro = d.get('logradouro', '')
    bairro = d.get('bairro', '')
    end_completo = (f"{logradouro} - {bairro}" if logradouro and bairro
                    else logradouro or bairro)

    # Detectar alertas urbanísticos
    to_val = _parse_numero(d.get('taxa_ocupacao', ''))
    perm_val = _parse_numero(d.get('taxa_permeabilidade', ''))
    terreno_val = _parse_numero(d.get('area_terreno', ''))
    alerta_zoneamento = (
        (to_val is not None and to_val > 70)
        and (perm_val is not None and perm_val < 20)
    )
    alerta_excecao = (terreno_val is not None and terreno_val < 220)

    linhas = [
        ("Endereço:",       end_completo,
         "Inscrição Mun.:", d.get("inscricao_municipal", "")),
        ("Proprietário:",   d.get("proprietario", d.get("requerente", "")),
         "Desenhista:",     d.get("desenhista", "")),
        ("Lote:",           d.get("lote", ""),
         "Quadra:",         d.get("quadra", "")),
        ("Área Terreno:",   d.get("area_terreno", ""),
         "Área Total:",     d.get("area_total_construida", "")),
        ("Taxa Ocupação:",  d.get("taxa_ocupacao", ""),
         "Coef. Aprov.:",   d.get("coef_aproveitamento", "")),
        ("Permeabilidade:", d.get("taxa_permeabilidade", ""),
         "Resp. Técnico:",  d.get("profissional_nome", "")),
    ]

    W_L1 = 2000
    W_V1 = 3500
    W_L2 = 2000
    W_V2 = 2700

    for lab1, val1, lab2, val2 in linhas:
        row = tbl.add_row()

        # Label 1
        c0 = row.cells[0]
        apply_label_cell(c0, W_L1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p0, line=220, before=25, after=25)
        r0 = p0.add_run(lab1)
        set_font(r0, size=SZ_TABELA, bold=True)
        r0.font.color.rgb = COR_LABEL_FONT

        # Valor 1
        c1 = row.cells[1]
        apply_value_cell(c1, W_V1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p1, line=220, before=25, after=25)
        
        cor_v1 = None
        alerta_msg_v1 = None

        # Alerta: Área do Terreno < 220m² → verde + "exceção da lei"
        if lab1 == "Área Terreno:" and alerta_excecao:
            cor_v1 = RGBColor(0x00, 0x80, 0x00)
            alerta_msg_v1 = "  (exceção da lei)"

        # Alerta: Taxa Ocupação > 70% → vermelho + "conferir zoneamento"
        if lab1 == "Taxa Ocupação:" and alerta_zoneamento:
            cor_v1 = RGBColor(0xCC, 0x00, 0x00)
            alerta_msg_v1 = "  (conferir zoneamento)"

        # Alerta: Permeabilidade < 20% → vermelho
        if lab1 == "Permeabilidade:" and alerta_zoneamento:
            cor_v1 = RGBColor(0xCC, 0x00, 0x00)
            alerta_msg_v1 = "  (conferir zoneamento)"

        rich_segments(p1, val1, size=SZ_TABELA, color=cor_v1)
        
        if alerta_msg_v1:
            r_alert = p1.add_run(alerta_msg_v1)
            set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = cor_v1

        # Label 2
        c2 = row.cells[2]
        apply_label_cell(c2, W_L2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p2, line=220, before=25, after=25)
        r2 = p2.add_run(lab2)
        set_font(r2, size=SZ_TABELA, bold=True)
        r2.font.color.rgb = COR_LABEL_FONT

        # Valor 2
        c3 = row.cells[3]
        apply_value_cell(c3, W_V2)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p3, line=220, before=25, after=25)
        rich_segments(p3, val2, size=SZ_TABELA)

    add_separator(doc)


# ═══════════════════════════════════════════════════════════
#  CORPO DO PARECER (abertura + considerandos)
# ═══════════════════════════════════════════════════════════

def build_corpo(doc, d):
    """Corpo do parecer: abertura, considerandos e parágrafos adicionais."""
    INDENT = 1.25

    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=300,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    if d.get("considerandos"):
        for cons in _ensure_list(d["considerandos"]):
            cons_limpo = cons
            # Remove o "Considerando que" se o GEM já tiver escrito na string
            if cons_limpo.lower().startswith("considerando que "):
                cons_limpo = cons_limpo[17:].strip()
            elif cons_limpo.lower().startswith("considerando "):
                cons_limpo = cons_limpo[13:].strip()
                
            p = add_para(doc, line=LINE_SPC, before=0,
                         after=PAR_AFTER, indent_cm=INDENT)
            add_run(p, "Considerando que ", bold=True, size=SZ_CORPO)
            rich_segments(p, cons_limpo, size=SZ_CORPO)

    if d.get("paragrafos_adicionais"):
        for txt in _ensure_list(d["paragrafos_adicionais"]):
            p_extra = add_para(doc, line=LINE_SPC, before=0,
                               after=PAR_AFTER, indent_cm=INDENT)
            rich_segments(p_extra, txt, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO LEGAL
# ═══════════════════════════════════════════════════════════

def build_fundamentacao(doc, d):
    """Seção de fundamentação legal com bullets."""
    if not d.get("fundamentacao_legal"):
        return

    INDENT = 1.25
    p_fund_tit = add_para(doc, line=LINE_SPC, before=200,
                          after=80, indent_cm=INDENT)
    r_tit = add_run(p_fund_tit, "Da Análise Legal e Técnica:",
                    bold=True, size=SZ_CORPO)
    r_tit.font.color.rgb = COR_LABEL_FONT

    for fund in _ensure_list(d["fundamentacao_legal"]):
        fund_limpo = fund.lstrip("•- \t") # Limpa marcadores se o GEM tiver colocado
        p_fund = add_para(doc, line=LINE_SPC, before=0,
                          after=PAR_AFTER, indent_cm=INDENT)
        add_run(p_fund, "• ", size=SZ_CORPO)
        rich_segments(p_fund, fund_limpo, size=SZ_CORPO)


# ═══════════════════════════════════════════════════════════
#  CARDS DE DOCUMENTOS A EMITIR
# ═══════════════════════════════════════════════════════════

def add_doc_item(doc, tipo, obs=None):
    """Adiciona um item de documento como card com borda lateral azul."""
    card = doc.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Bordas do card: borda esquerda grossa azul, restante cinza fina
    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        card._tbl.insert(0, tblPr)
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'D6DCE4')
        el.set(qn('w:space'), '0')
        tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '18')
    left_b.set(qn('w:color'), COR_INST)
    left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo levemente cinza
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F8F9FA')
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

    # Título do documento
    p_tipo = cell.paragraphs[0]
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_tipo, line=LINE_SPC, before=40, after=20)
    add_run(p_tipo, tipo, bold=True, size=11)

    # Observação
    if obs:
        p_obs = cell.add_paragraph()
        p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p_obs, line=280, before=0, after=40)
        rich_segments(p_obs, obs, size=10, color=COR_CINZA_TEXTO)

    # Espaçamento depois do card
    p_space = doc.add_paragraph()
    set_spacing(p_space, line=80, before=0, after=0)


# ═══════════════════════════════════════════════════════════
#  CONCLUSÃO + DOCUMENTOS + ASSINATURA
# ═══════════════════════════════════════════════════════════

def build_conclusao_e_docs(doc, d):
    """Conclusão completa: fundamentação + conclusão + lista de docs + assinatura."""
    INDENT = 1.25

    # 1º Fundamentação Legal
    build_fundamentacao(doc, d)

    # 2º Título "Emissão de Documentos"
    if d.get("documentos_emitir"):
        p_doc_tit = add_para(doc, line=LINE_SPC, before=200,
                             after=40, indent_cm=INDENT)
        r_dt = add_run(p_doc_tit, "Emissão de Documentos:",
                       bold=True, size=SZ_CORPO)
        set_font(r_dt, name=FONT_TITULO, size=12, bold=True)
        r_dt.font.color.rgb = COR_LABEL_FONT

    # 3º Conclusão ("Diante do exposto...")
    conclusao = d.get(
        "conclusao",
        "Diante do exposto, após análise das documentações apresentadas, conclui-se que o "
        "projeto está apto a ser aprovado quanto às questões técnicas pertinentes, podendo "
        "ser emitidos os seguintes documentos:"
    )
    p_conc = add_para(doc, line=LINE_SPC, before=0,
                      after=PAR_AFTER, indent_cm=INDENT)
    rich_segments(p_conc, conclusao, size=SZ_CORPO)

    # 4º Lista de documentos (cards estilizados)
    for item in d.get("documentos_emitir", []):
        add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    # 5º Assinatura
    build_assinatura(doc, d)


def build_conclusao_simples(doc, d):
    """Conclusão simples sem lista de documentos em cards."""
    INDENT = 1.25

    if d.get("conclusao"):
        p_conc = add_para(doc, line=LINE_SPC, before=200,
                          after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # Se houver documentos, usar cards
    if d.get("documentos_emitir"):
        for item in d.get("documentos_emitir", []):
            add_doc_item(doc, item.get("tipo", ""), item.get("obs") or None)

    build_assinatura(doc, d)


# ═══════════════════════════════════════════════════════════
#  ASSINATURA
# ═══════════════════════════════════════════════════════════

def build_assinatura(doc, d):
    """Bloco de assinatura: data + linha + nome/título/registro."""
    assinante = d.get("assinante", ASSINANTE)
    cidade = d.get("cidade", CIDADE)

    # Local e data (centralizado, negrito)
    p_local = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       line=LINE_SPC, before=200, after=0)
    add_run(p_local, f"{cidade}, {_data_hoje_extenso()}.", bold=True, size=SZ_CORPO)
    p_local.paragraph_format.keep_with_next = True

    # Espaço
    for _ in range(2):
        pb = doc.add_paragraph()
        set_spacing(pb, line=LINE_SPC, before=0, after=0)
        pb.paragraph_format.keep_with_next = True

    # Linha de assinatura
    p_line = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line=240, before=0, after=40)
    r_line = p_line.add_run('_' * 45)
    set_font(r_line, size=10)
    r_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p_line.paragraph_format.keep_with_next = True

    # Dados do assinante
    assinatura_items = [
        (assinante.get("nome",     ASSINANTE["nome"]),     True),
        (assinante.get("titulo",   ASSINANTE["titulo"]),   False),
        (assinante.get("registro", ASSINANTE["registro"]), False),
    ]
    for i, (texto, bold) in enumerate(assinatura_items):
        p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     line=LINE_SPC, before=0, after=0)
        add_run(p, texto, bold=bold, size=SZ_CORPO)
        if i < len(assinatura_items) - 1:
            p.paragraph_format.keep_with_next = True


# ═══════════════════════════════════════════════════════════
#  COMUNICADO DE PENDÊNCIA — Design Visual Dedicado
# ═══════════════════════════════════════════════════════════

def _box_colorido(doc, fill_hex, borda_hex, borda_sz='12'):
    """Cria uma tabela-card de 1 célula com fundo colorido e borda lateral."""
    card = doc.add_table(rows=1, cols=1)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        card._tbl.insert(0, tblPr)

    # Largura total
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')   # ~16,5 cm em twips
    tblW.set(qn('w:type'), 'dxa')
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)

    # Bordas
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), borda_hex)
        el.set(qn('w:space'), '0')
        tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), borda_sz)
    left_b.set(qn('w:color'), borda_hex)
    left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=100)

    return card, cell


def build_comunicado_pendencia(doc, d):
    """
    Layout visual dedicado para Comunicado de Pendência Documental.

    Estrutura:
      ① Parágrafo de abertura (texto institucional normal)
      ② Box LARANJA — lista numerada de pendências
      ③ Box VERDE   — orientação de continuidade ao requerente
      ④ Conclusão
      ⑤ Assinatura
    """
    INDENT = 1.25
    COR_ALERTA_FILL  = 'FFF3CD'   # amarelo-âmbar claro
    COR_ALERTA_BORDA = 'C07800'   # âmbar escuro
    COR_OK_FILL      = 'EBF5EB'   # verde muito claro
    COR_OK_BORDA     = '2E7D32'   # verde institucional

    # ── ① Parágrafo de abertura ────────────────────────────────────────────
    if d.get("paragrafo_abertura"):
        p_ab = add_para(doc, line=LINE_SPC, before=240,
                        after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    # ── ② Box de Pendências (laranja/âmbar) ───────────────────────────────
    p_esp = doc.add_paragraph()
    set_spacing(p_esp, line=80, before=0, after=0)

    card_alert, cell_alert = _box_colorido(doc, COR_ALERTA_FILL, COR_ALERTA_BORDA, '18')

    # Título do box
    p_tit = cell_alert.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_tit, line=260, before=60, after=60)
    r_icon = p_tit.add_run("⚠  DOCUMENTOS PENDENTES — ANÁLISE SUSPENSA")
    set_font(r_icon, name="Cambria", size=11, bold=True)
    r_icon.font.color.rgb = RGBColor(0x7B, 0x4F, 0x00)

    # Subtítulo
    p_sub = cell_alert.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_sub, line=240, before=0, after=80)
    r_sub = p_sub.add_run(
        "A análise técnica não pôde ser iniciada. "
        "Os documentos abaixo são imprescindíveis e devem ser apresentados "
        "para o prosseguimento do processo:"
    )
    set_font(r_sub, size=10, italic=True)
    r_sub.font.color.rgb = RGBColor(0x5C, 0x3A, 0x00)

    # Lista de pendências
    considerandos = _ensure_list(d.get("considerandos", []))
    for item in considerandos:
        p_item = cell_alert.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_item, line=260, before=20, after=20)
        # Ícone de marcador
        r_bullet = p_item.add_run("  ✗  ")
        set_font(r_bullet, size=10, bold=True)
        r_bullet.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
        # Texto do item (sem "Considerando que" — é uma lista de pendências)
        texto = item.lstrip("0123456789. ")  # remove "1. ", "2. " se houver
        if texto.lower().startswith("considerando que "):
            texto = texto[17:].strip()
        elif texto.lower().startswith("considerando "):
            texto = texto[13:].strip()
            
        rich_segments(p_item, texto, size=10, color=RGBColor(0x3C, 0x28, 0x00))

    # Espaço após o box
    p_esp2 = doc.add_paragraph()
    set_spacing(p_esp2, line=80, before=0, after=0)

    # ── ③ Box de Orientação (verde) ────────────────────────────────────────
    card_ok, cell_ok = _box_colorido(doc, COR_OK_FILL, COR_OK_BORDA, '18')

    p_ok_tit = cell_ok.paragraphs[0]
    p_ok_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_ok_tit, line=260, before=60, after=40)
    r_ok_icon = p_ok_tit.add_run("✔  COMO REGULARIZAR E RETOMAR O PROCESSO")
    set_font(r_ok_icon, name="Cambria", size=11, bold=True)
    r_ok_icon.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    orientacoes = [
        "Reúna todos os documentos listados acima em sua versão original ou digitalizada com legibilidade plena.",
        "Dirija-se ao balcão do protocolo da Prefeitura Municipal de Oliveira ou acesse o portal digital (Atende.Net) para acostar os documentos ao processo.",
        "Após o recebimento e verificação documental, a análise técnica será reiniciada automaticamente, sem necessidade de novo requerimento.",
        "Em caso de dúvidas, entre em contato com o Setor Técnico da SMOSU (Secretaria Municipal de Obras e Serviços Urbanos).",
    ]
    for i, orient in enumerate(orientacoes, 1):
        p_or = cell_ok.add_paragraph()
        p_or.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p_or, line=260, before=20, after=20)
        r_num = p_or.add_run(f"  {i}.  ")
        set_font(r_num, size=10, bold=True)
        r_num.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        r_or = p_or.add_run(orient)
        set_font(r_or, size=10)
        r_or.font.color.rgb = RGBColor(0x1A, 0x37, 0x1A)

    # Espaço após box verde
    p_esp3 = doc.add_paragraph()
    set_spacing(p_esp3, line=120, before=0, after=0)

    # ── ④ Conclusão ────────────────────────────────────────────────────────
    if d.get("conclusao"):
        p_conc = add_para(doc, line=LINE_SPC, before=120,
                          after=PAR_AFTER, indent_cm=INDENT)
        rich_segments(p_conc, d["conclusao"], size=SZ_CORPO)

    # ── ⑤ Assinatura ───────────────────────────────────────────────────────
    build_assinatura(doc, d)

