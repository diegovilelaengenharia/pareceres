"""
Gerador de Layout Exclusivo para os Alvarás Oficiais (Secretaria)
Gera o Alvará de Construção e a Carta de Habite-se final com a tipografia
de apresentação de CPF, Matrizes e Fim de Emissão em Papel A4 timbrado.
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from formatacao import set_font, add_run, apply_table_borders, set_spacing
from config import FONT_CORPO, COR_INST, COR_LABEL_FONT

def build_alvara_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'alvara_oficial.json'.
    Foca nos dados: Nome, CPF, Autores, CREA e dados da Obra.
    """
    _build_cabecalho_simples(doc, dados)

    # Bloco "PROPRIETÁRIO:"
    _add_labeled_pair(doc, "PROPRIETÁRIO:", "NOME", dados.get("proprietario_nome"), "CPF/CNPJ", dados.get("proprietario_cpf_cnpj"))
    
    # Bloco "AUTOR DO PROJETO:"
    _add_labeled_pair(doc, "AUTOR DO PROJETO:", "NOME", dados.get("autor_projeto_nome"), 
                      "CREA/ART", f"{dados.get('autor_projeto_crea', '')} / {dados.get('autor_projeto_art', '')}")
    
    # Bloco "RESPONSÁVEL TÉCNICO:"
    _add_labeled_pair(doc, "RESPONSÁVEL TÉCNICO:", "NOME", dados.get("responsavel_tecnico_nome"), 
                      "CREA/ART", f"{dados.get('responsavel_tecnico_crea', '')} / {dados.get('responsavel_tecnico_art', '')}")
    
    # Bloco "CONSTRUTORA:"
    if dados.get("construtora_nome"):
        _add_labeled_pair(doc, "CONSTRUTORA OU RESPONSÁVEL PELA EXECUÇÃO:", "NOME", dados.get("construtora_nome"), 
                          "CPF/CNPJ", dados.get("construtora_cpf_cnpj"))

    # Texto Livre (Concessão)
    t_processo = dados.get("numero_processo", "")
    t_data = dados.get("data_aprovacao", "")
    t_obra = dados.get("nome_obra", "")
    t_endereco = f"{dados.get('logradouro', '')}, {dados.get('bairro', '')}"
    texto_abertura = f"Tendo em vista o constante no processo nº {t_processo} fica concedida a licença para execução do projeto aprovado em {t_data} a obra denominada de {t_obra}, a ser executada no endereço: {t_endereco}."
    
    p = doc.add_paragraph()
    set_spacing(p, before=100, after=200, line=240)
    add_run(p, texto_abertura, size=11).font.color.rgb = RGBColor(0,0,0)
    
    # Matriz da Obra
    _build_matriz_areas(doc, dados)

    # Assinatura (Rodapé Simplificado do Alvará)
    _build_assinatura_balcao(doc)


def build_habitese_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'carta_habitese_oficial.json'.
    Foca no endereço em cima e despacho.
    """
    _build_cabecalho_simples(doc, dados)

    # Endereço
    p_end = doc.add_paragraph()
    add_run(p_end, "ENDEREÇO DA OBRA: ", bold=True)
    add_run(p_end, dados.get("logradouro", ""))
    p_end2 = doc.add_paragraph()
    add_run(p_end2, "BAIRRO: ", bold=True)
    add_run(p_end2, dados.get("bairro", ""))

    # Proprietário
    _add_labeled_pair(doc, "PROPRIETÁRIO DO IMÓVEL:", "NOME", dados.get("proprietario_nome"), "CPF", dados.get("proprietario_cpf_cnpj"))
    _add_labeled_pair(doc, "RESPONSÁVEL PELA EXECUÇÃO DA OBRA:", "NOME", dados.get("responsavel_execucao_nome"), "CPF", dados.get("responsavel_execucao_cpf_cnpj"))
    
    # Responsável
    p_resp = doc.add_paragraph()
    set_spacing(p_resp, before=100, after=50)
    add_run(p_resp, "RESPONSÁVEL TÉCNICO:", bold=True, size=11)
    ptx = doc.add_paragraph()
    set_spacing(ptx, before=0, after=200, line=240)
    add_run(ptx, dados.get("texto_despacho_responsavel_tecnico", ""))

    # Matriz de Áreas
    _build_matriz_areas(doc, dados)
    
    _build_assinatura_balcao(doc)


def build_certidao_oficial(doc, dados):
    """
    Roda a diagramação baseada no 'certidao_oficial.json'.
    Foca no texto corrido justificado e assinaturas lado a lado.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=100, after=400)
    add_run(p, dados.get("titulo_documento", "CERTIDÃO"), size=16, bold=True).font.color.rgb = COR_LABEL_FONT
    
    p_text = doc.add_paragraph()
    p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_text, before=200, after=400, line=360)
    add_run(p_text, "        " + dados.get("texto_certidao", ""), size=12).font.color.rgb = RGBColor(0,0,0)
    
    from datetime import datetime
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    d = datetime.now()
    hoje = f"Oliveira, {d.day} de {meses[d.month - 1]} de {d.year}."
    
    p_dt = doc.add_paragraph()
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(p_dt, before=200, after=600)
    add_run(p_dt, hoje, size=11).font.color.rgb = RGBColor(0,0,0)
    
    assinantes = dados.get("assinantes", [])
    if assinantes:
        p_ass = doc.add_paragraph()
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p_ass, before=400, after=200)
        for ass in assinantes:
            add_run(p_ass, str(ass.get("nome", "")) + "\n", bold=True, size=11)
            add_run(p_ass, str(ass.get("titulo", "")) + "\n\n", size=10)
    
    obs = dados.get("observacoes_finais", [])
    if isinstance(obs, list):
        for o in obs:
            po = doc.add_paragraph()
            po.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(po, o, size=9, bold=True).font.color.rgb = RGBColor(0,0,0)


# Helpers Visuais de Guias Finais

def _build_cabecalho_simples(doc, dados):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "MUNICIPIO DE OLIVEIRA\nSecretaria Municipal de Obras", size=12, bold=True)
    
    num_doc = dados.get("numero_documento", "XXX/XXXX")
    tit = str(dados.get("titulo_documento", "DOCUMENTO Nº {numero}")).replace("{numero}", num_doc)
    
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(pt, before=100, after=250)
    rt = add_run(pt, tit, size=14, bold=True)
    rt.font.color.rgb = COR_LABEL_FONT


def _add_labeled_pair(doc, main_label, l1, v1, l2, v2):
    p = doc.add_paragraph()
    set_spacing(p, before=150, after=30)
    add_run(p, main_label, bold=True, size=11)
    
    p1 = doc.add_paragraph()
    set_spacing(p1, before=0, after=0)
    add_run(p1, f"{l1}: ", bold=True, size=10)
    add_run(p1, v1 if v1 else "---", size=10)
    
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=0)
    add_run(p2, f"{l2}: ", bold=True, size=10)
    add_run(p2, v2 if v2 else "---", size=10)


def _build_matriz_areas(doc, dados):
    p = doc.add_paragraph()
    set_spacing(p, before=150, after=50)
    add_run(p, "Dados da obra:\nÁreas principais", bold=True, size=11)
    
    matriz = dados.get("areas_matriz", [])
    if matriz:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        apply_table_borders(tbl)
        
        # Cabeçalhos
        for i, h in enumerate(["Categoria", "Destinação", "Tipo de Obra", "Área (m²)"]):
            c = tbl.cell(0, i)
            pt = c.paragraphs[0]
            add_run(pt, h, bold=True, size=9)
            
        for row in matriz:
            tr = tbl.add_row()
            if isinstance(row, dict):
                add_run(tr.cells[0].paragraphs[0], str(row.get("categoria","")), size=9)
                add_run(tr.cells[1].paragraphs[0], str(row.get("destinacao","")), size=9)
                add_run(tr.cells[2].paragraphs[0], str(row.get("tipo_obra","")), size=9)
                add_run(tr.cells[3].paragraphs[0], str(row.get("area_m2","")), size=9)

    p_tot = doc.add_paragraph()
    set_spacing(p_tot, before=150, after=50)
    add_run(p_tot, "Área total da obra: ", bold=True, size=11)
    add_run(p_tot, str(dados.get("area_total_obra", "")), size=11)
    
    if dados.get("observacoes"):
        po = doc.add_paragraph()
        set_spacing(po, before=100, after=200, line=240)
        add_run(po, "Observações: ", bold=True, size=10)
        add_run(po, dados.get("observacoes", ""), size=10)

def _build_assinatura_balcao(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=400, after=400)
    
    p_ass = doc.add_paragraph()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_ass, "_" * 45, size=10)
    p_as2 = doc.add_paragraph()
    p_as2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_as2, "SECRETARIA MUNICIPAL DE OBRAS / PMO", bold=True, size=10)
