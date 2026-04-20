"""
Funções auxiliares de formatação DOCX.
Helpers de baixo nível para fontes, espaçamento, tabelas e campos.
"""

import re
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import (
    FONT_CORPO, SZ_CORPO, COR_INST, COR_LABEL_BG,
    COR_BORDA_TABELA, COR_LABEL_FONT, PAR_AFTER, LINE_SPC,
)


# ═══════════════════════════════════════════════════════════
#  ESPAÇAMENTO E FONTE
# ═══════════════════════════════════════════════════════════

def set_spacing(p, line=276, before=0, after=0):
    """Define espaçamento de um parágrafo."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    pPr.append(sp)


def set_font(run, name=None, size=SZ_CORPO, bold=False, underline=False, italic=False):
    """Aplica fonte, tamanho e estilo a um run."""
    fname = name or FONT_CORPO
    run.font.name      = fname
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.underline = underline
    run.font.italic    = italic
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), fname)
    rf.set(qn('w:hAnsi'), fname)
    rf.set(qn('w:cs'), fname)
    rPr.insert(0, rf)


# ═══════════════════════════════════════════════════════════
#  TEXTO E PARÁGRAFOS
# ═══════════════════════════════════════════════════════════

def clean(t):
    """Remove colchetes espúrios do texto."""
    if not isinstance(t, str):
        return t
    return re.sub(r'\s*\]+\]\s*', ' ', t)


def add_run(p, text, bold=False, size=None, underline=False, italic=False):
    """Adiciona run formatado a um parágrafo."""
    r = p.add_run(clean(text))
    set_font(r, bold=bold, size=size or SZ_CORPO, underline=underline, italic=italic)
    return r


def add_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_SPC,
             before=0, after=PAR_AFTER, indent_cm=None):
    """Cria parágrafo formatado."""
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, line=line, before=before, after=after)
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p


def bold_segments(p, text, size=None, color=None):
    """Processa **negrito** inline em um parágrafo."""
    for i, part in enumerate(text.split('**')):
        if part:
            r = add_run(p, part, bold=(i % 2 == 1), size=size)
            if color is not None:
                r.font.color.rgb = color


def rich_segments(p, text, size=None, color=None):
    """Processa **negrito** e __itálico__ inline em um parágrafo."""
    bold_parts = text.split('**')
    for bi, bpart in enumerate(bold_parts):
        is_bold = (bi % 2 == 1)
        italic_parts = bpart.split('__')
        for ii, ipart in enumerate(italic_parts):
            if ipart:
                is_italic = (ii % 2 == 1)
                r = add_run(p, ipart, bold=is_bold, italic=is_italic, size=size)
                if color is not None:
                    r.font.color.rgb = color


# ═══════════════════════════════════════════════════════════
#  TABELAS E CÉLULAS
# ═══════════════════════════════════════════════════════════

def no_borders(cell):
    """Remove todas as bordas de uma célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        b.append(el)
    tcPr.append(b)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """Define margens internas de uma célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def apply_table_borders(tbl):
    """Aplica bordas padrão cinza a uma tabela."""
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), COR_BORDA_TABELA)
        tblBrd.append(el)
    tblPr.append(tblBrd)


def apply_label_cell(cell, width_twips):
    """Formata célula de rótulo (fundo cinza, texto azul)."""
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), COR_LABEL_BG)
    tcPr.append(shd)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def apply_value_cell(cell, width_twips):
    """Formata célula de valor (fundo branco)."""
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ═══════════════════════════════════════════════════════════
#  ELEMENTOS VISUAIS
# ═══════════════════════════════════════════════════════════

def add_separator(doc, color=COR_INST):
    """Adiciona linha separadora colorida."""
    p_sep = doc.add_paragraph()
    set_spacing(p_sep, line=60, before=80, after=60)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pb_sep = OxmlElement('w:pBdr')
    bot_sep = OxmlElement('w:bottom')
    bot_sep.set(qn('w:val'), 'single')
    bot_sep.set(qn('w:sz'), '4')
    bot_sep.set(qn('w:color'), color)
    pb_sep.append(bot_sep)
    pPr_sep.append(pb_sep)


def add_field(r, instr_text):
    """Insere campo do Word (PAGE, NUMPAGES, etc.)."""
    for tag, txt in [('begin', None), ('instrText', instr_text), ('separate', None), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = txt
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        r._r.append(el)


def add_page_break(doc):
    """Insere uma quebra de página explícita."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_spacing(p, line=0, before=0, after=0)
