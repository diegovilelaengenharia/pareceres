"""
Gerador Automático de Pareceres – Prefeitura de Oliveira / SMOSU
Versão Final: Tabelas 8pt, corpo 11pt, citações 8pt, cabeçalho premium, quebra de página nos docs.
"""

import sys, json, os, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR      = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR     = os.path.dirname(SCRIPT_DIR)  # pasta raiz: 02. Pareceres
LOGO_BRASAO     = os.path.join(SCRIPT_DIR, "logos", "logo_brasao.jpeg")
LOGO_PREFEITURA = os.path.join(SCRIPT_DIR, "logos", "logo_prefeitura.png")

# ── Tipografia moderna ────────────────────────────────────
# • Corpo do texto:   Calibri 11pt  — sans-serif moderna, excelente leitura
# • Títulos/Labels:    Cambria bold  — serifa clássica, peso institucional
# • Cabeçalho:         Cambria (nome) + Calibri (detalhes)
# • Tabelas:           Calibri 8pt   — compacta e limpa
FONT_CORPO   = "Calibri"     # sans-serif para texto corrido
FONT_TITULO  = "Cambria"     # serifa para títulos e peso institucional
FONT_HEADER  = "Cambria"     # serifa para nome da prefeitura
FONT_DETALHE = "Calibri"     # sans-serif para contato, notas

SZ_CORPO   = 11
SZ_TABELA  = 9
SZ_CITACAO = 8
SZ_NOTA    = 8
SZ_RODAPE  = 8

ASSINANTE = {
    "nome":     "Diego Tarcísio Nunes Vilela",
    "titulo":   "Engenheiro Civil",
    "registro": "CREA 235.474/D",
}
CIDADE = "Oliveira"

# Cores padronizadas
COR_INST         = '1F3864'
COR_LABEL_BG     = 'D6DCE4'
COR_BORDA_TABELA = 'C0C0C0'
COR_LABEL_FONT   = RGBColor(0x1F, 0x38, 0x64)
COR_CINZA_TEXTO  = RGBColor(0x44, 0x44, 0x44)
COR_CINZA_LEVE   = RGBColor(0x66, 0x66, 0x66)

PAR_AFTER = 120
LINE_SPC  = 288

def _data_hoje_extenso():
    meses = ["janeiro","fevereiro","março","abril","maio","junho",
              "julho","agosto","setembro","outubro","novembro","dezembro"]
    d = datetime.now()
    return f"{d.day} de {meses[d.month-1]} de {d.year}"

# ═══════════════════════════════════════════════════════════
#  HELPERS E FORMATAÇÃO BÁSICA
# ═══════════════════════════════════════════════════════════
def _set_spacing(p, line=276, before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    pPr.append(sp)

def _set_font(run, name=None, size=SZ_CORPO, bold=False, underline=False, italic=False):
    fname = name or FONT_CORPO
    run.font.name      = fname
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.underline = underline
    run.font.italic    = italic
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')): rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), fname); rf.set(qn('w:hAnsi'), fname); rf.set(qn('w:cs'), fname)
    rPr.insert(0, rf)

def _clean(t):
    if not isinstance(t, str): return t
    return re.sub(r'\s*\]+\]\s*', ' ', t)

def _run(p, text, bold=False, size=None, underline=False, italic=False):
    r = p.add_run(_clean(text))
    _set_font(r, bold=bold, size=size or SZ_CORPO, underline=underline, italic=italic)
    return r

def _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=None):
    p = doc.add_paragraph()
    p.alignment = align
    _set_spacing(p, line=line, before=before, after=after)
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p

def _bold_segments(p, text, size=None, color=None):
    for i, part in enumerate(text.split('**')):
        if part:
            r = _run(p, part, bold=(i % 2 == 1), size=size)
            if color is not None:
                r.font.color.rgb = color

def _rich_segments(p, text, size=None, color=None):
    """Processa **negrito** e __itálico__ inline, na mesma linha."""
    bold_parts = text.split('**')
    for bi, bpart in enumerate(bold_parts):
        is_bold = (bi % 2 == 1)
        italic_parts = bpart.split('__')
        for ii, ipart in enumerate(italic_parts):
            if ipart:
                is_italic = (ii % 2 == 1)
                r = _run(p, ipart, bold=is_bold, italic=is_italic, size=size)
                if color is not None:
                    r.font.color.rgb = color

def _no_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); b.append(el)
    tcPr.append(b)

def _set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('start',left),('end',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def _apply_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'4'); el.set(qn('w:color'), COR_BORDA_TABELA)
        tblBrd.append(el)
    tblPr.append(tblBrd)

def _apply_label_cell(cell, width_twips):
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), COR_LABEL_BG)
    tcPr.append(shd)
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(width_twips)); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def _apply_value_cell(cell, width_twips):
    cell.width = Twips(width_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(width_twips)); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def _add_separator(doc, color=COR_INST):
    p_sep = doc.add_paragraph(); _set_spacing(p_sep, line=60, before=80, after=60)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pb_sep = OxmlElement('w:pBdr')
    bot_sep = OxmlElement('w:bottom')
    bot_sep.set(qn('w:val'),'single'); bot_sep.set(qn('w:sz'),'4')
    bot_sep.set(qn('w:color'), color)
    pb_sep.append(bot_sep); pPr_sep.append(pb_sep)

def _field(r, instr_text):
    for tag, txt in [('begin',None),('instrText',instr_text),('separate',None),('end',None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'),'preserve'); el.text = txt
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        r._r.append(el)

def _add_page_break(doc):
    """Insere uma quebra de página explícita."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    # Remove o parágrafo vazio que ficaria antes da quebra
    _set_spacing(p, line=0, before=0, after=0)

def _add_page_number_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs: p._element.getparent().remove(p._element)
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(fp, line=240, before=0, after=0)

    r1 = fp.add_run('Página ')
    _set_font(r1, size=SZ_RODAPE); r1.font.color.rgb = COR_CINZA_LEVE
    r2 = fp.add_run(); _set_font(r2, size=SZ_RODAPE); r2.font.color.rgb = COR_CINZA_LEVE
    _field(r2, ' PAGE ')
    r3 = fp.add_run(' de ')
    _set_font(r3, size=SZ_RODAPE); r3.font.color.rgb = COR_CINZA_LEVE
    r4 = fp.add_run(); _set_font(r4, size=SZ_RODAPE); r4.font.color.rgb = COR_CINZA_LEVE
    _field(r4, ' NUMPAGES ')

# ═══════════════════════════════════════════════════════════
#  CABEÇALHO MODERNO
# ═══════════════════════════════════════════════════════════
def _build_header(doc):
    section = doc.sections[0]
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs: p._element.getparent().remove(p._element)

    W_LOGO = 2400; W_TEXT = 7800
    total  = W_LOGO + W_TEXT

    tbl = hdr.add_table(rows=1, cols=2, width=Twips(total))
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = tbl._tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); tblBrd.append(el)
    tblPr.append(tblBrd)

    row = tbl.rows[0]; row.height = Twips(1200)

    # ── Logo Prefeitura (à esquerda) ──
    c0 = row.cells[0]; c0.width = Twips(W_LOGO)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; _no_borders(c0)
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p0, line=240, before=0, after=0)
    if os.path.exists(LOGO_PREFEITURA): p0.add_run().add_picture(LOGO_PREFEITURA, width=Cm(3.3))

    # ── Bloco de texto (à direita do logo) ──
    c1 = row.cells[1]; c1.width = Twips(W_TEXT)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; _no_borders(c1)
    for extra in c1.paragraphs[1:]: extra._element.getparent().remove(extra._element)

    # L1: PREFEITURA MUNICIPAL DE OLIVEIRA
    p1a = c1.paragraphs[0]; p1a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p1a, line=320, before=0, after=30)
    r_pref = p1a.add_run("PREFEITURA MUNICIPAL DE OLIVEIRA")
    _set_font(r_pref, name=FONT_HEADER, size=15, bold=True)
    r_pref.font.color.rgb = COR_LABEL_FONT

    # L2: Secretaria
    p1c = c1.add_paragraph(); p1c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p1c, line=300, before=0, after=30)
    r_sec = p1c.add_run("Secretaria Municipal de Obras e Serviços Urbanos")
    _set_font(r_sec, name=FONT_HEADER, size=9, bold=True)
    r_sec.font.color.rgb = COR_LABEL_FONT

    # L3: Endereço e telefone
    pb = c1.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(pb, line=260, before=0, after=0)
    r_info = pb.add_run("Praça XV de Novembro, 127  •  Centro  •  Oliveira/MG  •  (37) 3331-9800")
    _set_font(r_info, name=FONT_DETALHE, size=7)
    r_info.font.color.rgb = COR_CINZA_LEVE

    # ── Linha separadora dupla (fina cinza + grossa azul) ──
    p_thin = hdr.add_paragraph()
    _set_spacing(p_thin, line=20, before=30, after=0)
    pPr_thin = p_thin._p.get_or_add_pPr()
    pb_thin = OxmlElement('w:pBdr')
    bot_thin = OxmlElement('w:bottom')
    bot_thin.set(qn('w:val'),'single'); bot_thin.set(qn('w:sz'),'4')
    bot_thin.set(qn('w:color'), 'AAAAAA'); bot_thin.set(qn('w:space'), '1')
    pb_thin.append(bot_thin); pPr_thin.append(pb_thin)

    p_thick = hdr.add_paragraph()
    _set_spacing(p_thick, line=20, before=4, after=0)
    pPr_thick = p_thick._p.get_or_add_pPr()
    pb_thick = OxmlElement('w:pBdr')
    bot_thick = OxmlElement('w:bottom')
    bot_thick.set(qn('w:val'),'single'); bot_thick.set(qn('w:sz'),'18')
    bot_thick.set(qn('w:color'), COR_INST); bot_thick.set(qn('w:space'), '1')
    pb_thick.append(bot_thick); pPr_thick.append(pb_thick)

# ═══════════════════════════════════════════════════════════
#  IDENTIFICAÇÃO (tabela padronizada — fonte 8pt)
# ═══════════════════════════════════════════════════════════
def _build_identificacao(doc, d):
    W_LABEL = 2268; W_VALUE = 7938
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(tbl)

    linhas = [
        ("Processo nº", f"{d.get('numero_processo','')}  —  {d.get('data_processo','')}" if d.get('numero_processo') else ""),
        ("Assunto",     d.get("assunto","")),
        ("Requerente",  d.get("requerente","")),
    ]
    for i, (rot, val) in enumerate(linhas):
        row = tbl.rows[i]
        c_rot = row.cells[0]
        _apply_label_cell(c_rot, W_LABEL)
        p_rot = c_rot.paragraphs[0]; p_rot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p_rot, line=260, before=50, after=50)
        r_rot = p_rot.add_run(rot)
        _set_font(r_rot, size=SZ_TABELA, bold=True)
        r_rot.font.color.rgb = COR_LABEL_FONT

        c_val = row.cells[1]
        _apply_value_cell(c_val, W_VALUE)
        p_val = c_val.paragraphs[0]; p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p_val, line=260, before=50, after=50)
        r_val = p_val.add_run(val)
        _set_font(r_val, size=SZ_TABELA, bold=(i==2))

    _add_separator(doc)

# ═══════════════════════════════════════════════════════════
#  DADOS DO CARIMBO TÉCNICO (tabela padronizada — fonte 8pt)
# ═══════════════════════════════════════════════════════════
def _parse_numero(texto):
    """Extrai valor numérico de strings como '86,23%' ou '180,00m²' → float."""
    if not texto: return None
    try:
        limpo = re.sub(r'[^\d,.]', '', str(texto)).replace(',', '.')
        return float(limpo) if limpo else None
    except ValueError:
        return None

def _build_dados_carimbo(doc, d):
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p_tit, line=240, before=140, after=40)
    r_tit = _run(p_tit, "DADOS TÉCNICOS DO PROJETO (Ref. Decreto nº 4.149/2019)", bold=True, size=SZ_TABELA)
    r_tit.font.color.rgb = COR_LABEL_FONT

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(tbl)

    logradouro = d.get('logradouro', '')
    bairro = d.get('bairro', '')
    end_completo = f"{logradouro} - {bairro}" if logradouro and bairro else logradouro or bairro

    # Detectar alertas urbanísticos
    to_val = _parse_numero(d.get('taxa_ocupacao', ''))
    perm_val = _parse_numero(d.get('taxa_permeabilidade', ''))
    terreno_val = _parse_numero(d.get('area_terreno', ''))
    alerta_zoneamento = (to_val is not None and to_val > 70) and (perm_val is not None and perm_val < 20)
    alerta_excecao = (terreno_val is not None and terreno_val < 220)

    linhas = [
        ("Endereço:",       end_completo,                          "Inscrição Mun.:", d.get("inscricao_municipal", "")),
        ("Proprietário:",   d.get("proprietario", d.get("requerente", "")), "Desenhista:", d.get("desenhista", "")),
        ("Lote:",           d.get("lote", ""),                     "Quadra:",         d.get("quadra", "")),
        ("Área Terreno:",   d.get("area_terreno", ""),             "Área Total:",     d.get("area_total_construida", "")),
        ("Taxa Ocupação:",  d.get("taxa_ocupacao", ""),            "Coef. Aprov.:",   d.get("coef_aproveitamento", "")),
        ("Permeabilidade:", d.get("taxa_permeabilidade", ""),      "Resp. Técnico:",  d.get("profissional_nome", "")),
    ]

    W_L1 = 2000; W_V1 = 3500; W_L2 = 2000; W_V2 = 2700

    for lab1, val1, lab2, val2 in linhas:
        row = tbl.add_row()
        
        # Label 1
        c0 = row.cells[0]
        _apply_label_cell(c0, W_L1)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_spacing(p0, line=220, before=25, after=25)
        r0 = p0.add_run(lab1)
        _set_font(r0, size=SZ_TABELA, bold=True)
        r0.font.color.rgb = COR_LABEL_FONT
        
        # Valor 1
        c1 = row.cells[1]
        _apply_value_cell(c1, W_V1)
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p1, line=220, before=25, after=25)
        r_v1 = p1.add_run(val1)
        _set_font(r_v1, size=SZ_TABELA)

        # Alerta: Área do Terreno < 220m² → verde + "exceção da lei"
        if lab1 == "Área Terreno:" and alerta_excecao:
            r_v1.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            r_v1.font.bold = True
            r_alert = p1.add_run("  (exceção da lei)")
            _set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        # Alerta: Taxa Ocupação > 70% → vermelho + "conferir zoneamento"
        if lab1 == "Taxa Ocupação:" and alerta_zoneamento:
            r_v1.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            r_v1.font.bold = True
            r_alert = p1.add_run("  (conferir zoneamento)")
            _set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Alerta: Permeabilidade < 20% → vermelho + "conferir zoneamento"
        if lab1 == "Permeabilidade:" and alerta_zoneamento:
            r_v1.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            r_v1.font.bold = True
            r_alert = p1.add_run("  (conferir zoneamento)")
            _set_font(r_alert, size=7, bold=True, italic=True)
            r_alert.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Label 2
        c2 = row.cells[2]
        _apply_label_cell(c2, W_L2)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_spacing(p2, line=220, before=25, after=25)
        r2 = p2.add_run(lab2)
        _set_font(r2, size=SZ_TABELA, bold=True)
        r2.font.color.rgb = COR_LABEL_FONT

        # Valor 2
        c3 = row.cells[3]
        _apply_value_cell(c3, W_V2)
        p3 = c3.paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p3, line=220, before=25, after=25)
        _set_font(p3.add_run(val2), size=SZ_TABELA)

    _add_separator(doc)

# ═══════════════════════════════════════════════════════════
#  LISTA DE DOCUMENTOS (bullet list, sem numeração)
# ═══════════════════════════════════════════════════════════
def _add_doc_item(doc, tipo, obs=None):
    """Adiciona um item de documento como card estilizado com borda lateral azul."""
    # Tabela de 1 linha x 1 coluna para criar efeito de "card" com borda esquerda azul
    card = doc.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Bordas do card: só borda esquerda grossa azul, restante cinza fina
    tblPr = card._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); card._tbl.insert(0, tblPr)
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4'); el.set(qn('w:color'), 'D6DCE4')
        el.set(qn('w:space'), '0'); tblBrd.append(el)
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single'); left_b.set(qn('w:sz'), '18')
    left_b.set(qn('w:color'), COR_INST); left_b.set(qn('w:space'), '0')
    tblBrd.append(left_b)
    for old in tblPr.findall(qn('w:tblBorders')): tblPr.remove(old)
    tblPr.append(tblBrd)

    cell = card.rows[0].cells[0]
    # Fundo levemente cinza
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F8F9FA'); shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

    # Título do documento
    p_tipo = cell.paragraphs[0]
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p_tipo, line=LINE_SPC, before=40, after=20)
    _run(p_tipo, tipo, bold=True, size=9)

    # Observação
    if obs:
        p_obs = cell.add_paragraph()
        p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p_obs, line=260, before=0, after=40)
        r_obs = _run(p_obs, obs, size=9, italic=True)
        r_obs.font.color.rgb = COR_CINZA_TEXTO

    # Espaçamento depois do card
    p_space = doc.add_paragraph()
    _set_spacing(p_space, line=80, before=0, after=0)

# ═══════════════════════════════════════════════════════════
#  CORPO – UNIVERSAL E FLEXÍVEL (fonte 11pt)
# ═══════════════════════════════════════════════════════════
def _body_parecer(doc, d):
    INDENT = 1.25

    if d.get("paragrafo_abertura"):
        p_ab = _para(doc, line=LINE_SPC, before=300, after=PAR_AFTER, indent_cm=INDENT)
        _rich_segments(p_ab, d["paragrafo_abertura"], size=SZ_CORPO)

    if d.get("considerandos"):
        for cons in d["considerandos"]:
            p = _para(doc, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=INDENT)
            _run(p, "Considerando que ", bold=True, size=SZ_CORPO)
            _rich_segments(p, cons, size=SZ_CORPO)

    if d.get("paragrafos_adicionais"):
        for txt in d["paragrafos_adicionais"]:
            p_extra = _para(doc, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=INDENT)
            _rich_segments(p_extra, txt, size=SZ_CORPO)

# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO LEGAL (mesmo tamanho do corpo)
# ═══════════════════════════════════════════════════════════
def _build_fundamentacao(doc, d):
    if not d.get("fundamentacao_legal"):
        return
    
    INDENT = 1.25
    p_fund_tit = _para(doc, line=LINE_SPC, before=200, after=80, indent_cm=INDENT)
    r_tit = _run(p_fund_tit, "Da Análise Legal e Técnica:", bold=True, size=SZ_CORPO)
    r_tit.font.color.rgb = COR_LABEL_FONT
    
    for fund in d["fundamentacao_legal"]:
        p_fund = _para(doc, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=INDENT)
        _run(p_fund, "• ", size=SZ_CORPO)
        _rich_segments(p_fund, fund, size=SZ_CORPO)

# ═══════════════════════════════════════════════════════════
#  FUNDAMENTAÇÃO + CONCLUSÃO + DOCUMENTOS + ASSINATURA
# ═══════════════════════════════════════════════════════════
def _build_conclusao_e_docs(doc, d):
    INDENT = 1.25

    # 1º Fundamentação Legal
    _build_fundamentacao(doc, d)

    # 2º Título "Emissão de Documentos"
    p_doc_tit = _para(doc, line=LINE_SPC, before=200, after=40, indent_cm=INDENT)
    r_dt = _run(p_doc_tit, "Emissão de Documentos:", bold=True, size=SZ_CORPO)
    _set_font(r_dt, name=FONT_TITULO, size=12, bold=True)
    r_dt.font.color.rgb = COR_LABEL_FONT

    # 3º Conclusão ("Diante do exposto...")
    conclusao = d.get("conclusao",
        "Diante do exposto, após análise das documentações apresentadas, conclui-se que o "
        "projeto está apto a ser aprovado quanto às questões técnicas pertinentes, podendo "
        "ser emitidos os seguintes documentos:")
    p_conc = _para(doc, line=LINE_SPC, before=0, after=PAR_AFTER, indent_cm=INDENT)
    _run(p_conc, conclusao, size=SZ_CORPO)

    # 4º Lista de documentos (cards estilizados)
    for item in d.get("documentos_emitir", []):
        _add_doc_item(doc, item.get("tipo",""), item.get("obs") or None)

    # Local e data (centralizado, negrito) — mantém junto com assinatura
    p_local = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                    line=LINE_SPC, before=200, after=0)
    _run(p_local, f"{CIDADE}, {_data_hoje_extenso()}.", bold=True, size=SZ_CORPO)
    p_local.paragraph_format.keep_with_next = True

    # Espaço + assinatura (tudo keep_with_next para não separar)
    for _ in range(2):
        pb = doc.add_paragraph(); _set_spacing(pb, line=LINE_SPC, before=0, after=0)
        pb.paragraph_format.keep_with_next = True

    p_line = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=240, before=0, after=40)
    r_line = p_line.add_run('_' * 45)
    _set_font(r_line, size=10)
    r_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p_line.paragraph_format.keep_with_next = True

    assinatura_items = [
        (ASSINANTE["nome"],     True),
        (ASSINANTE["titulo"],   False),
        (ASSINANTE["registro"], False),
    ]
    for i, (texto, bold) in enumerate(assinatura_items):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=LINE_SPC, before=0, after=0)
        _run(p, texto, bold=bold, size=SZ_CORPO)
        # Não aplicar keep_with_next no último parágrafo
        if i < len(assinatura_items) - 1:
            p.paragraph_format.keep_with_next = True

# ═══════════════════════════════════════════════════════════
#  FUNÇÃO PRINCIPAL E HISTÓRICO
# ═══════════════════════════════════════════════════════════
def gerar_parecer(dados: dict, caminho_saida: str = None) -> str:
    d = dados
    if caminho_saida is None:
        proc = str(d.get("numero_processo", "RELATORIO")).replace("/", "-")
        req_str = str(d.get("requerente", "Desconhecido")).split()[0]
        nome = f"PARECER_{proc} {req_str}.docx"
        
        for ch in ['<','>',':','"','|','?','*']:
            nome = nome.replace(ch, '')
        caminho_saida = os.path.join(PROJECT_DIR, nome)

    doc = Document()
    section = doc.sections[0]
    M = Cm(1.5)
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    section.top_margin = section.bottom_margin = M
    section.left_margin = section.right_margin = M
    section.header_distance = Cm(1.0); section.footer_distance = Cm(1.0)

    style = doc.styles['Normal']
    style.font.name = FONT_CORPO; style.font.size = Pt(SZ_CORPO)

    _build_header(doc)
    _add_page_number_footer(doc)

    # Título — Cambria bold, sem sublinhado
    pt = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=288, before=200, after=140)
    r_tit = _run(pt, "PARECER SETOR TÉCNICO - SMOSU", bold=True, size=14)
    _set_font(r_tit, name=FONT_TITULO, size=14, bold=True)
    r_tit.font.color.rgb = COR_LABEL_FONT

    _build_identificacao(doc, d)
    _build_dados_carimbo(doc, d)
    _body_parecer(doc, d)
    _build_conclusao_e_docs(doc, d)

    doc.save(caminho_saida)
    print(f"[+] Parecer DOCX gerado: {caminho_saida}")

    # ── Geração automática de PDF ──
    caminho_pdf = caminho_saida.replace('.docx', '.pdf')
    try:
        from docx2pdf import convert
        convert(caminho_saida, caminho_pdf)
        print(f"[+] Parecer PDF  gerado: {caminho_pdf}")
    except ImportError:
        try:
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc_w = word.Documents.Open(os.path.abspath(caminho_saida))
            doc_w.SaveAs(os.path.abspath(caminho_pdf), FileFormat=17)
            doc_w.Close()
            word.Quit()
            print(f"[+] Parecer PDF  gerado: {caminho_pdf}")
        except Exception as e:
            print(f"[!] PDF não gerado (instale docx2pdf ou comtypes): {e}")

    return caminho_saida

# ═══════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    args = sys.argv[1:]
    if not args:
        print("Uso: python _engine\\compilador.py dados.json [saida.docx]")
        sys.exit(0)
    
    try:
        with open(args[0], encoding="utf-8") as f:
            dados = json.load(f)
        
        json_str = json.dumps(dados)
        if "⚠️ VERIFICAR" in json_str:
            print("\n[!] AVISO: O JSON contém marcações '⚠️ VERIFICAR'.")
            print("    Recomenda-se corrigir os dados antes da impressão final.\n")

        print(f"[>] Lendo arquivo: {args[0]}")
    except FileNotFoundError:
        print(f"[!] Não encontrado: {args[0]}"); sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"[!] JSON inválido: {e}"); sys.exit(1)
        
    gerar_parecer(dados, args[1] if len(args) > 1 else None)